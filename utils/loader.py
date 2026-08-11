#!/usr/bin/env python


import sys
import io
import os
import re
import getopt
import logging.config
import shutil
import subprocess
import tempfile
import threading
import warnings
import zlib
import zipfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path
from itertools import islice
from typing import Any, Iterable

import ebooklib
from ebooklib import epub
import pypdf
from docx import Document
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import chardet

from backend.es_manager import ESManager
from utils.stat import Stat
from utils.isbn import extract as extract_isbn
from utils.parser_timeout import ParserTimeout, time_limit

logging.config.fileConfig(Path(__file__).parent.parent / "logging.conf", disable_existing_loggers=False)
LOGGER = logging.getLogger()

# pypdf 및 ebooklib 내부 경고/복구 소음 로그 억제
logging.getLogger("pypdf").setLevel(logging.ERROR)
logging.getLogger("pypdf._reader").setLevel(logging.ERROR)
# pypdf._cmap은 미지원 CMap(예: /KSCms-UHC-H, /KSC-EUC-H)을 ERROR 레벨로 페이지마다 찍으므로 별도 차단
logging.getLogger("pypdf._cmap").setLevel(logging.CRITICAL)
# pdfplumber의 백엔드인 pdfminer는 손상 PDF에서 DEBUG 로그를 초당 10만 건 이상 쏟아낸다.
# logging.conf의 root logger가 DEBUG라 그대로 두면 레코드 생성만으로 CPU를 통째로 태운다.
# (실측: 손상 PDF 1건이 13시간 동안 100% CPU를 점유하며 /dev/null로 982GB를 씀)
logging.getLogger("pdfminer").setLevel(logging.WARNING)
warnings.filterwarnings("ignore", message=".*XML.*HTML.*")
warnings.filterwarnings("ignore", message=".*Ascii85.*")
warnings.filterwarnings("ignore", message=".*startxref.*")
with warnings.catch_warnings():
    try:
        warnings.filterwarnings("ignore", category=pypdf.errors.PdfReadWarning)
    except Exception:
        pass

if "TM_BOOK_DIR" not in os.environ:
    LOGGER.error("The environment variable TM_BOOK_DIR is not set.")
    sys.exit(-1)

if "TM_COMICS_DIR" not in os.environ:
    LOGGER.error("The environment variable TM_COMICS_DIR is not set.")
    sys.exit(-1)

HANGUL_RE = re.compile(r"[가-힣]")
# latin-1 상위 영역이 2글자 이상 연속된 구간 = 바이트를 latin-1로 잘못 읽은 흔적
PDF_LATIN1_RUN_RE = re.compile(r"[\u0080-\u00ff]{2,}")
# 상위바이트(0x80~0xFF) 연속 구간. 파이썬 루프 대신 C 속도로 인접 여부를 센다.
HIGH_BYTE_RUN_RE = re.compile(rb"[\x80-\xff]+")


class ProblemCollector(logging.Handler):
    """read_file 실행 중 발생한 WARNING 이상 로그를 모아두는 컨텍스트 매니저.

    콘솔 핸들러를 잠시 떼어 두어, 수집한 메시지를 파일 경로 아래에 묶어서 출력할 수 있게 한다.
    파일 로그(run.log)는 그대로 유지된다.
    """

    _lock = threading.RLock()
    _active_count = 0
    _detached_handlers: list[logging.Handler] = []

    def __init__(self) -> None:
        super().__init__(level=logging.WARNING)
        self.messages: list[str] = []
        self._thread_id: int | None = None

    def emit(self, record: logging.LogRecord) -> None:
        if record.thread == self._thread_id:
            self.messages.append(record.getMessage())

    def __enter__(self) -> "ProblemCollector":
        self.messages.clear()
        self._thread_id = threading.get_ident()
        with ProblemCollector._lock:
            root = logging.getLogger()
            if ProblemCollector._active_count == 0:
                # stdout 핸들러만 분리 (FileHandler 계열도 StreamHandler 하위지만 파일 로그는 유지)
                ProblemCollector._detached_handlers = [
                    h for h in root.handlers if isinstance(h, logging.StreamHandler) and not isinstance(h, logging.FileHandler)
                ]
                for handler in ProblemCollector._detached_handlers:
                    root.removeHandler(handler)
            root.addHandler(self)
            ProblemCollector._active_count += 1
        return self

    def __exit__(self, *_exc_info: Any) -> None:
        with ProblemCollector._lock:
            root = logging.getLogger()
            root.removeHandler(self)
            ProblemCollector._active_count -= 1
            if ProblemCollector._active_count == 0:
                for handler in ProblemCollector._detached_handlers:
                    root.addHandler(handler)
                ProblemCollector._detached_handlers = []
            self._thread_id = None


def report_problems(file_path: Path, messages: list[str], indexed: bool) -> bool:
    """문제가 있을 때만 '* 파일경로'와 그 아래에 사유를 출력. 문제 여부를 반환."""
    # read_from_epub/pdf/rtf 등은 경로와 예외를 각각 로깅하므로 경로 중복 줄은 제거
    problems = [msg for msg in messages if msg != str(file_path)]
    if not indexed:
        if file_path.is_file():
            problems.append(f"적재 제외: 지원하지 않는 파일 형식 '{file_path.suffix}'")
        else:
            problems.append("적재 제외: 파일이 사라짐")
    if not problems:
        return False
    print(f"* {file_path}")
    for msg in problems:
        print(f"    {msg}")
    return True


class Loader:
    TEXT_SIZE = 4096
    # 이 개수 이상의 mojibake 구간이 나오면 추출 실패로 보고 다음 파서로 넘긴다.
    PDF_MOJIBAKE_RUN_LIMIT = 3
    # TEXT_SIZE를 채우는 데 실측 최대 27페이지가 필요했고 첫 텍스트는 최대 4페이지에서
    # 나왔다. 여유를 둔 상한. 스캔본이 끝까지 스캔되는 것을 막는 역할도 한다.
    PDF_PAGE_LIMIT = 40
    # ZIP 컨테이너를 쓰는 포맷. 확장자가 이미 이 계열이면 매직바이트 검사에서
    # 내용물까지 열어보지 않는다(정상 파일의 비용을 0으로 두기 위함).
    ZIP_BASED_TYPES = frozenset({"epub", "cbz", "docx"})
    # 파서 단계 하나의 벽시계 상한. 실측 중앙 소요는 pdfplumber가 0.233s로 가장 느리고,
    # PDF_PAGE_LIMIT 덕에 작업량도 40페이지로 묶여 있어 30초면 정상 파일에는 충분하다.
    # 손상 PDF 1건이 무한히 도는 것을 막는 것이 목적이다(실측 87.7s~무한).
    PDF_STAGE_TIMEOUT = 30
    _pdftotext_path_checked = False
    _pdftotext_path: str | None = None
    # txt 디코딩 시도 순서. read_from_text와 reencode_text_file_to_utf8이 공유한다.
    TEXT_ENCODINGS = ["utf-8", "cp949", "euc-kr", "utf-16", "utf-16-le", "utf-16-be", "utf-8-sig"]
    # 서양어 단일바이트 인코딩(cp1252) 판정 기준.
    # cp949/euc-kr의 한글은 항상 상위바이트 2개가 연속이고, 서양어 악센트/구두점은
    # ASCII 사이에 고립된 단일 상위바이트로 나타난다. 실측 인접쌍 비율:
    # cp1252 영문 0.0000~0.0024 vs 한글(cp949/euc-kr/미상) 1.0000.
    TEXT_SINGLE_BYTE_ENCODING = "cp1252"
    TEXT_HIGH_BYTE_PAIRED_MIN = 0.5
    # chardet에 넘길 최대 바이트 수. 실측 1.4MB 전체는 2.08초, 64KB는 186ms, 8KB는 16ms.
    # 한국어 인코딩 판별에는 8KB로 충분하고, 어차피 후보 순서만 정한다.
    TEXT_DETECT_PREFIX = 8 * 1024
    # --reencode / --reencode-dry-run 으로 설정된다. 기본은 파일을 건드리지 않는다.
    reencode_txt_mode = False
    reencode_txt_dry_run = False
    # --reencode-backup-dir 로 설정. 지정하면 덮어쓰기 전에 원본을 이 아래에 복사하고,
    # 복사·검증이 실패하면 그 파일은 변환하지 않는다.
    reencode_backup_dir: Path | None = None
    path_prefix = Path(os.environ["TM_BOOK_DIR"])
    comics_path_prefix = Path(os.environ["TM_COMICS_DIR"])

    @staticmethod
    def get_path_prefix(file_path: Path) -> Path:
        """파일 경로에 해당하는 path_prefix를 반환"""
        if file_path.is_relative_to(Loader.comics_path_prefix):
            return Loader.comics_path_prefix
        return Loader.path_prefix

    @staticmethod
    def reencode_text_file_to_utf8(file_path: Path, dry_run: bool = False) -> tuple[bool, str]:
        """txt 파일 자체를 UTF-8로 재인코딩해 저장. 반환: (변경 여부, 사유).

        검증을 모두 통과하지 못하면 파일에 손대지 않는다.
          1. 바이트 왕복: 디코딩한 문자열을 원본 인코딩으로 되돌려 인코딩했을 때
             원본 바이트와 완전히 일치해야 한다(무손실·가역 확정).
          2. 타당성: 왕복만으로는 인코딩 선택이 옳은지 알 수 없다. BOM 없는 UTF-16은
             LE/BE 양쪽이 모두 왕복을 통과하므로, 한글/영숫자가 실제로 나오는지 본다.
          3. 글자수: 재인코딩 결과를 다시 읽어 원본 디코딩 결과와 글자수·내용이 일치해야 한다.

        ES 문서 _id가 inode이므로 임시 파일 교체(os.replace)를 쓰지 않고 같은 inode에
        덮어쓴다. 중간에 죽어도 복구할 수 있도록 검증된 내용을 먼저 임시 파일에 써 둔다.
        """
        try:
            raw = file_path.read_bytes()
        except OSError as e:
            return False, f"읽기 실패: {e}"

        if not raw:
            return False, "빈 파일"
        try:
            raw.decode("utf-8")
            return False, "이미 UTF-8"
        except UnicodeDecodeError:
            pass

        # chardet 라벨은 TEXT_ENCODINGS에 있는 것만 받아들인다. windows-1252 같은
        # 단일바이트 인코딩은 어떤 바이트열이든 디코딩되고 왕복도 무조건 통과하므로,
        # 후보로 허용하면 한글 파일에 mojibake를 영구 기록하게 된다.
        order = list(Loader.TEXT_ENCODINGS)
        try:
            # 전체를 넘기면 1.4MB당 2초가 걸린다(실측). chardet은 후보 순서만 정하고
            # 모든 후보는 전체 파일 왕복·타당성 검증을 통과해야 하므로 prefix로 충분하다.
            detected = chardet.detect(raw[: Loader.TEXT_DETECT_PREFIX])
            enc = (detected.get("encoding") or "") if detected else ""
            conf = (detected.get("confidence") or 0) if detected else 0
            allowed = {e.lower(): e for e in Loader.TEXT_ENCODINGS}
            if enc and conf >= 0.7 and enc.lower() in allowed:
                pick = allowed[enc.lower()]
                order = [pick] + [e for e in order if e != pick]
        except Exception:
            pass

        # UTF-16 후보는 근거가 있을 때만 시도한다. 길이가 짝수인 바이트열은 utf-16-be/le로
        # 거의 항상 디코딩되고 왕복까지 통과하므로, 근거 없이 허용하면 단일바이트 텍스트가
        # CJK 글자 뭉치로 바뀐다. 실제 UTF-16 텍스트는 BOM이 있거나 NUL 바이트가 많다.
        # 실측: BOM 없는 한글 UTF-16LE는 NUL 비율 0.15, 단일바이트 텍스트는 0.0000.
        has_bom = raw[:2] in (b"\xff\xfe", b"\xfe\xff")
        nul_ratio = raw.count(0) / len(raw)
        if not has_bom and nul_ratio < 0.02:
            order = [e for e in order if not e.startswith("utf-16")]

        # 상위바이트가 고립되어 나타나면 서양어 단일바이트 인코딩이다. 이때 cp949/euc-kr을
        # 남겨두면 상위바이트가 뒤따르는 ASCII까지 2바이트 시퀀스로 삼켜 영문자를 파괴한다
        # (예: 'Don Sabas—a man' → 'Don Sabas뾞 man', 'a' 소실). 그래서 대체가 아니라 교체한다.
        runs = HIGH_BYTE_RUN_RE.findall(raw)
        high_total = sum(len(r) for r in runs)
        if high_total:
            # 상위바이트가 길이 2 이상 구간에 속하면 '인접'이다(파이썬 루프와 동일한 정의).
            paired = sum(len(r) for r in runs if len(r) >= 2) / high_total
            if paired < Loader.TEXT_HIGH_BYTE_PAIRED_MIN:
                order = [e for e in order if e not in ("cp949", "euc-kr")] + [Loader.TEXT_SINGLE_BYTE_ENCODING]

        for source_enc in order:
            try:
                text = raw.decode(source_enc)
            except (UnicodeDecodeError, UnicodeError, LookupError):
                continue
            if "\ufffd" in text:
                continue
            # 1. 바이트 왕복 — 무손실·가역 확정
            try:
                if text.encode(source_enc) != raw:
                    continue
            except (UnicodeEncodeError, LookupError):
                continue
            # 2. 타당성 — 왕복은 무손실만 증명하고 인코딩 선택이 옳은지는 증명하지 못한다.
            #    (BOM 없는 UTF-16은 LE/BE 양쪽이 모두 왕복을 통과한다.)
            #    실제 글자가 있어야 하고, mojibake로 읽힌 결과여서는 안 된다.
            if not re.search(r"[\uac00-\ud7a3a-zA-Z0-9\u3400-\u9fff]", text):
                continue
            if Loader._count_mojibake_runs(text) >= Loader.PDF_MOJIBAKE_RUN_LIMIT:
                continue
            # 텍스트 파일에는 줄바꿈이 있다. 엔디언을 잘못 고른 UTF-16은 \n(0x000A)과
            # \r(0x000D)이 U+0A00/U+0D00으로 바뀌어 줄바꿈이 하나도 남지 않는다.
            if len(text) >= 200 and "\n" not in text and "\r" not in text:
                continue
            sample = text[: Loader.TEXT_SIZE]
            if sample and sum(1 for c in sample if "\u0080" <= c <= "\u00ff") / len(sample) > 0.1:
                continue
            break
        else:
            return False, "무손실 디코딩 가능한 인코딩을 찾지 못함"

        utf8_bytes = text.encode("utf-8")
        # 3. 글자수·내용 검증 — 새 바이트를 다시 읽었을 때 원본 디코딩 결과와 같아야 한다
        roundtrip = utf8_bytes.decode("utf-8")
        if len(roundtrip) != len(text):
            return False, f"글자수 불일치: {len(text)} → {len(roundtrip)}"
        if roundtrip != text:
            return False, "재인코딩 내용 불일치"

        if dry_run:
            return True, f"{source_enc} → utf-8 ({len(text)}자, {len(raw)}B → {len(utf8_bytes)}B) [dry-run]"

        # 원본 백업. 실패하면 되돌릴 수단이 없으므로 변환을 포기한다.
        if Loader.reencode_backup_dir is not None:
            try:
                prefix = Loader.get_path_prefix(file_path)
                rel = file_path.relative_to(prefix)
            except ValueError:
                rel = Path(file_path.name)
            backup_path = Loader.reencode_backup_dir / rel
            try:
                backup_path.parent.mkdir(parents=True, exist_ok=True)
                backup_path.write_bytes(raw)
                if backup_path.read_bytes() != raw:
                    return False, f"백업 검증 실패: {backup_path}"
            except OSError as e:
                return False, f"백업 실패: {e}"

        # 검증된 내용을 먼저 임시 파일에 보존한 뒤, 같은 inode에 덮어쓴다.
        tmp_path = file_path.with_name(file_path.name + ".utf8.tmp")
        try:
            tmp_path.write_bytes(utf8_bytes)
            with file_path.open("r+b") as outfile:
                outfile.write(utf8_bytes)
                outfile.truncate()
                outfile.flush()
                os.fsync(outfile.fileno())
        except OSError as e:
            return False, f"쓰기 실패: {e}"
        finally:
            tmp_path.unlink(missing_ok=True)

        # 저장된 파일을 실제로 다시 읽어 확인한다(중간 신호가 아니라 산출물로 확인).
        try:
            saved = file_path.read_bytes().decode("utf-8")
        except (OSError, UnicodeDecodeError) as e:
            return False, f"저장 후 재읽기 실패: {e}"
        if len(saved) != len(text) or saved != text:
            return False, f"저장 후 불일치: {len(text)}자 → {len(saved)}자"
        return True, f"{source_enc} → utf-8 ({len(text)}자, {len(raw)}B → {len(utf8_bytes)}B)"

    @staticmethod
    def read_from_text(file_path: Path) -> tuple[str, int, int, str]:
        """TXT 파일 읽기. 반환: (summary, line_count, page_count, raw_content)
        utf-8 -> cp949 -> euc-kr -> utf-16 -> utf-16-le -> utf-16-be -> utf-8-sig 순서로 인코딩 fallback 시도.
        chardet 감지를 활용하고 모든 인코딩 엄격 디코딩 실패 시 errors="replace" 옵션으로 재시도.
        """
        Stat.add("text_count", 1)
        start_time = datetime.now()

        line_count = 0
        data = ""
        raw_content = ""
        encodings = list(Loader.TEXT_ENCODINGS)
        replace_encodings = list(Loader.TEXT_ENCODINGS)
        last_exception = None

        try:
            raw_bytes = file_path.read_bytes()
            detected = chardet.detect(raw_bytes)
            detected_enc = detected.get("encoding") if detected else None
            confidence = detected.get("confidence", 0) if detected else 0
            if detected_enc and confidence >= 0.7 and detected_enc.lower() not in ["ascii", "iso-8859-1"]:
                encodings = [detected_enc] + [e for e in encodings if e.lower() != detected_enc.lower()]
                replace_encodings = [detected_enc] + [e for e in replace_encodings if e.lower() != detected_enc.lower()]
        except Exception:
            pass

        for enc in encodings:
            try:
                with file_path.open("r", encoding=enc) as infile:
                    content = infile.read()
                    # utf-16도 포함해야 한다. BOM 없는 단일바이트 파일이 utf-16으로
                    # 디코딩되면 CJK 글자 뭉치가 되는데, 빠뜨리면 검사 없이 채택된다.
                    if enc in ["cp949", "euc-kr", "utf-16", "utf-16-le", "utf-16-be"] and content:
                        if not re.search(r"[가-힣a-zA-Z0-9]", content) and not content.isspace():
                            continue
                    raw_content = content
                    break
            except (UnicodeDecodeError, UnicodeError) as e:
                last_exception = e
        else:
            for enc in replace_encodings:
                try:
                    with file_path.open("r", encoding=enc, errors="replace") as infile:
                        content = infile.read()
                        if enc in ["cp949", "euc-kr", "utf-16", "utf-16-le", "utf-16-be"] and content:
                            if not re.search(r"[가-힣a-zA-Z0-9]", content) and not content.isspace():
                                continue
                        raw_content = content
                        if raw_content:
                            break
                except Exception:
                    pass
            else:
                if last_exception:
                    LOGGER.error(f"can't read unicode text from file '{file_path}', {last_exception}")
                raw_content = ""

        if raw_content:
            line_count = raw_content.count("\n") + 1
            data = raw_content[: Loader.TEXT_SIZE]
            data = data.replace("\ufeff", "")
            data = re.sub(r"[^a-zA-Z0-9\sㄱ-힣]", " ", data)
            data = re.sub(r"\s+", " ", data).strip()

        end_time = datetime.now()
        Stat.add("text_total_time", (end_time - start_time).total_seconds())

        return data, line_count, 0, raw_content

    @staticmethod
    def read_from_epub_with_extracting_zip(file_path: Path) -> tuple[str, int]:
        """메모리에서 직접 zip 내용을 읽어 처리 (임시 디렉토리 사용 안 함)"""
        result = ""
        total_text = ""
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                # 1. container.xml에서 rootfile 찾기
                container_content = zf.read("META-INF/container.xml").decode("utf-8", errors="ignore")
                m = re.search(r'<rootfile[^>]*full-path="(?P<root_file>[^"]+)"', container_content)
                if not m:
                    return "", 0
                root_file = m.group("root_file")
                root_dir = "/".join(root_file.split("/")[:-1])

                # 2. OPF 파일에서 chapter 파일 목록 찾기
                opf_content = zf.read(root_file).decode("utf-8", errors="ignore")
                matches = re.findall(r'<(?:opf:)?item\s[^>]*href="(?P<chapter_file>[^"]*\.x?html)"[^>]*media-type="application/xhtml\+xml"', opf_content)

                # 3. 각 chapter 파일 읽기
                for chapter_file in matches:
                    chapter_path = f"{root_dir}/{chapter_file}" if root_dir else chapter_file
                    try:
                        content = zf.read(chapter_path).decode("utf-8", errors="ignore")
                        with warnings.catch_warnings():
                            warnings.filterwarnings("ignore", message=".*XML.*HTML.*")
                            soup = BeautifulSoup(content, "lxml")
                        text = soup.get_text()
                        total_text += text
                        if len(result) < Loader.TEXT_SIZE:
                            result += text
                    except Exception:
                        continue
        except zipfile.BadZipFile as e:
            LOGGER.error(file_path)
            LOGGER.error(e)

        line_count = total_text.count("\n") + 1 if total_text else 0
        return result[: Loader.TEXT_SIZE], line_count

    @staticmethod
    def read_from_epub(file_path: Path) -> tuple[str, int, int]:
        Stat.add("normal_epub_count", 1)
        start_time = datetime.now()

        result = ""
        total_text = ""
        line_count = 0
        try:
            book = epub.read_epub(file_path)
            titles = book.get_metadata("DC", "title")
            if titles:
                for title in titles:
                    if title[0]:
                        result += " " + title[0]
            creators = book.get_metadata("DC", "creator")
            if creators:
                for creator in creators:
                    if creator[0]:
                        result += " " + creator[0]

            for doc in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                with warnings.catch_warnings():
                    warnings.filterwarnings("ignore", message=".*XML.*HTML.*")
                    soup = BeautifulSoup(doc.get_body_content(), "lxml")
                text = soup.get_text()
                total_text += text
                if len(result) < Loader.TEXT_SIZE:
                    result += text

            line_count = total_text.count("\n") + 1 if total_text else 0

            end_time = datetime.now()
            Stat.add("normal_epub_total_time", (end_time - start_time).total_seconds())
        except Exception as e:
            Stat.add("normal_epub_count", -1)
            Stat.add("zipped_epub_count", 1)
            start_time = datetime.now()

            fallback_failed = False
            try:
                result, line_count = Loader.read_from_epub_with_extracting_zip(file_path)
            except Exception as e2:
                fallback_failed = True
                LOGGER.error(file_path)
                LOGGER.error(e2)

            if fallback_failed:
                LOGGER.error(file_path)
                LOGGER.error(e)

            end_time = datetime.now()
            Stat.add("zipped_epub_total_time", (end_time - start_time).total_seconds())

        result = re.sub(r"[^\w\sㄱ-힣]", " ", result)

        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def _is_usable_pdf_text(text: str, mojibake_runs: int | None = None) -> bool:
        """추출 결과를 그대로 써도 되는지 판정. 비어 있거나 mojibake면 False.

        EUC-KR 바이트를 latin-1로 잘못 해석한 텍스트는 비어 있지 않으므로,
        빈 문자열 검사만으로는 fallback이 발동하지 않고 깨진 텍스트가 그대로 적재된다.
        mojibake_runs를 넘기면 같은 텍스트를 두 번 스캔하지 않는다.
        """
        if not text[: Loader.TEXT_SIZE].strip():
            return False
        # 라틴 문자권 도서의 악센트 문자가 우연히 한글로 복원되는 경우가 있어 여유를 둔다.
        runs = Loader._count_mojibake_runs(text) if mojibake_runs is None else mojibake_runs
        return runs < Loader.PDF_MOJIBAKE_RUN_LIMIT

    @staticmethod
    def _count_mojibake_runs(text: str) -> int:
        """EUC-KR 바이트를 latin-1로 잘못 읽은 구간의 개수.

        전체 대비 비율로 판정하면 4096자 중 155자(3.8%)처럼 부분 손상을 놓치므로
        구간 개수로 센다. 각 구간은 실제 재해석으로 확정한다.
        """
        runs = 0
        for run in PDF_LATIN1_RUN_RE.findall(text[: Loader.TEXT_SIZE]):
            try:
                decoded = run.encode("latin-1").decode("euc-kr", "ignore")
            except Exception:
                continue
            if decoded and len(HANGUL_RE.findall(decoded)) / len(decoded) >= 0.5:
                runs += 1
        return runs

    @staticmethod
    def _zip_container_type(file_path: Path) -> str | None:
        """ZIP 컨테이너의 실제 종류를 내용물로 판별한다."""
        try:
            with zipfile.ZipFile(file_path) as zf:
                names = zf.namelist()
                if any(n.endswith(".opf") for n in names) or "META-INF/container.xml" in names:
                    return "epub"
                if any(n.startswith("word/") for n in names):
                    return "docx"
                image_suffixes = (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp")
                entries = [n for n in names if not n.endswith("/")]
                if entries and all(n.lower().endswith(image_suffixes) for n in entries):
                    return "cbz"
        except Exception:
            return None
        return None

    @staticmethod
    def detect_file_type(file_path: Path, declared: str) -> str | None:
        """매직바이트로 실제 포맷을 판별한다. 판별 불가면 None.

        확장자를 믿을 수 없는 파일이 실제로 존재한다(실측: 21,150건 중 EPUB 5건이
        .pdf 확장자를 달고 있어 PDF 파서로 가 전부 실패했다).

        비용을 아끼려고, 확장자와 매직바이트가 이미 같은 계열이면 더 파고들지 않는다.
        정상 파일에서는 앞 4바이트만 읽고 끝난다.
        """
        try:
            with file_path.open("rb") as f:
                head = f.read(4)
        except OSError:
            return None

        if head == b"%PDF":
            return "pdf"
        if head == b"PK\x03\x04":
            # 이미 zip 기반 포맷으로 선언되어 있으면 내용물까지 열어보지 않는다
            if declared in Loader.ZIP_BASED_TYPES:
                return declared
            return Loader._zip_container_type(file_path)
        return None

    @staticmethod
    def _pdfium_text(source: Any) -> tuple[str, int]:
        """pypdfium2로 PDF_PAGE_LIMIT까지 훑어 TEXT_SIZE만큼 텍스트를 모은다.

        source는 파일 객체나 BytesIO. 경로를 넘기면 pdfium이 fd를 쥐는데 문서 로드
        실패 시 그 fd가 닫히지 않고 GC로도 회수되지 않으므로(실측: 손상 PDF마다
        fd 1개씩 영구 누적), 호출자가 핸들을 소유해서 넘긴다.
        """
        import pypdfium2

        pdf = None
        try:
            pdf = pypdfium2.PdfDocument(source)
            page_count = len(pdf)
            texts: list[str] = []
            total = 0
            for i in range(min(page_count, Loader.PDF_PAGE_LIMIT)):
                page = None
                textpage = None
                try:
                    page = pdf[i]
                    textpage = page.get_textpage()
                    text = textpage.get_text_range() or ""
                finally:
                    close_textpage = getattr(textpage, "close", None)
                    if close_textpage:
                        close_textpage()
                    close_page = getattr(page, "close", None)
                    if close_page:
                        close_page()
                texts.append(text)
                total += len(text)
                if total >= Loader.TEXT_SIZE:
                    break
            return "".join(texts), page_count
        finally:
            if pdf is not None:
                pdf.close()

    @staticmethod
    def _pdf_text_by_pypdfium2(file_path: Path) -> tuple[str, int]:
        with file_path.open("rb") as fp:
            return Loader._pdfium_text(fp)

    @staticmethod
    def _pdf_text_by_pikepdf_repair(file_path: Path) -> tuple[str, int]:
        """xref/page tree가 깨진 PDF를 pikepdf(=qpdf 바인딩)로 재구성한 뒤 다시 추출한다.

        구조가 깨졌을 뿐 객체는 살아 있는 파일이 실제로 존재한다(실측: 9MB 손상본이
        4.0초 만에 470쪽으로 복구됨). 반대로 객체 자체가 소실된 파일은 여기서도
        복구되지 않고 빠르게 실패한다(실측: 63MB 손상본 5.6초, RSS 29MB).

        복구본은 디스크에 쓰지 않고 메모리에만 둔다. pikepdf 기본 설정은 압축을
        유지하므로 복구본이 원본보다 크게 부풀지 않는다(9MB -> 8.6MB).
        """
        import pikepdf

        buffer = io.BytesIO()
        with file_path.open("rb") as fp, pikepdf.open(fp) as pdf:
            pdf.save(buffer)
        buffer.seek(0)
        return Loader._pdfium_text(buffer)

    @staticmethod
    def _pdf_text_by_pdfplumber(file_path: Path) -> tuple[str, int]:
        import pdfplumber

        # 파일 핸들을 직접 소유해서, pdfplumber.open()이 열기 도중 실패하거나
        # 타임아웃이 끼어들어도 fd가 GC까지 남지 않게 한다.
        with file_path.open("rb") as fp, pdfplumber.open(fp) as pdf:
            page_count = len(pdf.pages)
            texts: list[str] = []
            total = 0
            for page in pdf.pages[: Loader.PDF_PAGE_LIMIT]:
                text = page.extract_text() or ""
                texts.append(text)
                total += len(text)
                # 페이지 객체가 캐시한 문자/도형 정보를 즉시 버려 메모리 누적을 막는다.
                page.close()
                if total >= Loader.TEXT_SIZE:
                    break
            return "".join(texts), page_count

    @staticmethod
    def _pdf_text_by_pdftotext(file_path: Path) -> tuple[str, int]:
        if not Loader._pdftotext_path_checked:
            Loader._pdftotext_path = shutil.which("pdftotext")
            Loader._pdftotext_path_checked = True
        pdftotext_path = Loader._pdftotext_path
        if not pdftotext_path:
            return "", 0
        # -enc UTF-8로 출력을 고정하고, locale에 의존하지 않도록 직접 디코딩한다.
        cmd = [pdftotext_path, "-enc", "UTF-8", "-l", str(Loader.PDF_PAGE_LIMIT), str(file_path), "-"]
        proc = subprocess.run(cmd, capture_output=True, timeout=10)
        if proc.returncode != 0:
            return "", 0
        return proc.stdout.decode("utf-8", "replace"), 0

    @staticmethod
    def read_from_pdf(file_path: Path) -> tuple[str, int, int]:
        Stat.add("pdf_count", 1)
        start_time = datetime.now()

        result = ""
        page_count = 0
        # 어느 파서도 온전한 텍스트를 못 주면 그중 손상이 가장 적은 결과를 쓴다.
        best_damage: int | None = None

        # 기본 경로는 in-process 파서만 쓴다. pdftotext는 빠르지만 외부 process spawn이고,
        # 실측 부분 손상 복구력은 pdfplumber가 더 높아 기본 fallback에서는 제외한다.
        stages = (Loader._pdf_text_by_pypdfium2, Loader._pdf_text_by_pdfplumber)
        timed_out_stages: list[str] = []
        for stage_no, extract in enumerate(stages):
            stage_name = getattr(extract, "__name__", str(extract))
            stage_start = datetime.now()
            stage_outcome = "error"
            try:
                with time_limit(Loader.PDF_STAGE_TIMEOUT, stage_name):
                    text, pages = extract(file_path)
                if page_count == 0:
                    page_count = pages
                if not text.strip():
                    stage_outcome = "empty"
                    # 1단계가 문서를 열었는데도 텍스트가 없으면 텍스트 레이어가 없는 스캔본이다.
                    # 이 경우 다른 파서도 결과가 없음을 확인했으므로(표본 28건 전수) 더 시도하지 않는다.
                    # 스캔본은 전체 PDF의 43.8%(9,261건, 1,545,537페이지)라 이 분기가 비용을 크게 줄인다.
                    if stage_no == 0 and page_count > 0:
                        break
                    continue
                damage = Loader._count_mojibake_runs(text)
                if Loader._is_usable_pdf_text(text, damage):
                    stage_outcome = "ok"
                    result = text
                    break
                stage_outcome = "damaged"
                if best_damage is None or damage < best_damage:
                    best_damage, result = damage, text
            except ParserTimeout:
                stage_outcome = "timeout"
                # 타임아웃은 조용히 넘기지 않고 파일별 문제로 보고한다.
                # (path는 report_problems가 헤더로 출력하므로 메시지에 넣지 않는다)
                timed_out_stages.append(stage_name)
                LOGGER.error(file_path)
                LOGGER.error("PDF 파서 타임아웃(%d초 초과): %s", Loader.PDF_STAGE_TIMEOUT, stage_name)
                continue
            except Exception:
                stage_outcome = "error"
                continue
            finally:
                Stat.record_pdf_stage(stage_name, (datetime.now() - stage_start).total_seconds(), stage_outcome)

        # 어느 파서도 문서를 열지 못했다면(page_count == 0) 파서 선택 문제가 아니라
        # xref/page tree 구조 손상이다. pikepdf로 재구성한 뒤 한 번 더 시도한다.
        # 텍스트가 조금이라도 나온 경우(mojibake 등)는 구조 문제가 아니므로 건너뛴다.
        # 실측 손상률은 표본 300건 중 1건(0.3%)이라 이 분기가 도는 일 자체가 드물다.
        if not result.strip() and page_count == 0:
            repair_stage = "_pdf_text_by_pikepdf_repair"
            repair_start = datetime.now()
            repair_outcome = "error"
            try:
                with time_limit(Loader.PDF_STAGE_TIMEOUT, repair_stage):
                    text, pages = Loader._pdf_text_by_pikepdf_repair(file_path)
                page_count = pages
                if Loader._is_usable_pdf_text(text):
                    repair_outcome = "ok"
                    result = text
                    LOGGER.warning("PDF 구조 복구 후 적재 (%d쪽)", pages)
                elif text.strip():
                    repair_outcome = "damaged"
                else:
                    repair_outcome = "empty"
            except ParserTimeout:
                repair_outcome = "timeout"
                timed_out_stages.append(repair_stage)
                LOGGER.error(file_path)
                LOGGER.error("PDF 구조 복구 타임아웃(%d초 초과)", Loader.PDF_STAGE_TIMEOUT)
            except Exception as e:
                repair_outcome = "error"
                LOGGER.error(file_path)
                LOGGER.error("PDF 구조 복구 실패(객체 소실 추정): %s", str(e)[:120])
            finally:
                Stat.record_pdf_stage(repair_stage, (datetime.now() - repair_start).total_seconds(), repair_outcome)

        if not result.strip() and timed_out_stages:
            LOGGER.error(file_path)
            LOGGER.error("PDF 텍스트 추출 실패: 모든 파서가 타임아웃 또는 오류 (타임아웃: %s)", ", ".join(timed_out_stages))
        elif not result.strip() and page_count == 0:
            LOGGER.error(file_path)

        result = re.sub(r"[^\w\sㄱ-힣]", " ", result)

        end_time = datetime.now()
        Stat.add("pdf_total_time", (end_time - start_time).total_seconds())

        return result[: Loader.TEXT_SIZE], 0, page_count

    @staticmethod
    def read_from_html(file_path: Path) -> tuple[str, int, int]:
        Stat.add("html_count", 1)
        start_time = datetime.now()

        content = ""
        line_count = 0
        with file_path.open("r") as infile:
            content = infile.read()
            line_count = content.count("\n") + 1

        # XMLParsedAsHTMLWarning 억제하고 lxml 파서 사용
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message=".*XML.*HTML.*")
            soup = BeautifulSoup(content, "lxml")
        result = soup.get_text()
        result = re.sub(r"[^\w\sㄱ-힣]", " ", result)

        end_time = datetime.now()
        Stat.add("html_total_time", (end_time - start_time).total_seconds())

        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def read_from_docx(file_path: Path) -> tuple[str, int, int]:
        Stat.add("docx_count", 1)
        start_time = datetime.now()

        result = ""
        total_text = ""
        doc = Document(str(file_path))
        for paragraph in doc.paragraphs:
            text = paragraph.text
            total_text += text + "\n"
            if len(result) < Loader.TEXT_SIZE:
                result += text
        result = re.sub(r"[^\w\sㄱ-힣]", " ", result)
        line_count = total_text.count("\n") if total_text else 0

        end_time = datetime.now()
        Stat.add("docx_total_time", (end_time - start_time).total_seconds())

        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def read_from_rtf(file_path: Path) -> tuple[str, int, int]:
        Stat.add("rtf_count", 1)
        start_time = datetime.now()

        result = ""
        line_count = 0
        try:
            with file_path.open("rb") as infile:
                raw_data = infile.read()
                doc = raw_data.decode("utf-8")
                result = rtf_to_text(doc, errors="ignore")
                line_count = result.count("\n") + 1 if result else 0
        except Exception as e:
            LOGGER.error(file_path)
            LOGGER.error(e)
        result = re.sub(r"[^\w\sㄱ-힣]", " ", result)

        end_time = datetime.now()
        Stat.add("rtf_total_time", (end_time - start_time).total_seconds())

        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def _find_libreoffice() -> str:
        """LibreOffice 실행 파일 경로를 반환"""
        for cmd in ["libreoffice", "soffice"]:
            path = shutil.which(cmd)
            if path:
                return path
        mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if Path(mac_path).exists():
            return mac_path
        return "libreoffice"

    @staticmethod
    def _convert_with_libreoffice(file_path: Path, output_format: str) -> str:
        """LibreOffice를 사용하여 파일을 변환하고 결과 텍스트를 반환"""
        lo_bin = Loader._find_libreoffice()
        with tempfile.TemporaryDirectory() as tmpdir:
            proc = subprocess.run([lo_bin, "--headless", "--convert-to", output_format, "--outdir", tmpdir, str(file_path)], capture_output=True, timeout=60)
            ext = output_format.split(":")[0]
            out_file = Path(tmpdir) / (file_path.stem + "." + ext)
            if out_file.exists():
                return out_file.read_text(encoding="utf-8", errors="replace")
            # stem 불일치 시 글로브로 검색
            out_files = list(Path(tmpdir).glob(f"*.{ext}"))
            if out_files:
                return out_files[0].read_text(encoding="utf-8", errors="replace")
            # 변환 결과 없음 — 진단 로그
            all_files = list(Path(tmpdir).iterdir())
            LOGGER.error("LibreOffice produced no output: file='%s', format='%s', returncode=%d, stderr=%s, tmpdir_files=%s", file_path, output_format, proc.returncode, proc.stderr.decode("utf-8", errors="replace")[:500], [f.name for f in all_files])
        return ""

    @staticmethod
    def read_from_doc(file_path: Path) -> tuple[str, int, int]:
        Stat.add("doc_count", 1)
        start_time = datetime.now()
        result = ""
        line_count = 0
        try:
            raw_text = Loader._convert_with_libreoffice(file_path, "txt:Text")
            line_count = raw_text.count("\n") + 1 if raw_text else 0
            result = raw_text[: Loader.TEXT_SIZE]
            result = re.sub(r"[^\w\sㄱ-힣]", " ", result)
        except Exception as e:
            LOGGER.error(f"can't read doc file '{file_path}': {e}")
        end_time = datetime.now()
        Stat.add("doc_total_time", (end_time - start_time).total_seconds())
        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def read_from_hwp(file_path: Path) -> tuple[str, int, int]:
        Stat.add("hwp_count", 1)
        start_time = datetime.now()
        result = ""
        line_count = 0
        try:
            raw_text = Loader._convert_with_libreoffice(file_path, "txt:Text")
            if not raw_text.strip():
                # LibreOffice 변환 실패 — 네이티브 HWP3 파서로 fallback
                from utils.hwp3_parser import extract_text_from_hwp3

                raw_text = extract_text_from_hwp3(file_path)
            line_count = raw_text.count("\n") + 1 if raw_text else 0
            result = raw_text[: Loader.TEXT_SIZE]
            result = re.sub(r"[^\w\sㄱ-힣]", " ", result)
        except Exception as e:
            LOGGER.error(f"can't read hwp file '{file_path}': {e}")
        end_time = datetime.now()
        Stat.add("hwp_total_time", (end_time - start_time).total_seconds())
        return result[: Loader.TEXT_SIZE], line_count, 0

    @staticmethod
    def read_from_image(_file_path: Path) -> tuple[str, int, int]:
        Stat.add("image_count", 1)
        # 이미지는 line_count, page_count 해당 없음
        return "", 0, 0

    @staticmethod
    def _find_xref_offset(xref_data: bytes, obj_num: int) -> int | None:
        """traditional xref 테이블에서 특정 object의 파일 내 offset을 찾는다."""
        xref_text = xref_data.replace(b"\r\n", b"\n").replace(b"\r", b"\n")
        lines = xref_text.split(b"\n")

        i = 0
        while i < len(lines) and lines[i].strip() != b"xref":
            i += 1
        i += 1  # 'xref' 건너뛰기

        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if line.startswith(b"trailer"):
                break
            parts = line.split()
            if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                sub_start = int(parts[0])
                sub_count = int(parts[1])
                if sub_start <= obj_num < sub_start + sub_count:
                    entry_idx = i + 1 + (obj_num - sub_start)
                    if entry_idx < len(lines):
                        entry_parts = lines[entry_idx].strip().split()
                        if len(entry_parts) >= 3 and entry_parts[2] == b"n":
                            return int(entry_parts[0])
                    return None
                i += 1 + sub_count
                continue
            i += 1
        return None

    @staticmethod
    def _apply_png_predictor(data: bytes, columns: int) -> bytes:
        """PNG predictor 해제 (xref stream 디코딩용). Sub/Up 필터 지원."""
        stride = 1 + columns  # filter byte + row data
        rows = len(data) // stride
        result = bytearray()
        prev_row = bytes(columns)
        for i in range(rows):
            offset = i * stride
            fb = data[offset]
            row = bytearray(data[offset + 1 : offset + 1 + columns])
            if fb == 1:  # Sub
                for j in range(1, columns):
                    row[j] = (row[j] + row[j - 1]) & 0xFF
            elif fb == 2:  # Up
                for j in range(columns):
                    row[j] = (row[j] + prev_row[j]) & 0xFF
            result.extend(row)
            prev_row = bytes(row)
        return bytes(result)

    @staticmethod
    def _xref_stream_find_entry(data: bytes, w: list[int], index_ranges: list[int], obj_num: int) -> tuple[int, int, int] | None:
        """xref stream에서 특정 object의 entry를 찾는다.
        반환: (type, field2, field3) 또는 None.
        type 0: free, type 1: (1, file_offset, gen), type 2: (2, obj_stream_num, index)"""
        w1, w2, w3 = w
        entry_size = w1 + w2 + w3
        if entry_size == 0:
            return None
        pos = 0
        for i in range(0, len(index_ranges) - 1, 2):
            start_obj = index_ranges[i]
            count = index_ranges[i + 1]
            for j in range(count):
                if pos + entry_size > len(data):
                    return None
                if start_obj + j == obj_num:
                    entry = data[pos : pos + entry_size]
                    type_val = int.from_bytes(entry[:w1], "big") if w1 > 0 else 1
                    field2 = int.from_bytes(entry[w1 : w1 + w2], "big") if w2 > 0 else 0
                    field3 = int.from_bytes(entry[w1 + w2 : w1 + w2 + w3], "big") if w3 > 0 else 0
                    return (type_val, field2, field3)
                pos += entry_size
        return None

    @staticmethod
    def _read_from_obj_stream(f, stream_offset: int, target_obj_num: int) -> bytes | None:
        """Object stream에서 특정 object의 딕셔너리 데이터를 읽는다."""
        f.seek(stream_offset)
        header = f.read(4096)

        first_match = re.search(rb"/First\s+(\d+)", header)
        length_match = re.search(rb"/Length\s+(\d+)", header)
        if not first_match or not length_match:
            return None
        if re.search(rb"/Length\s+\d+\s+\d+\s+R", header):
            return None

        first_offset = int(first_match.group(1))
        stream_length = int(length_match.group(1))

        stream_start = re.search(rb"stream\r?\n", header)
        if not stream_start:
            return None

        f.seek(stream_offset + stream_start.end())
        raw = f.read(stream_length)

        filter_match = re.search(rb"/Filter\s*/(\w+)", header)
        if filter_match:
            if filter_match.group(1) != b"FlateDecode":
                return None
            try:
                raw = zlib.decompress(raw)
            except zlib.error:
                return None

        # 헤더 파싱: "obj_num offset obj_num offset ..."
        header_text = raw[:first_offset].decode("ascii", errors="replace")
        tokens = header_text.split()

        target_offset = None
        next_offset = len(raw)
        for i in range(0, len(tokens) - 1, 2):
            if int(tokens[i]) == target_obj_num:
                target_offset = int(tokens[i + 1]) + first_offset
                if i + 3 < len(tokens):
                    next_offset = int(tokens[i + 3]) + first_offset
                break

        if target_offset is None:
            return None
        return raw[target_offset:next_offset]

    @staticmethod
    def _parse_one_xref_stream(f, xref_offset: int) -> tuple | None:
        """단일 xref stream을 파싱하여 (decompressed, w, index_ranges, prev_offset)를 반환.
        실패 시 None."""
        f.seek(xref_offset)
        stream_obj_data = f.read(4096)

        w_match = re.search(rb"/W\s*\[\s*(\d+)\s+(\d+)\s+(\d+)\s*\]", stream_obj_data)
        if not w_match:
            return None
        w = [int(w_match.group(i)) for i in (1, 2, 3)]

        size_match = re.search(rb"/Size\s+(\d+)", stream_obj_data)
        if not size_match:
            return None
        xref_size = int(size_match.group(1))

        index_match = re.search(rb"/Index\s*\[([^\]]+)\]", stream_obj_data)
        index_ranges = [int(x) for x in index_match.group(1).split()] if index_match else [0, xref_size]

        # /Length가 indirect reference이면 포기
        if re.search(rb"/Length\s+\d+\s+\d+\s+R", stream_obj_data):
            return None
        length_match = re.search(rb"/Length\s+(\d+)", stream_obj_data)
        if not length_match:
            return None
        stream_length = int(length_match.group(1))

        stream_start = re.search(rb"stream\r?\n", stream_obj_data)
        if not stream_start:
            return None

        f.seek(xref_offset + stream_start.end())
        raw_stream = f.read(stream_length)

        # 필터 처리
        filter_match = re.search(rb"/Filter\s*/(\w+)", stream_obj_data)
        if filter_match:
            if filter_match.group(1) != b"FlateDecode":
                return None
            try:
                decompressed = zlib.decompress(raw_stream)
            except zlib.error:
                return None
        else:
            decompressed = raw_stream

        # PNG predictor 처리
        parms_match = re.search(rb"/DecodeParms\s*<<([^>]*)>>", stream_obj_data)
        if parms_match:
            pred_match = re.search(rb"/Predictor\s+(\d+)", parms_match.group(1))
            if pred_match and int(pred_match.group(1)) >= 10:
                cols_match = re.search(rb"/Columns\s+(\d+)", parms_match.group(1))
                columns = int(cols_match.group(1)) if cols_match else sum(w)
                decompressed = Loader._apply_png_predictor(decompressed, columns)

        prev_match = re.search(rb"/Prev\s+(\d+)", stream_obj_data)
        prev_offset = int(prev_match.group(1)) if prev_match else None

        return decompressed, w, index_ranges, prev_offset

    @staticmethod
    def _fast_pdf_page_count(file_path: Path) -> int | None:
        """PDF xref/trailer에서 페이지 수만 경량 추출 (전체 파싱 없이).

        trailer → /Root → /Pages → /Count 경로를 따라 필요한 객체만 읽는다.
        traditional xref table과 xref stream을 모두 지원하며,
        /Prev 체인을 따라 이전 xref 섹션도 탐색한다.
        실패 시 None 반환 (pypdf fallback용).
        """
        try:
            with file_path.open("rb") as f:
                f.seek(0, 2)
                file_size = f.tell()

                # 1. 파일 끝에서 startxref 찾기
                tail_size = min(4096, file_size)
                f.seek(-tail_size, 2)
                tail = f.read()

                m = re.search(rb"startxref\s+(\d+)", tail)
                if not m:
                    return None

                # 2. xref 체인을 따라가며 lookup 함수들 수집
                # lookups: [('traditional', fn) | ('stream', fn)] newest→oldest
                root_obj_num = None
                lookups: list[tuple[str, Any]] = []
                xref_offset: int | None = int(m.group(1))
                max_depth = 10

                while xref_offset is not None and max_depth > 0:
                    max_depth -= 1
                    f.seek(xref_offset)
                    xref_header = f.read(256)

                    if xref_header.lstrip().startswith(b"xref"):
                        # === Traditional xref table ===
                        header_match = re.search(rb"xref\s+\d+\s+(\d+)", xref_header)
                        entry_count = int(header_match.group(1)) if header_match else 5000
                        xref_read_size = min(entry_count * 20 + 4096, file_size - xref_offset)
                        f.seek(xref_offset)
                        xref_data = f.read(xref_read_size)

                        if root_obj_num is None:
                            root_match = re.search(rb"/Root\s+(\d+)\s+\d+\s+R", xref_data)
                            if root_match:
                                root_obj_num = int(root_match.group(1))

                        lookups.append(("traditional", lambda obj_num, xd=xref_data: Loader._find_xref_offset(xd, obj_num)))

                        prev_match = re.search(rb"/Prev\s+(\d+)", xref_data)
                        xref_offset = int(prev_match.group(1)) if prev_match else None
                    else:
                        # === Xref stream ===
                        f.seek(xref_offset)
                        peek = f.read(512)

                        if root_obj_num is None:
                            root_match = re.search(rb"/Root\s+(\d+)\s+\d+\s+R", peek)
                            if root_match:
                                root_obj_num = int(root_match.group(1))

                        parsed = Loader._parse_one_xref_stream(f, xref_offset)
                        if parsed is None:
                            break
                        decompressed, w, index_ranges, prev_offset = parsed
                        lookups.append(("stream", lambda obj_num, d=decompressed, ww=list(w), ir=list(index_ranges): Loader._xref_stream_find_entry(d, ww, ir, obj_num)))
                        xref_offset = prev_offset

                if root_obj_num is None:
                    return None

                def find_file_offset(obj_num: int) -> int | None:
                    """type 1 (직접 저장) object의 파일 offset을 찾는다."""
                    for kind, lookup in lookups:
                        if kind == "traditional":
                            offset = lookup(obj_num)
                            if offset is not None:
                                return offset
                        else:
                            entry = lookup(obj_num)
                            if entry is not None and entry[0] == 1:
                                return entry[1]
                    return None

                def read_object_data(obj_num: int, read_size: int = 4096) -> bytes | None:
                    """object 데이터를 읽는다 (type 1: 직접, type 2: object stream)."""
                    # type 1 시도
                    offset = find_file_offset(obj_num)
                    if offset is not None:
                        f.seek(offset)
                        return f.read(read_size)
                    # type 2 시도 (object stream에 압축 저장)
                    for kind, lookup in lookups:
                        if kind != "stream":
                            continue
                        entry = lookup(obj_num)
                        if entry is not None and entry[0] == 2:
                            stream_offset = find_file_offset(entry[1])
                            if stream_offset is not None:
                                return Loader._read_from_obj_stream(f, stream_offset, obj_num)
                    return None

                # 3. Root object → /Pages 참조 찾기
                root_data = read_object_data(root_obj_num, 1024)
                if root_data is None:
                    return None
                pages_match = re.search(rb"/Pages\s+(\d+)\s+\d+\s+R", root_data)
                if not pages_match:
                    return None
                pages_obj_num = int(pages_match.group(1))

                # 4. Pages object → /Count 추출
                pages_data = read_object_data(pages_obj_num)
                if pages_data is None:
                    return None
                count_match = re.search(rb"/Count\s+(\d+)", pages_data)
                if count_match:
                    return int(count_match.group(1))
        except Exception:
            pass
        return None

    @staticmethod
    def read_file(file_path: Path, stat_result: os.stat_result | None = None, skip_text: bool = False) -> dict[int, dict[str, Any]]:
        if file_path.is_file():
            sys.stdout.flush()
            # read metadata of each file (stat 결과 재사용)
            st = stat_result if stat_result else file_path.stat()
            inode_num = st.st_ino
            file_size = st.st_size
            prefix = Loader.get_path_prefix(file_path)
            category = str(file_path.parent.relative_to(prefix))
            if category == ".":
                category = "_root"
            m = re.search(r"^\[(?P<author>[^\]]+)\]\s*(?P<title>.+)$", file_path.stem)
            if m:
                author = m.group("author")
                title = m.group("title")
            else:
                author = ""
                title = file_path.stem

            file_type = file_path.suffix[1:].lower()

            # 확장자가 실제 포맷과 다르면 실제 포맷 기준으로 파서를 고른다.
            # 확장자만 믿으면 EPUB이 .pdf를 달고 있을 때 PDF 파서로 가 반드시 실패한다.
            detected_type = Loader.detect_file_type(file_path, file_type)
            if detected_type and detected_type != file_type:
                LOGGER.warning("확장자와 실제 포맷 불일치: .%s 이지만 %s로 처리", file_type, detected_type)
                file_type = detected_type

            # skip_text: 텍스트 추출 건너뛰기 (만화 등 이미지 기반 파일)
            if skip_text:
                summary = ""
                line_count = 0
                page_count = 0
                isbn_list = []
                # PDF만 예외: page_count 추출 유지 (경량 방식 우선, 실패 시 pypdf fallback)
                if file_type == "pdf":
                    page_count = Loader._fast_pdf_page_count(file_path)
                    if page_count is None:
                        page_count = 0
                        try:
                            # 손상 PDF는 pypdf 안에서도 무한히 돌 수 있어 상한을 건다.
                            with time_limit(Loader.PDF_STAGE_TIMEOUT, "pypdf/page_count"), file_path.open("rb") as f:
                                page_count = len(pypdf.PdfReader(f).pages)
                        except ParserTimeout:
                            LOGGER.error(file_path)
                            LOGGER.error("PDF 페이지 수 추출 타임아웃(%d초 초과): pypdf", Loader.PDF_STAGE_TIMEOUT)
                        except Exception as e:
                            LOGGER.error(file_path)
                            LOGGER.error(e)
                    Stat.add("pdf_count", 1)
                # 지원하지 않는 확장자는 기존 동작 유지 (빈 dict 반환)
                supported_types = {"txt", "epub", "pdf", "docx", "doc", "hwp", "rtf", "html", "jpg", "jpeg", "png", "gif", "webp", "bmp", "tiff", "svg", "cbz"}
                if file_type not in supported_types:
                    return {}

                return {
                    inode_num: {"category": category, "title": title, "author": author, "file_path": str(file_path.relative_to(prefix)), "file_type": file_type, "file_size": int(file_size), "line_count": line_count, "page_count": page_count, "isbn": "", "summary": summary, "updated_time": datetime.now().isoformat()}
                }

            # read content of each file
            summary = ""
            line_count = 0
            page_count = 0
            isbn_list = []
            raw_content = ""  # TXT 파일용 원본 content
            if file_type == "txt":
                # --reencode 지정 시, 적재 전에 파일 자체를 UTF-8로 바꿔 둔다.
                # 검증을 통과하지 못하면 파일에 손대지 않으므로 적재 동작에는 영향이 없다.
                if Loader.reencode_txt_mode:
                    changed, reason = Loader.reencode_text_file_to_utf8(file_path, dry_run=Loader.reencode_txt_dry_run)
                    if changed:
                        Stat.add("text_reencoded_count", 1)
                        LOGGER.warning(f"UTF-8 재인코딩: {reason}")
                    elif reason not in ("이미 UTF-8", "빈 파일"):
                        LOGGER.warning(f"UTF-8 재인코딩 건너뜀: {reason}")
                summary, line_count, page_count, raw_content = Loader.read_from_text(file_path)
            elif file_type == "epub":
                summary, line_count, page_count = Loader.read_from_epub(file_path)
            elif file_type == "pdf":
                if file_path.is_relative_to(Loader.comics_path_prefix):
                    # comics: 텍스트 추출은 생략하되 페이지 수만 추출 (경량 방식 우선)
                    page_count = Loader._fast_pdf_page_count(file_path)
                    if page_count is None:
                        page_count = 0
                        try:
                            # 손상 PDF는 pypdf 안에서도 무한히 돌 수 있어 상한을 건다.
                            with time_limit(Loader.PDF_STAGE_TIMEOUT, "pypdf/page_count"), file_path.open("rb") as f:
                                page_count = len(pypdf.PdfReader(f).pages)
                        except ParserTimeout:
                            LOGGER.error(file_path)
                            LOGGER.error("PDF 페이지 수 추출 타임아웃(%d초 초과): pypdf", Loader.PDF_STAGE_TIMEOUT)
                        except Exception as e:
                            LOGGER.error(file_path)
                            LOGGER.error(e)
                    summary, line_count = "", 0
                    Stat.add("pdf_count", 1)
                else:
                    summary, line_count, page_count = Loader.read_from_pdf(file_path)
            elif file_type == "docx":
                summary, line_count, page_count = Loader.read_from_docx(file_path)
            elif file_type == "doc":
                summary, line_count, page_count = Loader.read_from_doc(file_path)
            elif file_type == "hwp":
                summary, line_count, page_count = Loader.read_from_hwp(file_path)
            elif file_type == "rtf":
                summary, line_count, page_count = Loader.read_from_rtf(file_path)
            elif file_type == "html":
                summary, line_count, page_count = Loader.read_from_html(file_path)
            elif file_type in ("jpg", "jpeg", "png", "gif", "webp", "bmp", "tiff", "svg"):
                summary, line_count, page_count = Loader.read_from_image(file_path)
            elif file_type == "cbz":
                summary, line_count, page_count = Loader.read_from_image(file_path)
            else:
                return {}

            # ISBN 추출 (TXT는 이미 읽은 content 재사용)
            if file_type in ("txt", "epub", "pdf", "djvu", "hwp"):
                isbn_content = raw_content if file_type == "txt" else summary if file_type == "pdf" else None
                isbn_list = extract_isbn(file_path, content=isbn_content)

            return {
                inode_num: {
                    "category": category,
                    "title": title,
                    "author": author,
                    "file_path": str(file_path.relative_to(prefix)),
                    "file_type": file_type,
                    "file_size": int(file_size),
                    "line_count": line_count,
                    "page_count": page_count,
                    "isbn": isbn_list[0] if isbn_list else "",
                    "summary": summary,
                    "updated_time": datetime.now().isoformat(),
                }
            }

        return {}

    @staticmethod
    def get_file_list(path: Path, num_files: int = sys.maxsize, recursive: bool = False) -> list[Path]:
        """파일 목록을 반환 (파싱 없이)

        recursive=False인 경우:
        1. 하위 디렉토리 각각에서 첫 번째 파일 1개씩
        2. 지정된 디렉토리에 바로 속한 파일들
        """
        if path.is_dir():
            if recursive:
                file_path_list = [p for p in path.rglob("*") if p.is_file() and not any(part.startswith(".") for part in p.relative_to(path).parts)]
            else:
                file_path_list = []
                # 1. 하위 디렉토리 각각에서 첫 번째 파일 1개씩
                for subdir in path.iterdir():
                    if subdir.is_dir() and not subdir.name.startswith("."):
                        first_file = next((p for p in subdir.iterdir() if p.is_file() and not p.name.startswith(".")), None)
                        if first_file:
                            file_path_list.append(first_file)
                # 2. 지정된 디렉토리에 바로 속한 파일들
                file_path_list.extend(p for p in path.iterdir() if p.is_file() and not p.name.startswith("."))
        else:
            file_path_list = [path]
        return file_path_list[:num_files]

    @staticmethod
    def get_stat(file_path: Path) -> os.stat_result:
        """파일의 stat 결과를 반환"""
        return file_path.stat()

    @staticmethod
    def read_files(path: Path, num_files: int = sys.maxsize, recursive: bool = False) -> dict[int, dict[str, Any]]:
        """파일들을 읽어서 데이터 딕셔너리로 반환 (테스트 및 하위 호환성용)"""
        data: dict[int, dict[str, Any]] = {}
        file_list = Loader.get_file_list(path, num_files, recursive)
        problem_collector = ProblemCollector()

        for child_path in file_list:
            with problem_collector as collector:
                data_item = Loader.read_file(child_path)
            report_problems(child_path, collector.messages, bool(data_item))
            if data_item:
                data.update(data_item)

        return data


def print_usage(program_name: str):
    print(f"Usage:\t{program_name}\t[ --delete ] [ --reload ] [ --recursive ] [ --reencode | --reencode-dry-run ] <index_name> [file or directory path ...]")
    print("\t\tindex_name: book | comics")
    print("\t\t--delete: delete index and exit (no file path required)")
    print("\t\t--reload: force reload even if file already exists in ES")
    print("\t\t--recursive: scan subdirectories recursively")
    print("\t\t--reencode: rewrite non-UTF-8 txt files as UTF-8 in place (verified; skipped if any check fails)")
    print("\t\t--reencode-dry-run: report what --reencode would change without writing")
    print("\t\t--reencode-backup-dir=DIR: copy each original under DIR before overwriting")
    print("\t\t        (a file is left untouched if its backup cannot be written and verified)")
    print("\t\t  Note: needs --reload (or an explicit file path). Files already in ES are otherwise skipped")
    print("\t\t        without being parsed, so re-encoding would never run for them.")
    print()
    print("\t\tNote: When a file (not directory) is specified, it will be force-reloaded automatically.")
    print("\t\tNote: --recursive --reload also deletes ES records under the path whose files no longer exist (orphans).")
    sys.exit(0)


def main() -> int:
    BATCH_SIZE = 100
    WORKER_COUNT = 2

    do_delete = False
    do_reload = False
    do_recursive = False
    args: list[str] = []
    try:
        opts, args = getopt.getopt(sys.argv[1:], "", ["delete", "reload", "recursive", "reencode", "reencode-dry-run", "reencode-backup-dir="])
        for opt, optarg in opts:
            if opt == "--delete":
                do_delete = True
            elif opt == "--reload":
                do_reload = True
            elif opt == "--recursive":
                do_recursive = True
            elif opt == "--reencode":
                Loader.reencode_txt_mode = True
            elif opt == "--reencode-dry-run":
                Loader.reencode_txt_mode = True
                Loader.reencode_txt_dry_run = True
            elif opt == "--reencode-backup-dir":
                Loader.reencode_backup_dir = Path(optarg)
    except getopt.GetoptError as e:
        LOGGER.error(e)
        print_usage(sys.argv[0])

    # 인덱스명 파싱
    INDEX_MAP = {"book": os.environ["TM_ES_BOOK_INDEX"], "comics": os.environ["TM_ES_COMICS_INDEX"]}

    if len(args) < 1:
        print_usage(sys.argv[0])
        return 1

    index_name_arg = args[0]
    if index_name_arg not in INDEX_MAP:
        LOGGER.error(f"유효하지 않은 인덱스명: '{index_name_arg}' (book 또는 comics만 가능)")
        print_usage(sys.argv[0])
        return 1

    file_args = args[1:]

    if not do_delete and len(file_args) < 1:
        print_usage(sys.argv[0])
        return 1

    start_time: datetime = datetime.now()

    # ESManager 초기화 (한 번의 실행에서 하나의 인덱스만 사용)
    idx = INDEX_MAP[index_name_arg]
    es_manager = ESManager(index_name=idx)
    try:
        if not es_manager.es.ping():
            LOGGER.error("Elasticsearch 서버에 연결할 수 없습니다.")
            sys.exit(-1)
    except Exception as e:
        LOGGER.error(f"Elasticsearch 접속 실패: {e}")
        sys.exit(-1)
    es_manager.create_index()
    print(f"ES 인덱스: {idx}")

    if do_delete:
        try:
            if es_manager.do_exist_index():
                es_manager.delete_index()
                print(f"인덱스 '{es_manager.index_name}' 삭제 완료")
            else:
                print(f"인덱스 '{es_manager.index_name}'가 존재하지 않습니다")
        except Exception as e:
            LOGGER.error(e)
            return -1
        return 0

    def process_file_iter(
        file_iter: Iterable[Path],
        skip_check: bool = False,
        skip_text: bool = False,
        seen_inodes: set[int] | None = None,
    ) -> tuple[int, int, int]:
        """파일 iterator를 배치 처리하여 ES에 저장. 반환: (처리 수, 건너뜀 수, 경로동기화 수)"""
        skipped_count = 0
        processed_count = 0
        synced_count = 0
        file_iterator = iter(file_iter)

        def can_parse_in_worker(file_path: Path) -> bool:
            declared_type = file_path.suffix[1:].lower()
            if declared_type == "pdf":
                return False
            return Loader.detect_file_type(file_path, declared_type) != "pdf"

        def parse_file(inode: int) -> tuple[Path, list[str], dict[int, dict[str, Any]]]:
            file_path, st = file_stat_map[inode]
            with ProblemCollector() as collector:
                data_item = Loader.read_file(file_path, stat_result=st, skip_text=skip_text)
            return file_path, list(collector.messages), data_item

        while True:
            # generator에서 배치 단위로 가져오기
            batch_files = list(islice(file_iterator, BATCH_SIZE))
            if not batch_files:
                break

            # stat 수집 (한 번만 호출하여 inode와 함께 저장)
            file_stat_map: dict[int, tuple[Path, os.stat_result]] = {}
            for file_path in batch_files:
                try:
                    st = Loader.get_stat(file_path)
                    file_stat_map[st.st_ino] = (file_path, st)
                except OSError:
                    continue

            # 디스크에 존재하는(=live) inode 수집: orphan 판정 기준 (파싱 성공 여부 무관)
            if seen_inodes is not None:
                seen_inodes.update(file_stat_map.keys())

            if skip_check:
                # --reload 모드: 존재 여부 무시, 전부 파싱
                new_inodes = set(file_stat_map.keys())
                path_changed_inodes: set[int] = set()
                existing_paths: dict[int, str] = {}
            else:
                # 기존 경로 조회 (inode → file_path)
                existing_paths = es_manager.get_existing_paths(list(file_stat_map.keys()))

                # 경로 변경 감지: ES에 있고 file_path가 다른 경우 → 재적재 대상
                path_changed_inodes: set[int] = set()
                for inode, es_file_path in existing_paths.items():
                    if inode not in file_stat_map:
                        continue
                    file_path, _ = file_stat_map[inode]
                    prefix = Loader.get_path_prefix(file_path)
                    current_file_path = str(file_path.relative_to(prefix))
                    if current_file_path != es_file_path:
                        path_changed_inodes.add(inode)
                        print(f"  [경로 변경 감지] inode={inode}: {es_file_path} → {current_file_path}")

                # 신규 파일 + 경로 변경 파일: 전체 재적재 대상
                new_inodes = (set(file_stat_map.keys()) - set(existing_paths.keys())) | path_changed_inodes
                skipped_count += len(existing_paths) - len(path_changed_inodes)
                synced_count += len(path_changed_inodes)

            # 파일 파싱 (stat 결과 재사용). 문제가 있는 파일만 경로와 사유를 출력
            batch_data: dict[int, dict[str, Any]] = {}
            new_inode_list = list(new_inodes)
            worker_inodes = [
                inode for inode in new_inode_list if can_parse_in_worker(file_stat_map[inode][0])
            ]
            worker_inode_set = set(worker_inodes)
            main_thread_inodes = [inode for inode in new_inode_list if inode not in worker_inode_set]
            parse_results: dict[
                int,
                tuple[Path, list[str], dict[int, dict[str, Any]]],
            ] = {}

            if worker_inodes:
                with ThreadPoolExecutor(max_workers=WORKER_COUNT) as executor:
                    futures = {inode: executor.submit(parse_file, inode) for inode in worker_inodes}
                    for inode in main_thread_inodes:
                        parse_results[inode] = parse_file(inode)
                    for inode, future in futures.items():
                        parse_results[inode] = future.result()
            else:
                for inode in main_thread_inodes:
                    parse_results[inode] = parse_file(inode)

            for inode in new_inode_list:
                file_path, messages, data_item = parse_results[inode]
                report_problems(file_path, messages, bool(data_item))
                if data_item:
                    batch_data.update(data_item)

            # 데이터 저장 (중복 file_path 문서 선제거 후 insert)
            if batch_data:
                new_file_paths = [v["file_path"] for v in batch_data.values() if "file_path" in v]
                new_ids = list(batch_data.keys())
                cleaned = es_manager.delete_by_file_paths(new_file_paths, exclude_ids=new_ids)
                if cleaned > 0:
                    print(f"  [중복 제거: {cleaned}개 기존 문서 삭제]")
                es_manager.insert(batch_data)
                processed_count += len(batch_data)
                Stat.add("index_count", len(batch_data))
                if skip_check:
                    print(f"  [배치 저장: {len(batch_data)}개]")
                else:
                    synced_msg = f", 경로변경 재적재: {len(path_changed_inodes)}개" if path_changed_inodes else ""
                    skip_batch = len(existing_paths) - len(path_changed_inodes)
                    print(f"  [배치 저장: {len(batch_data)}개, 건너뜀: {skip_batch}개{synced_msg}]")

        return processed_count, skipped_count, synced_count

    for arg in file_args:
        target_path = Path(arg)
        if not target_path.exists():
            LOGGER.error("can't find such a file or directory '%s'", target_path)
            return 0
        if not target_path.is_relative_to(Loader.path_prefix) and not target_path.is_relative_to(Loader.comics_path_prefix):
            LOGGER.error(f"{target_path} is not in $TM_BOOK_DIR({Loader.path_prefix}) or $TM_COMICS_DIR({Loader.comics_path_prefix}).")
            continue
        print(f"====== {target_path} ======")

        # comics 인덱스는 텍스트 추출 건너뛰기
        skip_text = index_name_arg == "comics"

        # 파일이 지정된 경우: 강제 재적재 (skip_check=True)
        if target_path.is_file():
            print(f"  [파일 강제 재적재] {target_path.name}")
            processed, _, _ = process_file_iter([target_path], skip_check=True, skip_text=skip_text)
            if processed > 0:
                print("  파일 재적재 완료")
            else:
                print("  파일 적재 실패 (지원하지 않는 형식일 수 있음)")
        elif do_recursive:
            # 전체 파일 등록 (generator 사용으로 메모리 효율화, hidden directory 제외)
            file_iter = (p for p in target_path.rglob("*") if p.is_file() and not any(part.startswith(".") for part in p.relative_to(target_path).parts))
            skip_check = do_reload
            # --reload 시 이번 실행에서 본 live inode를 수집하여 orphan 삭제에 사용
            seen_inodes: set[int] | None = set() if do_reload else None
            processed, skipped_count, synced_count = process_file_iter(file_iter, skip_check=skip_check, skip_text=skip_text, seen_inodes=seen_inodes)
            print(f"  총 {processed}개 파일 처리됨")
            if skipped_count > 0:
                print(f"  총 {skipped_count}개 중복 파일 건너뜀")
            if synced_count > 0:
                print(f"  총 {synced_count}개 경로 동기화")

            # --reload: 대상 경로 하위에서 디스크에 더 이상 없는 orphan 레코드 삭제
            if seen_inodes is not None:
                prefix = Loader.get_path_prefix(target_path)
                rel = str(target_path.relative_to(prefix))
                es_manager.refresh()  # 방금 insert한 문서가 scroll에 반영되도록
                existing_ids = es_manager.get_doc_ids_by_path_prefix(rel)
                orphan_ids = list(existing_ids - seen_inodes)
                if orphan_ids:
                    deleted = es_manager.delete_by_ids(orphan_ids)
                    print(f"  [orphan 삭제: {deleted}개] (디스크에서 사라진 예전 레코드, 경로: {rel})")
                else:
                    print("  orphan 없음 (삭제할 예전 레코드 없음)")
        else:
            skip_check = do_reload

            if len(file_args) == 1:
                # 1단계: 하위 디렉토리 각각에서 첫 번째 파일 1개씩 (모아서 한꺼번에 저장)
                print("  [1단계] 하위 디렉토리별 샘플 파일 등록")
                sample_files: list[tuple[str, Path]] = []  # (subdir_name, file_path)
                for subdir in target_path.iterdir():
                    if subdir.is_dir() and not subdir.name.startswith("."):
                        # 첫 번째 파일만 가져옴 (정렬 불필요, iterator 사용)
                        first_file = next((p for p in subdir.iterdir() if p.is_file() and not p.name.startswith(".")), None)
                        if first_file:
                            sample_files.append((subdir.name, first_file))
                        else:
                            print(f"    {subdir.name}/ -> (파일 없음)")

                if sample_files:
                    print(f"    {len(sample_files)}개 디렉토리에서 샘플 파일 선정 완료")
                    for subdir_name, sample_file in sample_files:
                        print(f"      {subdir_name}/ -> {sample_file.name}")

                    # 모아서 한꺼번에 ES에 저장
                    sample_file_paths = [f for _, f in sample_files]
                    processed, skipped1, synced1 = process_file_iter(sample_file_paths, skip_check=skip_check, skip_text=skip_text)
                    print(f"    {processed}개 카테고리 샘플 저장, {skipped1}개 중복 건너뜀")
                    if synced1 > 0:
                        print(f"    {synced1}개 경로 동기화")

                print("  [2단계] 현재 디렉토리 파일 등록")
            else:
                print("  [지정 디렉토리 파일 등록]")

            current_dir_files = [p for p in target_path.iterdir() if p.is_file() and not p.name.startswith(".")]
            print(f"    {len(current_dir_files)}개 파일 발견")
            _, skipped2, synced2 = process_file_iter(current_dir_files, skip_check=skip_check, skip_text=skip_text)
            if skipped2 > 0:
                print(f"    {skipped2}개 중복 파일 건너뜀")
            if synced2 > 0:
                print(f"    {synced2}개 경로 동기화")

        print("================================")

    # 모든 작업 완료 후 인덱스 refresh
    es_manager.refresh()

    end_time = datetime.now()
    Stat.set_value("index_total_time", (end_time - start_time).total_seconds())
    Stat.print()

    return 0


if __name__ == "__main__":
    sys.exit(main())
