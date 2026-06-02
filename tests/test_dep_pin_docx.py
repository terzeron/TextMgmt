#!/usr/bin/env python
"""python-docx dependency pinning.

utils 사용처:
- utils/loader.py:23,236-238  DOCX 본문 추출
    from docx import Document
    doc = Document(str(file_path))
    for paragraph in doc.paragraphs:
        text = paragraph.text

박제 API:
- docx.Document(str_path) -> document
- document.paragraphs -> 반복 가능
- paragraph.text -> str
"""

import tempfile
import unittest
from pathlib import Path


class TestPythonDocxDependencyPinning(unittest.TestCase):
    def test_document_importable(self):
        """from docx import Document"""
        from docx import Document

        self.assertTrue(callable(Document))

    def test_read_paragraphs_text(self):
        """loader.py: Document(str(path)) -> doc.paragraphs -> paragraph.text"""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = Path(f.name)
        try:
            doc = Document()
            doc.add_paragraph("First Line")
            doc.add_paragraph("둘째 줄")
            doc.save(str(docx_path))

            doc2 = Document(str(docx_path))
            texts = [paragraph.text for paragraph in doc2.paragraphs]
            self.assertIn("First Line", texts)
            self.assertIn("둘째 줄", texts)
        finally:
            docx_path.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
