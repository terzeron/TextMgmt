#!/usr/bin/env python
"""의존성 모듈 API 호환성 테스트

dependabot 등에 의한 의존성 업데이트 후에도
응용 코드가 사용하는 API가 정상 동작하는지 검증한다.
외부 서비스(DB, ES, 네트워크)가 필요 없는 단위 테스트만 포함한다.
"""

import io
import tempfile
import unittest
from pathlib import Path


# ---------------------------------------------------------------------------
# 1. beautifulsoup4 (bs4)
#    사용처: loader.py, bookstore.py, isbn.py, book_manager.py, main.py
# ---------------------------------------------------------------------------
class TestBeautifulSoup4(unittest.TestCase):
    """BeautifulSoup 파서별 생성, 탐색, 텍스트 추출 API 테스트"""

    HTML = """
    <html><body>
      <h2 class="gd_name">Test Title</h2>
      <span class="gd_auth"><a href="/author/1">Author Name</a></span>
      <ul id="ulCategory">
        <li><a href="/cat/1">Category 1</a></li>
        <li><a href="/cat/2">Category 2</a></li>
      </ul>
      <img src="cover.jpg" />
      <image xlink:href="cover.svg" />
      <link rel="stylesheet" href="style.css" />
    </body></html>
    """

    def _assert_common(self, soup):
        from bs4 import BeautifulSoup

        self.assertIsInstance(soup, BeautifulSoup)

        # find() - 단일 요소
        h2 = soup.find("h2", class_="gd_name")
        self.assertIsNotNone(h2)
        self.assertEqual(h2.get_text(strip=True), "Test Title")

        # find_all() - 다수 요소
        li_list = soup.find_all("li")
        self.assertEqual(len(li_list), 2)

        # select() - CSS 선택자
        links = soup.select("#ulCategory a")
        self.assertEqual(len(links), 2)
        self.assertEqual(links[0]["href"], "/cat/1")

        # 중첩 find
        span = soup.find("span", class_="gd_auth")
        a_tag = span.find("a")
        self.assertEqual(a_tag.get_text(), "Author Name")
        self.assertTrue(a_tag.get("href"))

        # find_all with lambda
        imgs = soup.find_all("img")
        self.assertGreaterEqual(len(imgs), 1)

        # get() - 속성 접근
        image = soup.find("image")
        self.assertEqual(image.get("xlink:href"), "cover.svg")

        # get_text() - 전체 텍스트 추출
        text = soup.get_text()
        self.assertIn("Test Title", text)

    def test_html_parser(self):
        """html.parser 백엔드로 파싱 (bookstore.py, main.py 패턴)"""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(self.HTML, "html.parser")
        self._assert_common(soup)

    def test_lxml_parser(self):
        """lxml 백엔드로 파싱 (loader.py EPUB 패턴)"""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(self.HTML, "lxml")
        self._assert_common(soup)

    def test_find_all_with_regex(self):
        """정규식 href 매칭 (bookstore.py 패턴)"""
        import re
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(self.HTML, "html.parser")
        links = soup.find_all("a", href=re.compile(r"/cat/"))
        self.assertEqual(len(links), 2)

    def test_find_all_string_lambda(self):
        """문자열 lambda 검색 (bookstore.py 패턴)"""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(self.HTML, "html.parser")
        matches = soup.find_all(string=lambda t: t and "Title" in t)
        self.assertGreaterEqual(len(matches), 1)


# ---------------------------------------------------------------------------
# 2. chardet
#    사용처: book_manager.py - 텍스트 파일 인코딩 감지
# ---------------------------------------------------------------------------
class TestChardet(unittest.TestCase):
    """chardet.detect() API 테스트"""

    def test_detect_utf8(self):
        import chardet

        result = chardet.detect("안녕하세요 hello".encode("utf-8"))
        self.assertIn("encoding", result)
        self.assertIn("confidence", result)
        self.assertIsInstance(result["confidence"], float)
        self.assertGreater(result["confidence"], 0.5)

    def test_detect_euc_kr(self):
        import chardet

        result = chardet.detect("대한민국".encode("euc-kr"))
        self.assertIsNotNone(result["encoding"])
        self.assertGreater(result["confidence"], 0)

    def test_detect_returns_encoding_key(self):
        """encoding 키가 None이 아닌 문자열 반환"""
        import chardet

        result = chardet.detect(b"plain ascii text")
        self.assertIsInstance(result["encoding"], str)


# ---------------------------------------------------------------------------
# 3. ebooklib
#    사용처: loader.py - EPUB 파일 읽기/메타데이터 추출
# ---------------------------------------------------------------------------
class TestEbooklib(unittest.TestCase):
    """ebooklib EPUB 읽기/쓰기 및 메타데이터 API 테스트"""

    def _create_epub(self, path: Path):
        """테스트용 최소 EPUB 생성"""
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier("test-id-12345")
        book.set_title("Test Book Title")
        book.set_language("ko")
        book.add_author("Test Author")

        chapter = epub.EpubHtml(title="Chapter 1", file_name="chap01.xhtml", lang="ko")
        chapter.content = b"<html><body><h1>Chapter 1</h1><p>Hello World</p></body></html>"
        book.add_item(chapter)
        book.spine = ["nav", chapter]

        nav = epub.EpubNav()
        book.add_item(nav)
        book.add_item(epub.EpubNcx())

        epub.write_epub(str(path), book)

    def test_read_epub_metadata(self):
        """epub.read_epub(), get_metadata() API"""
        from ebooklib import epub

        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
            epub_path = Path(f.name)

        try:
            self._create_epub(epub_path)
            book = epub.read_epub(str(epub_path))

            titles = book.get_metadata("DC", "title")
            self.assertGreaterEqual(len(titles), 1)
            self.assertIn("Test Book Title", titles[0][0])

            creators = book.get_metadata("DC", "creator")
            self.assertGreaterEqual(len(creators), 1)
        finally:
            epub_path.unlink(missing_ok=True)

    def test_get_items_of_type(self):
        """get_items_of_type(ITEM_DOCUMENT) API"""
        import ebooklib
        from ebooklib import epub

        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
            epub_path = Path(f.name)

        try:
            self._create_epub(epub_path)
            book = epub.read_epub(str(epub_path))
            docs = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
            self.assertGreaterEqual(len(docs), 1)
            body = docs[0].get_body_content()
            self.assertIsInstance(body, bytes)
            self.assertGreater(len(body), 0)
        finally:
            epub_path.unlink(missing_ok=True)

    def test_epub_exception_exists(self):
        """EpubException 클래스 존재 확인"""
        from ebooklib import epub

        self.assertTrue(hasattr(epub, "EpubException"))


# ---------------------------------------------------------------------------
# 4. elastic_transport
#    사용처: es_manager.py - ES 작업 시 예외 처리
# ---------------------------------------------------------------------------
class TestElasticTransport(unittest.TestCase):
    """elastic_transport 예외 클래스 import 및 상속 관계 테스트"""

    def test_exception_classes_importable(self):
        from elastic_transport import SerializationError, ConnectionError, ConnectionTimeout

        self.assertTrue(issubclass(SerializationError, Exception))
        self.assertTrue(issubclass(ConnectionError, Exception))
        self.assertTrue(issubclass(ConnectionTimeout, Exception))

    def test_exceptions_are_catchable(self):
        from elastic_transport import SerializationError, ConnectionError, ConnectionTimeout

        for exc_cls in (SerializationError, ConnectionError, ConnectionTimeout):
            with self.assertRaises(exc_cls):
                raise exc_cls("test")


# ---------------------------------------------------------------------------
# 5. elasticsearch
#    사용처: es_manager.py - 인덱스 CRUD, 검색, bulk, scroll
# ---------------------------------------------------------------------------
class TestElasticsearch(unittest.TestCase):
    """Elasticsearch 클라이언트 클래스 및 주요 메서드 시그니처 테스트"""

    def test_client_class_importable(self):
        from elasticsearch import Elasticsearch

        self.assertTrue(callable(Elasticsearch))

    def test_bad_request_error_importable(self):
        from elasticsearch import BadRequestError

        self.assertTrue(issubclass(BadRequestError, Exception))

    def test_client_has_required_methods(self):
        """응용 코드에서 사용하는 메서드가 존재하는지 확인"""
        from elasticsearch import Elasticsearch

        # __init__ 없이 속성 검사하면 descriptor가 동작하지 않으므로
        # 클래스 레벨에서 확인
        for method_name in ("search", "bulk", "update", "delete", "mget", "scroll", "clear_scroll"):
            self.assertTrue(hasattr(Elasticsearch, method_name), f"Missing method: {method_name}")

    def test_indices_namespace_exists(self):
        """IndicesClient 클래스 import 가능 확인"""
        from elasticsearch._sync.client.indices import IndicesClient

        for method_name in ("exists", "create", "delete", "get_mapping", "refresh"):
            self.assertTrue(hasattr(IndicesClient, method_name), f"Missing indices.{method_name}")

    def test_cluster_namespace_exists(self):
        """ClusterClient 클래스 import 가능 확인"""
        from elasticsearch._sync.client.cluster import ClusterClient

        self.assertTrue(hasattr(ClusterClient, "health"))

    def test_delete_by_query_exists(self):
        """delete_by_query 메서드 존재 확인 (conftest.py 패턴)"""
        from elasticsearch import Elasticsearch

        self.assertTrue(hasattr(Elasticsearch, "delete_by_query"))


# ---------------------------------------------------------------------------
# 6. fastapi
#    사용처: main.py - REST API 서버, 미들웨어, 예외 처리
# ---------------------------------------------------------------------------
class TestFastAPI(unittest.TestCase):
    """FastAPI 프레임워크 주요 API 테스트"""

    def test_core_imports(self):
        from fastapi import FastAPI

        app = FastAPI()
        self.assertIsNotNone(app)

    def test_middleware_imports(self):
        from fastapi.middleware.cors import CORSMiddleware
        from fastapi.middleware.gzip import GZipMiddleware

        self.assertTrue(callable(CORSMiddleware))
        self.assertTrue(callable(GZipMiddleware))

    def test_response_classes(self):
        from fastapi.responses import FileResponse, Response, JSONResponse

        self.assertTrue(callable(FileResponse))
        self.assertTrue(callable(Response))
        self.assertTrue(callable(JSONResponse))

    def test_exception_classes(self):
        from fastapi.exceptions import RequestValidationError
        from fastapi import HTTPException

        self.assertTrue(issubclass(HTTPException, Exception))
        self.assertTrue(issubclass(RequestValidationError, Exception))

    def test_encoders(self):
        from fastapi.encoders import jsonable_encoder

        result = jsonable_encoder({"key": "value", "num": 42})
        self.assertEqual(result["key"], "value")

    def test_route_decorator(self):
        """@app.get / @app.put / @app.delete / @app.post 데코레이터 동작"""
        from fastapi import FastAPI

        app = FastAPI()

        @app.get("/test")
        async def get_test():
            return {"status": "ok"}

        @app.put("/test/{item_id}")
        async def put_test(item_id: int):
            return {"id": item_id}

        @app.delete("/test/{item_id}")
        async def delete_test(item_id: int):
            return {"deleted": item_id}

        @app.post("/test")
        async def post_test():
            return {"created": True}

        self.assertTrue(len(app.routes) > 0)

    def test_cors_middleware_registration(self):
        """CORSMiddleware 등록"""
        from fastapi import FastAPI
        from fastapi.middleware.cors import CORSMiddleware

        app = FastAPI()
        app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"])

    def test_exception_handler_decorator(self):
        """exception_handler 데코레이터"""
        from fastapi import FastAPI
        from fastapi.exceptions import RequestValidationError

        app = FastAPI()

        @app.exception_handler(RequestValidationError)
        async def handler(request, exc):
            return {"error": str(exc)}

    def test_custom_json_response_subclass(self):
        """JSONResponse 서브클래싱 (main.py CustomJSONResponse 패턴)"""
        import json
        from fastapi.responses import JSONResponse

        class CustomJSONResponse(JSONResponse):
            def render(self, content) -> bytes:
                return json.dumps(content, ensure_ascii=False, separators=(",", ":")).encode("utf-8")

        resp = CustomJSONResponse(content={"한글": "테스트"})
        body = resp.body
        self.assertIn("한글".encode("utf-8"), body)


# ---------------------------------------------------------------------------
# 7. pydantic
#    사용처: main.py - 요청/응답 모델 정의
# ---------------------------------------------------------------------------
class TestPydantic(unittest.TestCase):
    """Pydantic BaseModel API 테스트"""

    def test_base_model_definition(self):
        """BaseModel 서브클래스 정의 및 인스턴스 생성"""
        from pydantic import BaseModel

        class BookModel(BaseModel):
            book_id: int
            category: str
            title: str
            author: str
            file_path: str
            file_type: str
            file_size: int
            line_count: int = 0
            page_count: int = 0
            isbn: str = ""
            updated_time: str
            score: float = 0.0

        book = BookModel(book_id=1, category="fiction", title="Test", author="Author", file_path="/tmp/test.txt", file_type="txt", file_size=100, updated_time="2024-01-01")
        self.assertEqual(book.book_id, 1)
        self.assertEqual(book.line_count, 0)
        self.assertEqual(book.isbn, "")

    def test_model_with_list_field(self):
        """list 필드 모델 (CategoryKeywordsModel 패턴)"""
        from pydantic import BaseModel

        class CategoryKeywordsModel(BaseModel):
            keywords: list[str]

        m = CategoryKeywordsModel(keywords=["a", "b"])
        self.assertEqual(m.keywords, ["a", "b"])

    def test_model_with_dict_field(self):
        """dict 필드 모델 (CategoryMappingsModel 패턴)"""
        from pydantic import BaseModel

        class CategoryMappingsModel(BaseModel):
            mappings: dict[str, list[str]]

        m = CategoryMappingsModel(mappings={"cat": ["kw1", "kw2"]})
        self.assertEqual(m.mappings["cat"], ["kw1", "kw2"])

    def test_model_validation_error(self):
        """타입 불일치 시 ValidationError 발생"""
        from pydantic import BaseModel, ValidationError

        class SimpleModel(BaseModel):
            value: int

        with self.assertRaises(ValidationError):
            SimpleModel(value="not_a_number")

    def test_model_dict_method(self):
        """model_dump() / dict() 변환"""
        from pydantic import BaseModel

        class SimpleModel(BaseModel):
            name: str
            count: int = 0

        m = SimpleModel(name="test")
        d = m.model_dump()
        self.assertIsInstance(d, dict)
        self.assertEqual(d["name"], "test")


# ---------------------------------------------------------------------------
# 8. httpx
#    사용처: main.py - Google OAuth 토큰 검증 (AsyncClient)
# ---------------------------------------------------------------------------
class TestHttpx(unittest.TestCase):
    """httpx AsyncClient API 존재 확인 테스트"""

    def test_async_client_importable(self):
        import httpx

        self.assertTrue(hasattr(httpx, "AsyncClient"))
        self.assertTrue(callable(httpx.AsyncClient))

    def test_async_client_is_context_manager(self):
        """async with 사용 가능 여부"""
        import httpx

        client = httpx.AsyncClient()
        self.assertTrue(hasattr(client, "__aenter__"))
        self.assertTrue(hasattr(client, "__aexit__"))
        self.assertTrue(hasattr(client, "get"))
        self.assertTrue(hasattr(client, "post"))

    def test_response_has_json_method(self):
        """Response 객체에 json() 메서드 존재"""
        import httpx

        self.assertTrue(hasattr(httpx.Response, "json"))


# ---------------------------------------------------------------------------
# 9. lxml
#    사용처: book_manager.py - EPUB OPF/NCX XML 파싱 및 조작
# ---------------------------------------------------------------------------
class TestLxml(unittest.TestCase):
    """lxml etree XML 파싱, 탐색, 조작, 직렬화 API 테스트"""

    OPF_XML = b"""<?xml version="1.0" encoding="UTF-8"?>
    <package xmlns="http://www.idpf.org/2007/opf" version="3.0">
        <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
            <dc:title>Test Book</dc:title>
        </metadata>
        <manifest>
            <item id="chap01" href="chap01.xhtml" media-type="application/xhtml+xml"/>
            <item id="chap02" href="chap02.xhtml" media-type="application/xhtml+xml"/>
        </manifest>
        <spine toc="ncx">
            <itemref idref="chap01"/>
            <itemref idref="chap02"/>
        </spine>
    </package>
    """

    NS = "http://www.idpf.org/2007/opf"

    def test_fromstring_with_xml_parser(self):
        """fromstring + XMLParser(recover=True) 패턴"""
        from lxml import etree

        tree = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        self.assertIsNotNone(tree)

    def test_namespace_find(self):
        """namespace 포함 find() (book_manager.py 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        manifest = opf.find(f"{{{self.NS}}}manifest")
        self.assertIsNotNone(manifest)
        spine = opf.find(f"{{{self.NS}}}spine")
        self.assertIsNotNone(spine)

    def test_findall_and_get_attribute(self):
        """findall() + get() 속성 접근"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        manifest = opf.find(f"{{{self.NS}}}manifest")
        items = manifest.findall(f"{{{self.NS}}}item")
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0].get("id"), "chap01")
        self.assertEqual(items[0].get("href"), "chap01.xhtml")

    def test_subelement_and_set_attribute(self):
        """SubElement 생성 및 속성 설정 (EPUB preview 생성 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        spine = opf.find(f"{{{self.NS}}}spine")

        # 기존 itemref 제거
        for ref in spine.findall(f"{{{self.NS}}}itemref"):
            spine.remove(ref)
        self.assertEqual(len(spine.findall(f"{{{self.NS}}}itemref")), 0)

        # 새 itemref 추가
        new_ref = etree.SubElement(spine, f"{{{self.NS}}}itemref")
        new_ref.set("idref", "chap_new")
        self.assertEqual(new_ref.get("idref"), "chap_new")

    def test_tostring_unicode(self):
        """tostring(encoding='unicode') 직렬화"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        result = etree.tostring(opf, encoding="unicode")
        self.assertIsInstance(result, str)
        self.assertIn("manifest", result)

    def test_remove_element(self):
        """요소 제거 (spine 항목 제거 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        spine = opf.find(f"{{{self.NS}}}spine")
        refs = spine.findall(f"{{{self.NS}}}itemref")
        original_count = len(refs)
        spine.remove(refs[0])
        self.assertEqual(len(spine.findall(f"{{{self.NS}}}itemref")), original_count - 1)

    def test_delete_attribute(self):
        """속성 삭제 (spine toc 속성 제거 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        spine = opf.find(f"{{{self.NS}}}spine")
        self.assertEqual(spine.get("toc"), "ncx")
        del spine.attrib["toc"]
        self.assertIsNone(spine.get("toc"))

    def test_iter_deep_traversal(self):
        """iter()로 깊은 탐색 (NCX navPoint 순회 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        items = list(opf.iter(f"{{{self.NS}}}item"))
        self.assertEqual(len(items), 2)

    def test_getparent(self):
        """getparent() 호출 (NCX navPoint 제거 패턴)"""
        from lxml import etree

        opf = etree.fromstring(self.OPF_XML, etree.XMLParser(recover=True))
        spine = opf.find(f"{{{self.NS}}}spine")
        ref = spine.findall(f"{{{self.NS}}}itemref")[0]
        parent = ref.getparent()
        self.assertEqual(parent.tag, f"{{{self.NS}}}spine")
        parent.remove(ref)


# ---------------------------------------------------------------------------
# 10. pymysql
#     사용처: category_mapping.py - MySQL 접속, CRUD
# ---------------------------------------------------------------------------
class TestPyMySQL(unittest.TestCase):
    """pymysql 모듈 API 존재 확인 테스트 (실제 DB 연결 불필요)"""

    def test_connect_function_exists(self):
        import pymysql

        self.assertTrue(callable(pymysql.connect))

    def test_dict_cursor_importable(self):
        from pymysql.cursors import DictCursor

        self.assertTrue(callable(DictCursor))

    def test_integrity_error_importable(self):
        import pymysql

        self.assertTrue(issubclass(pymysql.IntegrityError, Exception))

    def test_connect_signature(self):
        """connect()에 사용하는 파라미터가 수용되는지 확인"""
        import inspect
        import pymysql

        sig = inspect.signature(pymysql.connect)
        params = set(sig.parameters.keys())
        for required in ("host", "port", "database", "user", "password", "charset", "cursorclass"):
            self.assertIn(required, params, f"pymysql.connect missing parameter: {required}")


# ---------------------------------------------------------------------------
# 11. requests
#     사용처: bookstore.py, test_api.py - HTTP GET, Session, JSON 파싱
# ---------------------------------------------------------------------------
class TestRequests(unittest.TestCase):
    """requests 라이브러리 API 존재 확인 테스트"""

    def test_session_class(self):
        import requests

        session = requests.Session()
        self.assertTrue(hasattr(session, "get"))
        self.assertTrue(hasattr(session, "post"))
        self.assertTrue(hasattr(session, "headers"))

    def test_session_headers_update(self):
        import requests

        session = requests.Session()
        session.headers.update({"User-Agent": "test/1.0"})
        self.assertEqual(session.headers["User-Agent"], "test/1.0")

    def test_response_attributes(self):
        """Response 객체가 사용하는 속성을 가지는지 확인"""
        import requests

        resp = requests.models.Response()
        for attr in ("status_code", "text", "encoding", "json"):
            self.assertTrue(hasattr(resp, attr), f"Response missing: {attr}")

    def test_raise_for_status_method(self):
        import requests

        resp = requests.models.Response()
        resp.status_code = 200
        resp.raise_for_status()  # 200이면 예외 없음

    def test_exception_classes(self):
        import requests.exceptions

        self.assertTrue(issubclass(requests.exceptions.Timeout, Exception))
        self.assertTrue(issubclass(requests.exceptions.ConnectionError, Exception))

    def test_get_function_exists(self):
        import requests

        self.assertTrue(callable(requests.get))
        self.assertTrue(callable(requests.post))


# ---------------------------------------------------------------------------
# 12. pypdf
#     사용처: loader.py, isbn.py, book_manager.py - PDF 읽기/쓰기
# ---------------------------------------------------------------------------
class TestPyPDF(unittest.TestCase):
    """pypdf PdfReader/PdfWriter API 테스트"""

    def _create_pdf_bytes(self, page_count=3):
        from pypdf import PdfWriter

        writer = PdfWriter()
        for _ in range(page_count):
            writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return buf

    def test_pdf_writer_create(self):
        """PdfWriter로 빈 페이지 생성 후 write()"""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        self.assertGreater(buf.tell(), 0)

    def test_pdf_reader_pages(self):
        """PdfReader.pages 접근 및 페이지 수"""
        from pypdf import PdfReader

        buf = self._create_pdf_bytes(5)
        reader = PdfReader(buf)
        self.assertEqual(len(reader.pages), 5)

    def test_extract_text(self):
        """page.extract_text() 호출 가능"""
        from pypdf import PdfReader

        buf = self._create_pdf_bytes(1)
        reader = PdfReader(buf)
        text = reader.pages[0].extract_text()
        self.assertIsInstance(text, str)

    def test_add_page_from_reader(self):
        """Reader에서 Writer로 페이지 복사 (preview 생성 패턴)"""
        from pypdf import PdfReader, PdfWriter

        buf = self._create_pdf_bytes(3)
        reader = PdfReader(buf)
        writer = PdfWriter()
        writer.add_page(reader.pages[0])
        out = io.BytesIO()
        writer.write(out)
        out.seek(0)
        reader2 = PdfReader(out)
        self.assertEqual(len(reader2.pages), 1)


# ---------------------------------------------------------------------------
# 13. pikepdf
#     사용처: book_manager.py - PDF 검증, 메타데이터 추출
# ---------------------------------------------------------------------------
class TestPikePDF(unittest.TestCase):
    """pikepdf PDF 열기, 메타데이터, 구문 검사 API 테스트"""

    def _create_pdf_file(self):
        from pypdf import PdfWriter

        f = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_blank_page(width=612, height=792)
        writer.write(f)
        f.close()
        return Path(f.name)

    def test_open_and_close(self):
        """pikepdf.open() / pdf.close()"""
        import pikepdf

        pdf_path = self._create_pdf_file()
        try:
            pdf = pikepdf.open(pdf_path)
            self.assertIsNotNone(pdf)
            pdf.close()
        finally:
            pdf_path.unlink(missing_ok=True)

    def test_page_count(self):
        """len(pdf.pages) 페이지 수 확인"""
        import pikepdf

        pdf_path = self._create_pdf_file()
        try:
            pdf = pikepdf.open(pdf_path)
            self.assertEqual(len(pdf.pages), 2)
            pdf.close()
        finally:
            pdf_path.unlink(missing_ok=True)

    def test_pdf_version(self):
        """pdf.pdf_version 속성"""
        import pikepdf

        pdf_path = self._create_pdf_file()
        try:
            pdf = pikepdf.open(pdf_path)
            self.assertIsNotNone(pdf.pdf_version)
            pdf.close()
        finally:
            pdf_path.unlink(missing_ok=True)

    def test_docinfo(self):
        """pdf.docinfo 메타데이터 접근"""
        import pikepdf

        pdf_path = self._create_pdf_file()
        try:
            pdf = pikepdf.open(pdf_path)
            docinfo = pdf.docinfo
            # docinfo가 존재하고 get() 호출 가능
            self.assertIsNotNone(docinfo)
            _ = docinfo.get("/Title")  # None 가능
            pdf.close()
        finally:
            pdf_path.unlink(missing_ok=True)

    def test_check_pdf_syntax(self):
        """pdf.check() 구문 검사 (이전 check_pdf_syntax -> check)"""
        import pikepdf

        pdf_path = self._create_pdf_file()
        try:
            pdf = pikepdf.open(pdf_path)
            # check() 또는 check_pdf_syntax() 중 존재하는 것 확인
            has_check = hasattr(pdf, "check") or hasattr(pdf, "check_pdf_syntax")
            self.assertTrue(has_check, "pikepdf lacks check/check_pdf_syntax method")
            pdf.close()
        finally:
            pdf_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# 14. python-docx
#     사용처: loader.py - DOCX 파일 읽기
# ---------------------------------------------------------------------------
class TestPythonDocx(unittest.TestCase):
    """python-docx Document/Paragraph API 테스트"""

    def test_create_and_read_document(self):
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = Path(f.name)

        try:
            # 생성
            doc = Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("Second Paragraph")
            doc.save(str(docx_path))

            # 읽기 (loader.py 패턴)
            doc2 = Document(str(docx_path))
            texts = [p.text for p in doc2.paragraphs]
            self.assertIn("Hello World", texts)
            self.assertIn("Second Paragraph", texts)
        finally:
            docx_path.unlink(missing_ok=True)

    def test_paragraph_text_attribute(self):
        """paragraph.text 속성 존재"""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("test")
        self.assertEqual(p.text, "test")


# ---------------------------------------------------------------------------
# 15. striprtf
#     사용처: loader.py - RTF → 텍스트 변환
# ---------------------------------------------------------------------------
class TestStripRTF(unittest.TestCase):
    """striprtf rtf_to_text() API 테스트"""

    def test_rtf_to_text_basic(self):
        from striprtf.striprtf import rtf_to_text

        rtf_content = r"{\rtf1\ansi Hello World}"
        result = rtf_to_text(rtf_content)
        self.assertIn("Hello", result)

    def test_rtf_to_text_with_errors_ignore(self):
        """errors='ignore' 파라미터 (loader.py 패턴)"""
        from striprtf.striprtf import rtf_to_text

        rtf_content = r"{\rtf1\ansi Test content}"
        result = rtf_to_text(rtf_content, errors="ignore")
        self.assertIsInstance(result, str)


# ---------------------------------------------------------------------------
# 16. uvicorn
#     사용처: CLI 서버 실행 (import 없이 사용)
# ---------------------------------------------------------------------------
class TestUvicorn(unittest.TestCase):
    """uvicorn 모듈 존재 확인"""

    def test_importable(self):
        import uvicorn

        self.assertTrue(hasattr(uvicorn, "run"))


if __name__ == "__main__":
    unittest.main()
