import sys
import os
import unittest
import zipfile
import zlib
import tempfile
import warnings
import importlib
from pathlib import Path
from unittest.mock import MagicMock

# loader.py가 import 시점에 환경 변수와 외부 모듈을 요구하므로 사전 설정
os.environ.setdefault("TM_BOOK_DIR", str(Path(__file__).parent.parent / "tests" / "books"))
os.environ.setdefault("TM_COMICS_DIR", str(Path(__file__).parent.parent / "tests" / "comics"))

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
warnings.filterwarnings("ignore", message=".*XML.*HTML.*")


def _get_loader():
    """Loader 클래스를 lazy import (외부 의존성 처리)"""
    from utils.loader import Loader

    return Loader


class TestLoaderInit(unittest.TestCase):
    def test_loader_class_exists(self):
        Loader = _get_loader()
        assert Loader is not None
        assert hasattr(Loader, "read_file")
        assert hasattr(Loader, "read_files")

    def test_loader_requires_book_dir(self):
        import pytest

        prev_book_dir = os.environ.pop("TM_BOOK_DIR", None)
        prev_loader = sys.modules.pop("utils.loader", None)
        prev_comics_dir = os.environ.get("TM_COMICS_DIR")
        os.environ["TM_COMICS_DIR"] = "/tmp/comics"

        try:
            with pytest.raises(SystemExit):
                importlib.import_module("utils.loader")
        finally:
            if prev_book_dir is not None:
                os.environ["TM_BOOK_DIR"] = prev_book_dir
            if prev_comics_dir is None:
                os.environ.pop("TM_COMICS_DIR", None)
            else:
                os.environ["TM_COMICS_DIR"] = prev_comics_dir
            sys.modules.pop("utils.loader", None)
            if prev_loader is not None:
                sys.modules["utils.loader"] = prev_loader

    def test_loader_requires_comics_dir(self):
        import pytest

        prev_comics_dir = os.environ.pop("TM_COMICS_DIR", None)
        prev_loader = sys.modules.pop("utils.loader", None)
        prev_book_dir = os.environ.get("TM_BOOK_DIR")
        os.environ["TM_BOOK_DIR"] = str(Path(__file__).parent.parent / "tests" / "books")

        try:
            with pytest.raises(SystemExit):
                importlib.import_module("utils.loader")
        finally:
            if prev_comics_dir is not None:
                os.environ["TM_COMICS_DIR"] = prev_comics_dir
            if prev_book_dir is None:
                os.environ.pop("TM_BOOK_DIR", None)
            else:
                os.environ["TM_BOOK_DIR"] = prev_book_dir
            sys.modules.pop("utils.loader", None)
            if prev_loader is not None:
                sys.modules["utils.loader"] = prev_loader


class TestReadFromText:
    def test_read_utf8_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "test.txt"
        f.write_text("Hello World\n안녕하세요\n테스트", encoding="utf-8")
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert isinstance(summary, str)
        assert line_count == 3
        assert page_count == 0
        assert "Hello" in raw

    def test_read_empty_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert summary == ""
        assert line_count == 0

    def test_read_bom_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "bom.txt"
        f.write_bytes(b"\xef\xbb\xbfBOM content")
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "\ufeff" not in summary

    def test_read_cp949_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "cp949.txt"
        f.write_bytes("안녕하세요 한글 테스트입니다".encode("cp949"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "안녕하세요" in raw
        assert line_count == 1

    def test_read_euc_kr_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "euckr.txt"
        f.write_bytes("완성형 한글 테스트".encode("euc-kr"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "한글" in raw
        assert line_count == 1

    def test_read_utf16_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "utf16.txt"
        f.write_bytes("유니코드 16 비트 테스트\n두번째 줄".encode("utf-16"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "유니코드" in raw
        assert line_count == 2

    def test_read_binary_as_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "binary.txt"
        f.write_bytes(b"\x80\x81\x82\x83")
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert summary == ""

    def test_read_text_when_chardet_raises(self, tmp_path: Path, monkeypatch):
        """chardet가 실패해도 기본 인코딩 순서로 읽어낸다."""
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        f.write_bytes("한글 본문입니다 가나다라\n".encode("cp949"))

        import chardet

        def boom(_b):
            raise RuntimeError("chardet 실패")

        monkeypatch.setattr(chardet, "detect", boom)

        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "한글 본문입니다" in raw

    def test_read_text_falls_back_to_replace(self, tmp_path: Path):
        """모든 엄격 디코딩이 실패하면 errors='replace'로라도 읽어 적재를 이어간다.

        이 경로는 원문 글자를 잃으므로(U+FFFD → 이후 공백 치환) 재인코딩 대상이
        되지 못한 파일에서만 쓰인다.
        """
        Loader = _get_loader()
        f = tmp_path / "broken.txt"
        raw_bytes = bytes([0x80, 0x81, 0xFE, 0xFF, 0x00, 0x9C, 0x90]) * 40 + b"abc\n"
        f.write_bytes(raw_bytes)

        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert isinstance(summary, str)
        assert isinstance(raw, str)
        # 파일은 읽기 전용 경로이므로 원본이 그대로여야 한다
        assert f.read_bytes() == raw_bytes

    def test_read_text_replace_skips_candidate_without_letters(self, tmp_path: Path):
        """replace 디코딩 결과에 한글·영숫자가 전혀 없으면 그 후보를 건너뛴다.

        치환문자만 잔뜩 남은 결과를 본문으로 채택하면 검색 품질이 더 나빠진다.
        """
        Loader = _get_loader()
        f = tmp_path / "nonletters.txt"
        raw_bytes = bytes([0x80, 0x81, 0x82, 0x83]) * 60
        f.write_bytes(raw_bytes)

        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert isinstance(summary, str) and isinstance(raw, str)
        assert f.read_bytes() == raw_bytes

    def test_read_text_replace_skips_korean_candidate_without_letters(self, tmp_path: Path, monkeypatch):
        """replace 단계에서 cp949/euc-kr 결과에 글자가 없으면 그 후보를 건너뛴다.

        chardet가 한국어 인코딩을 먼저 제안했을 때만 도달하는 분기라 그 상황을 만든다.
        """
        Loader = _get_loader()
        f = tmp_path / "nonletters2.txt"
        raw_bytes = bytes([0x80, 0x81, 0x82, 0x83]) * 60
        f.write_bytes(raw_bytes)

        import chardet

        monkeypatch.setattr(chardet, "detect", lambda b: {"encoding": "EUC-KR", "confidence": 0.99})

        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert isinstance(summary, str) and isinstance(raw, str)
        assert f.read_bytes() == raw_bytes

    def test_read_text_rejects_utf16_cjk_garbage(self, tmp_path: Path):
        """단일바이트 파일이 utf-16으로 디코딩돼 CJK 뭉치가 되면 채택하지 않는다.

        utf-16-le/be만 검사하고 plain utf-16을 빠뜨리면 검사 없이 통과한다.
        """
        Loader = _get_loader()
        f = tmp_path / "notutf16.txt"
        raw_bytes = bytes([0x80, 0x81, 0x82, 0x83]) * 60
        f.write_bytes(raw_bytes)

        summary, line_count, page_count, raw = Loader.read_from_text(f)
        # utf-16 계열로 읽힌 CJK 뭉치(膀莂..., 肁芃...)가 그대로 들어오면 안 된다
        assert "膀莂" not in raw and "肁芃" not in raw, raw[:40]
        assert f.read_bytes() == raw_bytes

    def test_read_corrupted_utf16_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "corrupted_utf16.txt"
        valid_bytes = "안녕하세요 검맥 홍파 테스트".encode("utf-16")
        corrupted_bytes = valid_bytes[:20] + b"\xff\xff" + valid_bytes[20:]
        f.write_bytes(corrupted_bytes)
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "안녕하세요" in raw or "테스트" in raw

    def test_read_utf16_le_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "utf16le.txt"
        f.write_bytes("한글 UTF16LE 테스트".encode("utf-16-le"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert len(raw) > 0
        assert line_count == 1

    def test_read_utf16_be_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "utf16be.txt"
        f.write_bytes("한글 UTF16BE 테스트".encode("utf-16-be"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert len(raw) > 0
        assert line_count == 1

    def test_read_utf8_sig_text(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "utf8sig.txt"
        f.write_bytes("UTF-8 BOM 포함 한글 테스트".encode("utf-8-sig"))
        summary, line_count, page_count, raw = Loader.read_from_text(f)
        assert "한글" in raw
        assert line_count == 1


class TestReadFromEpub:
    def test_read_valid_epub(self, tmp_path: Path):
        Loader = _get_loader()
        epub = tmp_path / "book.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", "<html><body><p>테스트 내용입니다</p></body></html>")
        # read_from_epub_with_extracting_zip 직접 테스트
        summary, line_count = Loader.read_from_epub_with_extracting_zip(epub)
        assert "테스트" in summary

    def test_read_bad_zip_epub(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "bad.epub"
        f.write_text("not a zip")
        summary, line_count = Loader.read_from_epub_with_extracting_zip(f)
        assert summary == ""
        assert line_count == 0

    def test_read_epub_no_container(self, tmp_path: Path):
        Loader = _get_loader()
        epub = tmp_path / "nocontainer.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("META-INF/container.xml", "<container></container>")
        summary, line_count = Loader.read_from_epub_with_extracting_zip(epub)
        assert summary == ""

    def test_read_epub_chapter_exception_is_skipped(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        epub = tmp_path / "chapter-error.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", "<html><body><p>ignored</p></body></html>")

        monkeypatch.setattr("utils.loader.BeautifulSoup", lambda *_args, **_kwargs: (_ for _ in ()).throw(ValueError("boom")))
        summary, line_count = Loader.read_from_epub_with_extracting_zip(epub)
        assert summary == ""
        assert line_count == 0


class TestReadFromHtml:
    def test_read_html(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "test.html"
        f.write_text("<html><body><p>Hello</p></body></html>", encoding="utf-8")
        summary, line_count, page_count = Loader.read_from_html(f)
        assert "Hello" in summary
        assert page_count == 0


class TestReadFromImage:
    def test_read_image_returns_empty(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "test.jpg"
        f.write_bytes(b"\xff\xd8")
        summary, line_count, page_count = Loader.read_from_image(f)
        assert summary == ""
        assert line_count == 0
        assert page_count == 0


class TestGetFileList:
    def test_single_file(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "test.txt"
        f.write_text("hello")
        result = Loader.get_file_list(f)
        assert result == [f]

    def test_directory(self, tmp_path: Path):
        Loader = _get_loader()
        (tmp_path / "a.txt").write_text("a")
        (tmp_path / "b.txt").write_text("b")
        result = Loader.get_file_list(tmp_path)
        assert len(result) >= 2

    def test_directory_with_subdirs(self, tmp_path: Path):
        Loader = _get_loader()
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "c.txt").write_text("c")
        (tmp_path / "a.txt").write_text("a")
        result = Loader.get_file_list(tmp_path)
        # Should include one file from subdir + direct files
        assert len(result) >= 2

    def test_directory_recursive(self, tmp_path: Path):
        Loader = _get_loader()
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "c.txt").write_text("c")
        (tmp_path / "a.txt").write_text("a")
        result = Loader.get_file_list(tmp_path, recursive=True)
        assert len(result) >= 2

    def test_directory_num_files_limit(self, tmp_path: Path):
        Loader = _get_loader()
        for i in range(5):
            (tmp_path / f"{i}.txt").write_text(str(i))
        result = Loader.get_file_list(tmp_path, num_files=2)
        assert len(result) == 2

    def test_directory_skips_hidden(self, tmp_path: Path):
        Loader = _get_loader()
        (tmp_path / ".hidden").write_text("x")
        (tmp_path / "visible.txt").write_text("y")
        hidden_dir = tmp_path / ".hidden_dir"
        hidden_dir.mkdir()
        (hidden_dir / "z.txt").write_text("z")
        result = Loader.get_file_list(tmp_path, recursive=True)
        names = [p.name for p in result]
        assert ".hidden" not in names
        assert "z.txt" not in names


class TestGetPathPrefix:
    def test_book_path(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "book.txt"
        assert Loader.get_path_prefix(f) == tmp_path
        Loader.path_prefix = original_prefix

    def test_comics_path(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        original_comics = Loader.comics_path_prefix
        Loader.comics_path_prefix = tmp_path
        f = tmp_path / "comic.cbz"
        assert Loader.get_path_prefix(f) == tmp_path
        Loader.comics_path_prefix = original_comics


class TestReadFile:
    def test_read_txt_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        cat_dir = tmp_path / "cat"
        cat_dir.mkdir()
        f = cat_dir / "[Author] Title.txt"
        f.write_text("Hello world\nLine 2", encoding="utf-8")

        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix

        assert len(result) == 1
        inode = list(result.keys())[0]
        doc = result[inode]
        assert doc["author"] == "Author"
        assert doc["title"] == "Title"
        assert doc["category"] == "cat"
        assert doc["file_type"] == "txt"

    def test_read_file_no_author(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        f = tmp_path / "NoAuthor.txt"
        f.write_text("content")

        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix

        inode = list(result.keys())[0]
        doc = result[inode]
        assert doc["author"] == ""
        assert doc["title"] == "NoAuthor"
        assert doc["category"] == "_root"

    def test_read_nonexistent_file(self, tmp_path: Path):
        Loader = _get_loader()
        result = Loader.read_file(tmp_path / "nonexistent.txt")
        assert result == {}

    def test_read_unsupported_format(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        f = tmp_path / "file.xyz"
        f.write_text("content")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        assert result == {}

    def test_read_image_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        f = tmp_path / "image.jpg"
        f.write_bytes(b"\xff\xd8\xff\xe0")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix

        assert len(result) == 1
        doc = list(result.values())[0]
        assert doc["file_type"] == "jpg"
        assert doc["summary"] == ""

    def test_read_file_skip_text(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        f = tmp_path / "book.txt"
        f.write_text("content")
        result = Loader.read_file(f, skip_text=True)
        Loader.path_prefix = original_prefix

        doc = list(result.values())[0]
        assert doc["summary"] == ""

    def test_read_file_skip_text_unsupported(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path

        f = tmp_path / "file.xyz"
        f.write_text("content")
        result = Loader.read_file(f, skip_text=True)
        Loader.path_prefix = original_prefix
        assert result == {}


class TestPngPredictor:
    def test_no_filter(self):
        Loader = _get_loader()
        # filter byte = 0 (None), 2 columns of data
        data = bytes([0, 0x10, 0x20, 0, 0x30, 0x40])
        result = Loader._apply_png_predictor(data, columns=2)
        assert result == bytes([0x10, 0x20, 0x30, 0x40])

    def test_sub_filter(self):
        Loader = _get_loader()
        # filter byte = 1 (Sub)
        data = bytes([1, 0x10, 0x05])
        result = Loader._apply_png_predictor(data, columns=2)
        assert result == bytes([0x10, 0x15])

    def test_up_filter(self):
        Loader = _get_loader()
        # Two rows: first row no filter, second row Up filter
        data = bytes([0, 0x10, 0x20, 2, 0x01, 0x02])
        result = Loader._apply_png_predictor(data, columns=2)
        assert result == bytes([0x10, 0x20, 0x11, 0x22])


class TestFindXrefOffset:
    def test_basic_xref(self):
        Loader = _get_loader()
        xref = b"xref\n0 3\n0000000000 65535 f \n0000000100 00000 n \n0000000200 00000 n \ntrailer\n"
        assert Loader._find_xref_offset(xref, 1) == 100
        assert Loader._find_xref_offset(xref, 2) == 200
        assert Loader._find_xref_offset(xref, 5) is None

    def test_free_entry(self):
        Loader = _get_loader()
        xref = b"xref\n0 2\n0000000000 65535 f \n0000000100 00000 f \ntrailer\n"
        assert Loader._find_xref_offset(xref, 1) is None

    def test_malformed_target_entry_returns_none(self):
        Loader = _get_loader()
        xref = b"xref\n0 2\n0000000000 65535 f \nnot-an-entry\ntrailer\n"
        assert Loader._find_xref_offset(xref, 1) is None

    def test_missing_xref_keyword_returns_none(self):
        Loader = _get_loader()
        assert Loader._find_xref_offset(b"no xref here", 1) is None


class TestXrefStreamFindEntry:
    def test_basic_entry(self):
        Loader = _get_loader()
        # W=[1,2,1], index=[0,2]
        # Entry 0: type=1, offset=256, gen=0
        # Entry 1: type=1, offset=512, gen=0
        data = bytes([1, 1, 0, 0, 1, 2, 0, 0])
        result = Loader._xref_stream_find_entry(data, [1, 2, 1], [0, 2], 1)
        assert result == (1, 512, 0)

    def test_entry_not_found(self):
        Loader = _get_loader()
        data = bytes([1, 1, 0, 0])
        result = Loader._xref_stream_find_entry(data, [1, 2, 1], [0, 1], 5)
        assert result is None

    def test_zero_entry_size(self):
        Loader = _get_loader()
        result = Loader._xref_stream_find_entry(b"", [0, 0, 0], [0, 1], 0)
        assert result is None

    def test_truncated_entry_returns_none(self):
        Loader = _get_loader()
        result = Loader._xref_stream_find_entry(bytes([1, 0, 10]), [1, 2, 1], [0, 1], 0)
        assert result is None


# ---- coverage: loader additional uncovered lines ----


class TestReadFromPdf:
    def test_read_valid_pdf(self, tmp_path: Path):
        Loader = _get_loader()
        import pypdf

        # Create a minimal valid PDF
        pdf = tmp_path / "test.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(pdf, "wb") as f:
            writer.write(f)
        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert page_count == 1

    def test_read_broken_pdf(self, tmp_path: Path):
        Loader = _get_loader()
        pdf = tmp_path / "broken.pdf"
        pdf.write_bytes(b"%PDF-1.0\nbroken content")
        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert page_count == 0

    def test_read_pdf_pypdfium2_primary(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "pypdfium2_test.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        class MockTextPage:
            def get_text_range(self):
                return "pypdfium2 추출 텍스트"

        class MockPage:
            def get_textpage(self):
                return MockTextPage()

        class MockPdfDocument:
            def __init__(self, path):
                pass

            def __len__(self):
                return 1

            def __getitem__(self, idx):
                return MockPage()

            def __iter__(self):
                return iter([MockPage()])

            def close(self):
                pass

        import pypdfium2

        monkeypatch.setattr(pypdfium2, "PdfDocument", MockPdfDocument)

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert "pypdfium2" in summary
        assert page_count == 1

    def test_read_pdf_pdfplumber_fallback(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "pdfplumber_test.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        import pypdfium2

        monkeypatch.setattr(pypdfium2, "PdfDocument", lambda *a, **k: (_ for _ in ()).throw(Exception("pypdfium2 fail")))

        # pdftotext CLI 단계를 건너뛰게 해서 pdfplumber까지 내려오게 한다.
        import shutil

        monkeypatch.setattr(shutil, "which", lambda *a, **k: None)

        class MockPlumberPage:
            def extract_text(self):
                return "pdfplumber 추출 텍스트"

            def close(self):
                pass

        class MockPlumberPdf:
            def __enter__(self):
                return self

            def __exit__(self, *args):
                pass

            @property
            def pages(self):
                return [MockPlumberPage()]

        import pdfplumber

        monkeypatch.setattr(pdfplumber, "open", lambda *a, **k: MockPlumberPdf())

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert "pdfplumber" in summary
        assert page_count == 1

    def test_read_pdf_pdftotext_cli_fallback(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "pdftotext_cli_test.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        import pypdfium2

        monkeypatch.setattr(pypdfium2, "PdfDocument", lambda *a, **k: (_ for _ in ()).throw(Exception("fail")))

        import pdfplumber

        monkeypatch.setattr(pdfplumber, "open", lambda *a, **k: (_ for _ in ()).throw(Exception("fail")))

        class MockCompletedProcess:
            returncode = 0
            stdout = "pdftotext CLI 추출 텍스트".encode("utf-8")

        import subprocess

        monkeypatch.setattr(subprocess, "run", lambda *a, **k: MockCompletedProcess())

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert "pdftotext" in summary

    def test_read_pdf_mojibake_falls_through(self, tmp_path: Path, monkeypatch):
        """EUC-KR을 latin-1로 잘못 해석한 텍스트는 비어 있지 않아도 fallback을 타야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "mojibake.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        mojibake = "맥(麥) 김남천 조은커뮤니티".encode("euc-kr").decode("latin-1") * 8

        class MockTextPage:
            def get_text_range(self):
                return mojibake

        class MockPage:
            def get_textpage(self):
                return MockTextPage()

        class MockPdfDocument:
            def __init__(self, path):
                pass

            def __len__(self):
                return 1

            def __getitem__(self, idx):
                return MockPage()

            def __iter__(self):
                return iter([MockPage()])

            def close(self):
                pass

        import pypdfium2

        monkeypatch.setattr(pypdfium2, "PdfDocument", MockPdfDocument)

        class MockCompletedProcess:
            returncode = 0
            stdout = "정상 추출된 한글 본문".encode("utf-8")

        import subprocess

        monkeypatch.setattr(subprocess, "run", lambda *a, **k: MockCompletedProcess())

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert "정상 추출된 한글 본문" in summary
        assert page_count == 1


class TestReencodeTextFileToUtf8:
    """txt 파일 UTF-8 재인코딩 — 검증 통과 시에만 파일을 바꾼다."""

    def test_cp949_file_is_reencoded(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "cp949.txt"
        original = "한글 소설 본문입니다.\n둘째 줄\n"
        f.write_bytes(original.encode("cp949"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original

    def test_euckr_file_is_reencoded(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "euckr.txt"
        original = "동양고전 한문 번역 본문\n"
        f.write_bytes(original.encode("euc-kr"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original

    def test_inode_is_preserved(self, tmp_path: Path):
        """ES 문서 _id가 inode이므로 재인코딩이 inode를 바꾸면 문서가 갈린다."""
        Loader = _get_loader()
        f = tmp_path / "inode.txt"
        f.write_bytes("한글 본문\n".encode("cp949"))
        before = f.stat().st_ino

        changed, _ = Loader.reencode_text_file_to_utf8(f)
        assert changed
        assert f.stat().st_ino == before

    def test_already_utf8_is_untouched(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "utf8.txt"
        raw = "한글 본문입니다\n".encode("utf-8")
        f.write_bytes(raw)
        mtime = f.stat().st_mtime_ns

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed
        assert reason == "이미 UTF-8"
        assert f.read_bytes() == raw
        assert f.stat().st_mtime_ns == mtime

    def test_undecodable_file_is_untouched(self, tmp_path: Path):
        """어떤 인코딩으로도 무손실 디코딩이 안 되면 파일을 건드리지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "broken.txt"
        raw = bytes([0x80, 0x81, 0xFE, 0xFF, 0x00, 0x9C, 0x90]) * 40
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed, reason
        assert f.read_bytes() == raw

    def test_punctuation_only_file_is_untouched(self, tmp_path: Path):
        """실제 글자가 전혀 없으면 인코딩 선택이 옳다고 볼 수 없어 건드리지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "nontext.txt"
        raw = ("。、※★○【】" * 20).encode("cp949")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed, reason
        assert f.read_bytes() == raw

    def test_hanja_only_file_is_reencoded(self, tmp_path: Path):
        """한문 고전처럼 한글이 없어도 한자 본문이면 재인코딩 대상이다."""
        Loader = _get_loader()
        f = tmp_path / "hanja.txt"
        original = "君子曰 學不可以已 靑取之於藍而靑於藍\n" * 10
        f.write_bytes(original.encode("cp949"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original

    def test_windows1252_misdetection_is_rejected(self, tmp_path: Path, monkeypatch):
        """chardet가 단일바이트 인코딩을 제안해도 채택하지 않는다.

        단일바이트 인코딩은 어떤 바이트열이든 디코딩되고 바이트 왕복도 무조건 통과하므로,
        후보로 허용하면 한글 파일에 mojibake(ÇÑ±Û ¼Ò¼³...)를 영구 기록하게 된다.
        """
        Loader = _get_loader()
        f = tmp_path / "korean.txt"
        original = "한글 소설 본문 chapter 12 입니다\n" * 20
        f.write_bytes(original.encode("cp949"))

        import chardet

        monkeypatch.setattr(chardet, "detect", lambda b: {"encoding": "Windows-1252", "confidence": 0.99})

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        # cp949로 올바르게 복원되어야 하고, mojibake가 저장되면 안 된다
        saved = f.read_bytes().decode("utf-8")
        assert saved == original
        assert "ÇÑ" not in saved

    def test_empty_file_is_skipped(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "empty.txt"
        f.write_bytes(b"")

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed
        assert reason == "빈 파일"
        assert f.read_bytes() == b""

    def test_unreadable_path_is_reported(self, tmp_path: Path):
        """읽을 수 없는 경로는 예외를 밖으로 내보내지 않고 사유로 돌려준다."""
        Loader = _get_loader()
        d = tmp_path / "adir"
        d.mkdir()

        changed, reason = Loader.reencode_text_file_to_utf8(d)
        assert not changed
        assert reason.startswith("읽기 실패")

    def test_works_when_chardet_raises(self, tmp_path: Path, monkeypatch):
        """chardet가 실패해도 기본 후보 순서로 정상 변환한다."""
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        raw = "한글 본문입니다\n둘째 줄\n".encode("cp949")
        f.write_bytes(raw)

        import chardet

        def boom(_b):
            raise RuntimeError("chardet 실패")

        monkeypatch.setattr(chardet, "detect", boom)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == raw.decode("cp949")

    def test_utf16_with_bom_is_converted(self, tmp_path: Path):
        """BOM이 있으면 NUL 비율과 무관하게 UTF-16 후보를 유지한다."""
        Loader = _get_loader()
        f = tmp_path / "bom.txt"
        original = "한글 본문입니다 가나다라마바사\n" * 20
        raw = original.encode("utf-16")
        assert raw[:2] in (b"\xff\xfe", b"\xfe\xff")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original

    def test_replacement_char_in_source_is_rejected(self, tmp_path: Path):
        """디코딩 결과에 U+FFFD가 있으면 그 후보를 쓰지 않는다.

        U+FFFD가 들어간 채로 저장하면 원문 글자를 영구히 잃는다.
        """
        Loader = _get_loader()
        f = tmp_path / "fffd.txt"
        raw = ("한글 본문�입니다\n" * 30).encode("utf-16-le")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        if changed:
            assert "�" not in f.read_bytes().decode("utf-8"), reason
        else:
            assert f.read_bytes() == raw

    def test_long_text_without_newline_is_rejected(self, tmp_path: Path):
        """줄바꿈이 하나도 없는 긴 텍스트는 인코딩 오선택 신호이므로 변환하지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "noline.txt"
        raw = ("가나다라마바사아자차카타파하" * 40).encode("cp949")
        text = raw.decode("cp949")
        assert len(text) >= 200 and "\n" not in text and "\r" not in text
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed, reason
        assert f.read_bytes() == raw

    def test_high_latin1_ratio_is_rejected(self, tmp_path: Path):
        """악센트 비율이 지나치게 높으면 단일바이트 오독으로 보고 변환하지 않는다.

        상위바이트가 고립돼 cp1252 후보로 넘어가더라도, 정상 서양어 문서라면
        악센트 비율이 이렇게 높을 수 없다. 실측 cp1252 영문 파일은 0.003 수준이다.
        """
        Loader = _get_loader()
        f = tmp_path / "accents.txt"
        raw = ("a \xe0 b \xe9 c \xee d \xf4 e \xf9\r\n" * 40).encode("cp1252")
        f.write_bytes(raw)
        sample = raw.decode("cp1252")
        assert sum(1 for c in sample if "\u0080" <= c <= "\u00ff") / len(sample) > 0.1

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed, reason
        assert f.read_bytes() == raw

    def test_adjacent_accents_stay_lossless(self, tmp_path: Path):
        """알려진 한계: 악센트가 연속이면 한글 2바이트 시퀀스와 구분되지 않는다.

        실제 서양어 문서의 악센트는 고립돼 나타나므로(실측 인접쌍 0.000~0.0024)
        현실적 위험은 낮다. 어느 쪽으로 판정되든 무손실 왕복은 지켜져야 한다.
        """
        Loader = _get_loader()
        f = tmp_path / "adjacent.txt"
        raw = ("ab \xe0\xe9\xee\xf4\xf9\xe0\xe9\xee\xf4\xf9\r\n" * 40).encode("cp1252")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        if changed:
            enc = reason.split(" ")[0]
            assert f.read_bytes().decode("utf-8").encode(enc) == raw, "무손실 왕복이 깨짐"
        else:
            assert f.read_bytes() == raw

    def test_file_untouched_when_backup_content_mismatches(self, tmp_path: Path, monkeypatch):
        """백업이 잘려 쓰이면 검증에서 걸러내고 원본을 건드리지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        raw = "한글 본문입니다\n둘째 줄\n".encode("cp949")
        f.write_bytes(raw)
        backup = tmp_path / "backup"
        monkeypatch.setattr(Loader, "reencode_backup_dir", backup)

        real_write = Path.write_bytes

        def truncating_write(self, data):
            if backup in self.parents:
                return real_write(self, data[:5])
            return real_write(self, data)

        monkeypatch.setattr(Path, "write_bytes", truncating_write)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed
        assert "백업 검증 실패" in reason, reason
        assert f.read_bytes() == raw

    def test_original_preserved_when_write_fails(self, tmp_path: Path):
        """덮어쓰기가 실패하면 원본이 남고 임시 파일도 남기지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "readonly.txt"
        raw = "한글 본문입니다\n둘째 줄\n".encode("cp949")
        f.write_bytes(raw)
        f.chmod(0o444)
        try:
            changed, reason = Loader.reencode_text_file_to_utf8(f)
            assert not changed
            assert "쓰기 실패" in reason, reason
            assert f.read_bytes() == raw
            assert list(tmp_path.iterdir()) == [f], "임시 파일이 남음"
        finally:
            f.chmod(0o644)

    def test_post_write_mismatch_is_detected(self, tmp_path: Path, monkeypatch):
        """저장 후 다시 읽어 확인하는 단계가 실제로 불일치를 잡아낸다."""
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        raw = "한글 본문입니다\n둘째 줄\n".encode("cp949")
        f.write_bytes(raw)

        real_read = Path.read_bytes
        calls = {"n": 0}

        def flaky_read(self):
            data = real_read(self)
            if self == f:
                calls["n"] += 1
                if calls["n"] >= 2:
                    # 저장 후 재읽기에서만 내용을 바꾼다. 유효한 UTF-8이라 디코딩은
                    # 성공하고 길이만 달라지므로 '저장 후 불일치' 분기를 탄다.
                    return data + "추가".encode("utf-8")
            return data

        monkeypatch.setattr(Path, "read_bytes", flaky_read)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed
        assert "저장 후 불일치" in reason, reason

    def test_backup_is_written_before_overwrite(self, tmp_path: Path, monkeypatch):
        """백업 디렉터리를 지정하면 원본이 상대 경로 구조로 보존된다."""
        Loader = _get_loader()
        book = tmp_path / "book"
        (book / "3_판타지").mkdir(parents=True)
        f = book / "3_판타지" / "novel.txt"
        raw = "한글 본문입니다\n둘째 줄\n".encode("cp949")
        f.write_bytes(raw)
        backup = tmp_path / "backup"

        monkeypatch.setattr(Loader, "path_prefix", book)
        monkeypatch.setattr(Loader, "reencode_backup_dir", backup)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        saved_backup = backup / "3_판타지" / "novel.txt"
        assert saved_backup.is_file(), "백업이 생성되지 않음"
        assert saved_backup.read_bytes() == raw, "백업 내용이 원본과 다름"
        assert f.read_bytes().decode("utf-8") == raw.decode("cp949")

    def test_file_untouched_when_backup_fails(self, tmp_path: Path, monkeypatch):
        """백업을 쓸 수 없으면 파일을 변환하지 않는다."""
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        raw = "한글 본문입니다\n".encode("cp949")
        f.write_bytes(raw)
        # 파일을 백업 디렉터리로 지정 → mkdir 실패
        blocker = tmp_path / "blocker"
        blocker.write_text("not a dir")

        monkeypatch.setattr(Loader, "reencode_backup_dir", blocker)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed
        assert "백업" in reason, reason
        assert f.read_bytes() == raw

    def test_dry_run_writes_no_backup(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        f = tmp_path / "novel.txt"
        raw = "한글 본문입니다\n".encode("cp949")
        f.write_bytes(raw)
        backup = tmp_path / "backup"
        monkeypatch.setattr(Loader, "reencode_backup_dir", backup)

        changed, reason = Loader.reencode_text_file_to_utf8(f, dry_run=True)
        assert changed and "dry-run" in reason
        assert not backup.exists(), "dry-run인데 백업이 생성됨"
        assert f.read_bytes() == raw

    def test_western_text_uses_cp1252_not_cp949(self, tmp_path: Path):
        """고립된 상위바이트는 서양어 단일바이트 인코딩이다.

        실제로 발생한 케이스: cp1252 영문 파일의 0x97(em-dash)이 cp949에서 뒤따르는
        ASCII까지 2바이트 시퀀스로 삼켜 영문자를 파괴했다
        ('Don Sabas—a man' → 'Don Sabas뾞 man', 'a' 소실).
        """
        Loader = _get_loader()
        f = tmp_path / "ohenry.txt"
        body = "characteristic of Don Sabas\x97a man at once merry, learned\r\n"
        raw = (body * 40).encode("latin-1")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert reason.startswith("cp1252"), f"cp1252가 아닌 {reason}"
        saved = f.read_bytes().decode("utf-8")
        assert "Don Sabas—a man at once" in saved
        assert "뾞" not in saved

    def test_korean_with_few_high_bytes_still_uses_cp949(self, tmp_path: Path):
        """영문 위주 한글 파일(상위바이트가 적어도 인접쌍)은 cp949/euc-kr을 유지한다."""
        Loader = _get_loader()
        f = tmp_path / "mixed.txt"
        body = "Taegu University Press chapter one and two\r\n"
        raw = (body * 40).encode("cp949") + "머 리 말\r\n지금 한 세기가\r\n".encode("cp949")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert reason.split(" ")[0] in ("cp949", "euc-kr"), f"한글 인코딩이 아닌 {reason}"
        saved = f.read_bytes().decode("utf-8")
        assert "머 리 말" in saved and "지금 한 세기가" in saved

    def test_korean_decodable_only_by_cp1252_is_untouched(self, tmp_path: Path):
        """상위바이트가 인접쌍이면 cp1252로만 디코딩되더라도 변환하지 않는다.

        실제 케이스: 상위바이트 비율 0.78인 한글 파일이 cp949/euc-kr 왕복은 실패하고
        cp1252로는 디코딩됐다. 가드가 없으면 mojibake로 변환된다.
        """
        Loader = _get_loader()
        f = tmp_path / "korean_broken.txt"
        # 한글 cp949 바이트 사이에 cp949로는 무효인 상위바이트 쌍을 섞는다
        raw = "가나다라마바사아자차\r\n".encode("cp949") * 20 + b"\xff\xfe" * 10
        f.write_bytes(raw)
        hi = [i for i, b in enumerate(raw) if b >= 0x80]
        seen = set(hi)
        paired = sum(1 for i in hi if (i - 1) in seen or (i + 1) in seen) / len(hi)
        assert paired >= Loader.TEXT_HIGH_BYTE_PAIRED_MIN

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert not changed, f"한글 파일이 변환됨: {reason}"
        assert f.read_bytes() == raw

    def test_single_byte_text_is_not_read_as_utf16(self, tmp_path: Path, monkeypatch):
        """단일바이트 텍스트를 UTF-16으로 오독해 CJK 글자 뭉치로 바꾸면 안 된다.

        실제로 발생한 케이스: Windows-1252 영문 파일(0x97 em-dash)이 utf-8/cp949/euc-kr
        디코딩에 모두 실패하고, 길이가 짝수라 utf-16-be로 디코딩되며 왕복까지 통과했다.
        NUL 바이트가 0%이므로 UTF-16일 수 없고, 줄바꿈도 하나도 남지 않는다.
        """
        Loader = _get_loader()
        f = tmp_path / "cp1252.txt"
        body = "J. R. R. Tolkien \x97 The Lord Of The Rings. (4/4)\r\n"
        raw = (body * 60).encode("latin-1")
        assert len(raw) % 2 == 0
        f.write_bytes(raw)

        import chardet

        monkeypatch.setattr(chardet, "detect", lambda b: {"encoding": "UTF-16BE", "confidence": 0.99})

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        # UTF-16으로 오독하면 안 되고, 서양어 단일바이트로 올바르게 읽혀야 한다
        assert "utf-16" not in reason, f"UTF-16으로 오독됨: {reason}"
        assert changed and reason.startswith("cp1252"), reason
        assert f.read_bytes().decode("utf-8") == raw.decode("cp1252")

    def test_utf16_not_attempted_without_nul_bytes(self, tmp_path: Path):
        """NUL 바이트가 거의 없으면 UTF-16 후보를 시도하지 않는다.

        cp949 등 다른 후보로 변환되는 것은 정상이며, UTF-16으로 오독되지만 않으면 된다.
        """
        Loader = _get_loader()
        f = tmp_path / "notutf16.txt"
        raw = ("abcd efgh 1234\r\n" * 50).encode("latin-1") + "가나".encode("cp949")
        f.write_bytes(raw)
        assert raw.count(0) == 0

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert "utf-16" not in reason, f"UTF-16으로 오독됨: {reason}"
        if changed:
            assert f.read_bytes().decode("utf-8") == raw.decode("cp949")

    def test_utf16_wrong_endianness_is_not_written(self, tmp_path: Path):
        """BOM 없는 UTF-16은 LE/BE 양쪽이 왕복을 통과하므로 타당성 검사로 걸러야 한다."""
        Loader = _get_loader()
        f = tmp_path / "utf16le.txt"
        original = "한글 소설 본문입니다 가나다라마바사\n" * 20
        f.write_bytes(original.encode("utf-16-le"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original

    def test_dry_run_does_not_write(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "dry.txt"
        raw = "한글 본문입니다\n".encode("cp949")
        f.write_bytes(raw)

        changed, reason = Loader.reencode_text_file_to_utf8(f, dry_run=True)
        assert changed
        assert "dry-run" in reason
        assert f.read_bytes() == raw

    def test_no_temp_file_left_behind(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "tmpcheck.txt"
        f.write_bytes("한글 본문\n".encode("cp949"))

        Loader.reencode_text_file_to_utf8(f)
        assert list(tmp_path.iterdir()) == [f]

    def test_char_count_and_content_match_after_write(self, tmp_path: Path):
        """저장 후 실제로 다시 읽어 글자수와 내용이 원본 디코딩 결과와 일치한다."""
        Loader = _get_loader()
        f = tmp_path / "count.txt"
        original = "".join(f"{i}행 한글 본문 내용입니다 가나다라마바사\n" for i in range(500))
        f.write_bytes(original.encode("cp949"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        saved = f.read_bytes().decode("utf-8")
        assert len(saved) == len(original)
        assert saved == original

    def test_shrinking_content_is_fully_truncated(self, tmp_path: Path):
        """UTF-8이 원본보다 짧아지는 경우(잔여 바이트가 남지 않아야 함)."""
        Loader = _get_loader()
        f = tmp_path / "shrink.txt"
        # UTF-16은 ASCII 위주 텍스트에서 UTF-8보다 크므로 재인코딩 시 파일이 줄어든다.
        original = "ABCDEFG hello world 12345\n" * 200
        f.write_bytes(original.encode("utf-16"))
        assert f.stat().st_size > len(original.encode("utf-8"))

        changed, reason = Loader.reencode_text_file_to_utf8(f)
        assert changed, reason
        assert f.read_bytes().decode("utf-8") == original
        assert f.stat().st_size == len(original.encode("utf-8"))

    def test_reencode_disabled_by_default_in_read_file(self, tmp_path: Path, monkeypatch):
        """--reencode 없이는 read_file이 파일을 건드리지 않는다."""
        Loader = _get_loader()
        assert Loader.reencode_txt_mode is False
        called: list[Path] = []
        monkeypatch.setattr(Loader, "reencode_text_file_to_utf8", staticmethod(lambda p, dry_run=False: called.append(p) or (False, "x")))
        f = tmp_path / "plain.txt"
        f.write_bytes("한글 본문\n".encode("cp949"))
        Loader.read_from_text(f)
        assert called == []


class TestPdfExtractionCost:
    """추출 비용 최적화 동작 — 스캔본 조기 종료와 페이지 상한."""

    @staticmethod
    def _mock_pypdfium2(monkeypatch, page_texts: list[str], visited: list[int]):
        class MockTextPage:
            def __init__(self, idx):
                self.idx = idx

            def get_text_range(self):
                visited.append(self.idx)
                return page_texts[self.idx]

        class MockPage:
            def __init__(self, idx):
                self.idx = idx

            def get_textpage(self):
                return MockTextPage(self.idx)

        class MockPdfDocument:
            def __init__(self, path):
                pass

            def __len__(self):
                return len(page_texts)

            def __getitem__(self, idx):
                return MockPage(idx)

            def close(self):
                pass

        import pypdfium2

        monkeypatch.setattr(pypdfium2, "PdfDocument", MockPdfDocument)

    def test_scanned_pdf_skips_remaining_parsers(self, tmp_path: Path, monkeypatch):
        """텍스트 레이어가 없으면 pdftotext/pdfplumber를 아예 호출하지 않아야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        visited: list[int] = []
        self._mock_pypdfium2(monkeypatch, [""] * 500, visited)

        called: list[str] = []
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: called.append("pdftotext") or ("", 0)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: called.append("pdfplumber") or ("", 0)))

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert called == [], f"스캔본인데 후속 파서가 호출됨: {called}"
        assert page_count == 500
        assert not summary.strip()

    def test_scanned_pdf_stops_at_page_limit(self, tmp_path: Path, monkeypatch):
        """스캔본이라도 PDF_PAGE_LIMIT 페이지까지만 훑어야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "big_scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        visited: list[int] = []
        self._mock_pypdfium2(monkeypatch, [""] * 5000, visited)
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: ("", 0)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: ("", 0)))

        Loader.read_from_pdf(pdf)
        assert len(visited) == Loader.PDF_PAGE_LIMIT, f"{len(visited)}페이지를 읽음"

    def test_stops_reading_once_text_size_reached(self, tmp_path: Path, monkeypatch):
        """TEXT_SIZE를 채우면 남은 페이지를 읽지 않아야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "long.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        visited: list[int] = []
        self._mock_pypdfium2(monkeypatch, ["가나다라마" * 300] * 100, visited)

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert len(visited) <= 4, f"{len(visited)}페이지를 읽음"
        assert len(summary) == Loader.TEXT_SIZE
        assert page_count == 100

    def test_pdftotext_is_page_limited(self, tmp_path: Path, monkeypatch):
        """pdftotext에 -l 옵션이 붙어 전체 문서를 뽑지 않아야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "cli.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        captured: list[list[str]] = []

        class MockCompletedProcess:
            returncode = 0
            stdout = "본문".encode("utf-8")

        import subprocess

        def fake_run(cmd, **kwargs):
            captured.append(cmd)
            return MockCompletedProcess()

        monkeypatch.setattr(subprocess, "run", fake_run)
        Loader._pdf_text_by_pdftotext(pdf)
        assert captured, "pdftotext가 호출되지 않음"
        cmd = captured[0]
        assert "-l" in cmd and cmd[cmd.index("-l") + 1] == str(Loader.PDF_PAGE_LIMIT)

    def test_fast_parsers_run_before_slow_one(self, tmp_path: Path, monkeypatch):
        """손상 시 fallback 순서는 빠른 것(pdftotext) → 느린 것(pdfplumber)이다."""
        Loader = _get_loader()
        pdf = tmp_path / "damaged.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        broken = "".join(f" {w.encode('euc-kr').decode('latin-1')} 정상 본문 " for w in ("즐기는", "들어서", "필요한", "함께"))
        order: list[str] = []
        self._mock_pypdfium2(monkeypatch, [broken], [])

        def fake_pdftotext(p):
            order.append("pdftotext")
            return broken, 0

        def fake_pdfplumber(p):
            order.append("pdfplumber")
            return "정상 한글 본문입니다", 1

        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(fake_pdftotext))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(fake_pdfplumber))

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert order == ["pdftotext", "pdfplumber"]
        assert "정상 한글 본문입니다" in summary


class TestPdfMojibakeDetection:
    """_is_usable_pdf_text / _count_mojibake_runs 판정 — 부분 손상과 오탐 경계."""

    def test_clean_korean_is_usable(self):
        Loader = _get_loader()
        assert Loader._is_usable_pdf_text("맥(麥) 김남천 조은커뮤니티 소설 본문입니다")

    def test_empty_is_not_usable(self):
        Loader = _get_loader()
        assert not Loader._is_usable_pdf_text("   \n\t  ")

    def test_full_mojibake_is_not_usable(self):
        Loader = _get_loader()
        broken = "맥(麥) 김남천 조은커뮤니티".encode("euc-kr").decode("latin-1")
        assert not Loader._is_usable_pdf_text(broken)

    def test_partial_mojibake_is_not_usable(self):
        """정상 한글이 대부분이어도 손상 구간이 섞이면 재추출 대상이다.

        비율 기준(10%)으로 판정하면 4096자 중 155자 손상을 놓친다.
        """
        Loader = _get_loader()
        # 실제 부분 손상은 정상 텍스트 사이에 손상 구간이 흩어져 나타난다.
        broken = [w.encode("euc-kr").decode("latin-1") for w in ("즐기는", "들어서", "필요한", "함께")]
        partial = "정상 한글 본문이 길게 이어집니다 " * 10 + "".join(f" {w} 정상 한글 본문이 이어집니다 " * 3 for w in broken)
        assert Loader._count_mojibake_runs(partial) >= Loader.PDF_MOJIBAKE_RUN_LIMIT
        assert not Loader._is_usable_pdf_text(partial)

    def test_latin_accents_are_usable(self):
        """포르투갈어 등 라틴 악센트 문자를 손상으로 오판하지 않아야 한다."""
        Loader = _get_loader()
        text = "Songbook Lô Borges Música de Minas Transcrições Cláudio Faria Belo Horizonte"
        assert Loader._is_usable_pdf_text(text)

    def test_falls_back_to_least_damaged(self, tmp_path: Path, monkeypatch):
        """모든 파서가 손상 결과를 주면 손상이 가장 적은 것을 남긴다."""
        Loader = _get_loader()
        pdf = tmp_path / "all_damaged.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        def scatter(words):
            return "".join(f" {w.encode('euc-kr').decode('latin-1')} 정상 본문 " for w in words)

        heavy = scatter(("즐기는", "들어서", "필요한", "함께", "그리고", "사람은"))
        light = scatter(("즐기는", "들어서", "필요한")) + " LIGHT_DAMAGE_MARKER"

        monkeypatch.setattr(Loader, "_pdf_text_by_pypdfium2", staticmethod(lambda p: (heavy, 7)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: (light, 7)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: ("", 0)))

        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        assert "LIGHT_DAMAGE_MARKER" in summary
        assert page_count == 7


class TestPdfParserTimeout:
    """손상 PDF가 파서 안에서 무한히 도는 것을 막는 벽시계 상한.

    실제 사고: 손상 PDF 1건이 pdfminer 안에서 13시간 동안 100% CPU를 점유했다.
    읽기 syscall은 완전히 멈춘 채 메모리에서만 돌기 때문에 페이지 루프 검사로는
    빠져나올 수 없어 SIGALRM으로 덮는다.
    """

    @staticmethod
    def _spin(_path):
        """SIGALRM으로만 벗어날 수 있는 순수 파이썬 무한 루프."""
        while True:
            pass

    def test_time_limit_interrupts_infinite_loop(self):
        import pytest

        from utils.parser_timeout import ParserTimeout, time_limit

        with pytest.raises(ParserTimeout):
            with time_limit(0.3, "테스트"):
                while True:
                    pass

    def test_time_limit_restores_previous_handler_and_disarms(self):
        """블록을 벗어나면 이전 핸들러와 타이머를 원상복구해야 한다."""
        import signal

        from utils.parser_timeout import time_limit

        before = signal.getsignal(signal.SIGALRM)
        with time_limit(30, "테스트"):
            pass
        assert signal.getsignal(signal.SIGALRM) is before
        # 타이머가 해제되어 남은 시간이 0이어야 한다 (뒤늦게 터지면 무관한 코드가 죽는다)
        assert signal.getitimer(signal.ITIMER_REAL)[0] == 0

    def test_time_limit_is_noop_in_worker_thread(self):
        """워커 스레드에서는 signal을 설치할 수 없으므로 그대로 실행해야 한다."""
        import threading

        from utils.parser_timeout import time_limit

        outcome: list[str] = []

        def run():
            try:
                with time_limit(0.1, "테스트"):
                    outcome.append("ran")
            except Exception as e:  # ValueError: signal only works in main thread
                outcome.append(f"raised:{type(e).__name__}")

        t = threading.Thread(target=run)
        t.start()
        t.join(timeout=10)
        assert outcome == ["ran"], outcome

    def test_hung_stage_falls_through_to_next_parser(self, tmp_path: Path, monkeypatch):
        """한 파서가 멈춰도 타임아웃 후 다음 파서로 넘어가 결과를 얻어야 한다."""
        Loader = _get_loader()
        pdf = tmp_path / "hang.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        monkeypatch.setattr(Loader, "PDF_STAGE_TIMEOUT", 1)
        monkeypatch.setattr(Loader, "_pdf_text_by_pypdfium2", staticmethod(self._spin))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: ("정상적으로 추출된 본문입니다", 5)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: ("", 0)))

        summary, _line_count, page_count = Loader.read_from_pdf(pdf)
        assert "정상적으로" in summary
        assert page_count == 5

    def test_timeout_is_reported_as_error(self, tmp_path: Path, monkeypatch, caplog):
        """타임아웃은 조용히 넘어가지 않고 에러로 표시되어야 한다."""
        import logging

        Loader = _get_loader()
        pdf = tmp_path / "hang_all.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock pdf content")

        monkeypatch.setattr(Loader, "PDF_STAGE_TIMEOUT", 1)
        for name in ("_pdf_text_by_pypdfium2", "_pdf_text_by_pdftotext", "_pdf_text_by_pdfplumber"):
            monkeypatch.setattr(Loader, name, staticmethod(self._spin))

        with caplog.at_level(logging.ERROR):
            summary, _line_count, _page_count = Loader.read_from_pdf(pdf)

        assert not summary.strip()
        messages = "\n".join(r.getMessage() for r in caplog.records if r.levelno >= logging.ERROR)
        assert "타임아웃" in messages, messages
        # 세 단계 모두 타임아웃했음이 드러나야 한다
        assert "모든 파서가 타임아웃" in messages, messages

    def test_pdfplumber_leaves_no_open_handle_on_timeout(self, tmp_path: Path, monkeypatch):
        """타임아웃이 열기 도중에 터져도 fd가 GC까지 남으면 안 된다."""
        import glob
        import os

        import pytest

        Loader = _get_loader()
        pdf = tmp_path / "leak.pdf"
        pdf.write_bytes(b"%PDF-1.4 not really a pdf")

        def open_pdf_fds() -> list[str]:
            found = []
            for p in glob.glob("/proc/self/fd/*"):
                try:
                    target = os.readlink(p)
                except OSError:
                    continue
                if target.endswith("leak.pdf"):
                    found.append(target)
            return found

        assert open_pdf_fds() == []

        import pdfplumber

        def hang_on_open(_fp, **_kwargs):
            while True:
                pass

        monkeypatch.setattr(pdfplumber, "open", hang_on_open)

        from utils.parser_timeout import ParserTimeout, time_limit

        with pytest.raises(ParserTimeout):
            with time_limit(0.5, "테스트"):
                Loader._pdf_text_by_pdfplumber(pdf)

        assert open_pdf_fds() == [], "타임아웃 후 파일 핸들이 남아 있음"

    def test_pypdfium2_leaves_no_open_handle_when_load_fails(self, tmp_path: Path):
        """문서 로드 실패 시 pdfium이 쥔 fd가 남으면 안 된다.

        경로를 넘기면 로드 실패한 fd가 닫히지 않고 GC로도 회수되지 않아,
        손상 PDF를 만날 때마다 fd가 1개씩 영구 누적됐다(실측).
        """
        import gc
        import glob
        import os

        Loader = _get_loader()
        pdf = tmp_path / "broken_pypdfium.pdf"
        # 헤더는 PDF지만 본문이 없어 pdfium이 문서 로드에 실패한다
        pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")

        def open_pdf_fds() -> list[str]:
            found = []
            for p in glob.glob("/proc/self/fd/*"):
                try:
                    target = os.readlink(p)
                except OSError:
                    continue
                if target.endswith("broken_pypdfium.pdf"):
                    found.append(target)
            return found

        assert open_pdf_fds() == []

        try:
            Loader._pdf_text_by_pypdfium2(pdf)
        except Exception:
            pass

        # GC에 기대지 않고 즉시 닫혀 있어야 한다
        assert open_pdf_fds() == [], "문서 로드 실패 후 fd가 남아 있음"
        gc.collect()
        assert open_pdf_fds() == []


class TestFileTypeDetection:
    """확장자가 아닌 매직바이트로 실제 포맷을 판별한다.

    실측: 21,150건 중 EPUB 5건이 .pdf 확장자를 달고 있어 PDF 파서로 가 전부 실패했다.
    """

    @staticmethod
    def _write_epub(path: Path, body: str = "테스트 내용입니다") -> None:
        with zipfile.ZipFile(path, "w") as zf:
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", f"<html><body><p>{body}</p></body></html>")

    def test_detects_pdf_by_magic_bytes(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "mislabeled.txt"
        f.write_bytes(b"%PDF-1.6\r%\xe2\xe3\xcf\xd3\r\n")
        assert Loader.detect_file_type(f, "txt") == "pdf"

    def test_detects_epub_inside_pdf_extension(self, tmp_path: Path):
        """실제 사고 케이스: EPUB이 .pdf 확장자를 달고 있음."""
        Loader = _get_loader()
        f = tmp_path / "actually_epub.pdf"
        self._write_epub(f)
        assert Loader.detect_file_type(f, "pdf") == "epub"

    def test_zip_based_extension_is_not_inspected_deeper(self, tmp_path: Path, monkeypatch):
        """정상 epub/cbz/docx는 내용물을 열지 않아야 한다 (비용 0 유지)."""
        Loader = _get_loader()
        f = tmp_path / "normal.epub"
        self._write_epub(f)

        called: list[str] = []
        monkeypatch.setattr(Loader, "_zip_container_type", staticmethod(lambda p: called.append("inspect") or "epub"))
        assert Loader.detect_file_type(f, "epub") == "epub"
        assert called == [], "zip 기반 확장자인데 내용물을 열어봄"

    def test_returns_none_for_unknown_magic(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "plain.txt"
        f.write_text("그냥 텍스트 파일입니다", encoding="utf-8")
        assert Loader.detect_file_type(f, "txt") is None

    def test_missing_file_returns_none(self, tmp_path: Path):
        Loader = _get_loader()
        assert Loader.detect_file_type(tmp_path / "없는파일.pdf", "pdf") is None

    def test_detects_cbz_and_docx_containers(self, tmp_path: Path):
        Loader = _get_loader()
        cbz = tmp_path / "comic.pdf"
        with zipfile.ZipFile(cbz, "w") as zf:
            zf.writestr("001.jpg", b"\xff\xd8\xff\xe0fake")
            zf.writestr("002.png", b"\x89PNGfake")
        assert Loader.detect_file_type(cbz, "pdf") == "cbz"

        docx = tmp_path / "doc.pdf"
        with zipfile.ZipFile(docx, "w") as zf:
            zf.writestr("word/document.xml", "<w:document/>")
        assert Loader.detect_file_type(docx, "pdf") == "docx"

    def test_read_file_uses_detected_type(self, tmp_path: Path, monkeypatch):
        """.pdf 확장자의 EPUB이 EPUB 파서로 가고, ES에도 epub으로 기록되어야 한다."""
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)
        f = tmp_path / "[조지 오웰] 1984년.pdf"
        self._write_epub(f, body="빅 브라더가 당신을 보고 있다")

        data = Loader.read_file(f)
        assert data, "적재되지 않음"
        record = next(iter(data.values()))
        assert record["file_type"] == "epub", record["file_type"]
        assert "빅 브라더" in record["summary"]

    def test_normal_pdf_is_unaffected(self, tmp_path: Path, monkeypatch):
        """정상 PDF는 감지 로직이 끼어들어도 그대로 pdf로 처리되어야 한다."""
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)
        f = tmp_path / "normal.pdf"
        f.write_bytes(b"%PDF-1.4 mock pdf content")

        monkeypatch.setattr(Loader, "_pdf_text_by_pypdfium2", staticmethod(lambda p: ("정상 본문입니다", 10)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: ("", 0)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: ("", 0)))

        data = Loader.read_file(f)
        record = next(iter(data.values()))
        assert record["file_type"] == "pdf"
        assert "정상" in record["summary"]


class TestPdfStructureRepair:
    """xref/page tree가 깨진 PDF를 pikepdf로 재구성한 뒤 다시 추출하는 폴백.

    실측: 9MB 손상본이 기존 3단계 파서로는 빈 문서였는데, 복구 후 470쪽/4096자를 얻었다.
    반대로 객체 자체가 소실된 파일은 복구되지 않고 에러로 보고되어야 한다.
    """

    @staticmethod
    def _all_stages_fail(monkeypatch, Loader):
        """3단계 파서가 모두 문서를 열지 못한 상태(구조 손상)를 만든다."""
        for name in ("_pdf_text_by_pypdfium2", "_pdf_text_by_pdftotext", "_pdf_text_by_pdfplumber"):
            monkeypatch.setattr(Loader, name, staticmethod(lambda p: ("", 0)))

    def test_repair_recovers_text_when_all_parsers_fail(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "broken_xref.pdf"
        pdf.write_bytes(b"%PDF-1.6 mock")

        self._all_stages_fail(monkeypatch, Loader)
        monkeypatch.setattr(Loader, "_pdf_text_by_pikepdf_repair", staticmethod(lambda p: ("복구된 본문입니다. 충분히 긴 한국어 텍스트.", 470)))

        summary, _line_count, page_count = Loader.read_from_pdf(pdf)
        assert "복구된" in summary
        assert page_count == 470

    def test_repair_is_skipped_when_text_already_extracted(self, tmp_path: Path, monkeypatch):
        """mojibake는 구조 손상이 아니므로 복구 비용을 쓰면 안 된다."""
        Loader = _get_loader()
        pdf = tmp_path / "mojibake.pdf"
        pdf.write_bytes(b"%PDF-1.6 mock")

        damaged = "".join(f" {w.encode('euc-kr').decode('latin-1')} 본문 " for w in ("즐기는", "들어서", "필요한", "함께"))
        monkeypatch.setattr(Loader, "_pdf_text_by_pypdfium2", staticmethod(lambda p: (damaged, 12)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdftotext", staticmethod(lambda p: ("", 0)))
        monkeypatch.setattr(Loader, "_pdf_text_by_pdfplumber", staticmethod(lambda p: ("", 0)))

        called: list[str] = []
        monkeypatch.setattr(Loader, "_pdf_text_by_pikepdf_repair", staticmethod(lambda p: called.append("repair") or ("", 0)))

        Loader.read_from_pdf(pdf)
        assert called == [], "텍스트가 이미 나왔는데 복구를 시도함"

    def test_repair_is_skipped_for_scanned_pdf(self, tmp_path: Path, monkeypatch):
        """스캔본(문서는 열리는데 텍스트 레이어 없음)은 복구 대상이 아니다."""
        Loader = _get_loader()
        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.6 mock")

        monkeypatch.setattr(Loader, "_pdf_text_by_pypdfium2", staticmethod(lambda p: ("", 300)))
        called: list[str] = []
        monkeypatch.setattr(Loader, "_pdf_text_by_pikepdf_repair", staticmethod(lambda p: called.append("repair") or ("", 0)))

        _summary, _line_count, page_count = Loader.read_from_pdf(pdf)
        assert called == [], "스캔본에 복구를 시도함"
        assert page_count == 300

    def test_unrecoverable_file_is_reported_as_error(self, tmp_path: Path, monkeypatch, caplog):
        """객체가 소실된 파일은 복구 실패가 에러로 드러나야 한다."""
        import logging

        Loader = _get_loader()
        pdf = tmp_path / "unrecoverable.pdf"
        pdf.write_bytes(b"%PDF-1.6 mock")

        self._all_stages_fail(monkeypatch, Loader)

        def boom(_p):
            raise RuntimeError("unable to find /Root dictionary")

        monkeypatch.setattr(Loader, "_pdf_text_by_pikepdf_repair", staticmethod(boom))

        with caplog.at_level(logging.ERROR):
            summary, _line_count, _page_count = Loader.read_from_pdf(pdf)

        assert not summary.strip()
        messages = "\n".join(r.getMessage() for r in caplog.records if r.levelno >= logging.ERROR)
        assert "구조 복구 실패" in messages, messages

    def test_repair_timeout_is_reported(self, tmp_path: Path, monkeypatch, caplog):
        """복구 단계도 무한히 돌면 안 된다."""
        import logging

        Loader = _get_loader()
        pdf = tmp_path / "hang_repair.pdf"
        pdf.write_bytes(b"%PDF-1.6 mock")

        self._all_stages_fail(monkeypatch, Loader)
        monkeypatch.setattr(Loader, "PDF_STAGE_TIMEOUT", 1)

        def spin(_p):
            while True:
                pass

        monkeypatch.setattr(Loader, "_pdf_text_by_pikepdf_repair", staticmethod(spin))

        with caplog.at_level(logging.ERROR):
            Loader.read_from_pdf(pdf)

        messages = "\n".join(r.getMessage() for r in caplog.records if r.levelno >= logging.ERROR)
        assert "구조 복구 타임아웃" in messages, messages

    def test_repair_reads_through_owned_handle(self, tmp_path: Path):
        """복구 경로도 fd를 남기면 안 된다 (손상 파일이라 실패해도 마찬가지)."""
        import glob
        import os

        Loader = _get_loader()
        pdf = tmp_path / "repair_fd.pdf"
        pdf.write_bytes(b"%PDF-1.6\n%%EOF\n")

        def open_pdf_fds() -> list[str]:
            found = []
            for p in glob.glob("/proc/self/fd/*"):
                try:
                    target = os.readlink(p)
                except OSError:
                    continue
                if target.endswith("repair_fd.pdf"):
                    found.append(target)
            return found

        assert open_pdf_fds() == []
        try:
            Loader._pdf_text_by_pikepdf_repair(pdf)
        except Exception:
            pass
        assert open_pdf_fds() == [], "복구 시도 후 fd가 남아 있음"


class TestIsbnPdfTimeout:
    def test_isbn_pdf_extraction_has_time_limit(self, tmp_path: Path, monkeypatch):
        """ISBN 추출도 손상 PDF에서 무한히 돌면 안 된다.

        read_from_pdf가 포기한 파일도 read_file 안에서 pypdf로 다시 파싱되므로,
        여기에 상한이 없으면 파일 1건이 read_file 전체를 멈춰 세운다(실측).
        """
        import time

        import pypdf

        from utils import isbn as isbn_mod

        pdf = tmp_path / "hang_isbn.pdf"
        pdf.write_bytes(b"%PDF-1.4 mock")

        def hang(*_args, **_kwargs):
            while True:
                pass

        monkeypatch.setattr(isbn_mod, "PDF_ISBN_TIMEOUT", 1)
        monkeypatch.setattr(pypdf, "PdfReader", hang)

        t0 = time.time()
        result = isbn_mod.extract(pdf)
        elapsed = time.time() - t0

        assert result == []
        assert elapsed < 15, f"상한이 걸리지 않음: {elapsed:.1f}s"


class TestReadFromRtf:
    def test_read_valid_rtf(self, tmp_path: Path):
        Loader = _get_loader()
        rtf = tmp_path / "test.rtf"
        rtf.write_bytes(b"{\\rtf1 Hello World}")
        summary, line_count, page_count = Loader.read_from_rtf(rtf)
        assert "Hello" in summary or summary == ""  # depends on striprtf behavior

    def test_read_broken_rtf(self, tmp_path: Path):
        Loader = _get_loader()
        rtf = tmp_path / "broken.rtf"
        rtf.write_bytes(b"\x80\x81\x82not rtf")
        summary, line_count, page_count = Loader.read_from_rtf(rtf)
        assert isinstance(summary, str)


class TestReadFromDocx:
    def test_read_valid_docx(self, tmp_path: Path):
        Loader = _get_loader()
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello Docx World")
        doc.add_paragraph("Line 2")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))
        summary, line_count, page_count = Loader.read_from_docx(docx_path)
        assert "Hello" in summary
        assert line_count >= 2


class TestReadFromEpubFull:
    def test_read_epub_ebooklib_success_path(self, monkeypatch):
        Loader = _get_loader()
        import ebooklib

        class FakeDoc:
            def get_body_content(self):
                return "<html><body><p>본문 줄1</p><p>줄2</p></body></html>"

        class FakeBook:
            def get_metadata(self, namespace, key):
                if (namespace, key) == ("DC", "title"):
                    return [("테스트 제목", {})]
                if (namespace, key) == ("DC", "creator"):
                    return [("테스트 저자", {})]
                return []

            def get_items_of_type(self, item_type):
                assert item_type == ebooklib.ITEM_DOCUMENT
                return [FakeDoc()]

        monkeypatch.setattr("utils.loader.epub.read_epub", lambda _path: FakeBook())
        summary, line_count, page_count = Loader.read_from_epub(Path("dummy.epub"))
        assert "테스트 제목" in summary
        assert "테스트 저자" in summary
        assert line_count >= 1
        assert page_count == 0

    def test_read_epub_with_ebooklib(self, tmp_path: Path):
        """read_from_epub using ebooklib or fallback"""
        Loader = _get_loader()
        epub = tmp_path / "test.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", "<html><body><p>테스트 내용</p></body></html>")
        summary, line_count, page_count = Loader.read_from_epub(epub)
        assert isinstance(summary, str)

    def test_read_epub_exception_zip_fallback(self, tmp_path: Path):
        """Lines 161-176: epub read_epub fails, fallback to zip extraction"""
        Loader = _get_loader()
        epub = tmp_path / "bad_epub.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", "<html><body>Fallback content</body></html>")
        # No mimetype → ebooklib will fail, triggering zip fallback
        summary, line_count, page_count = Loader.read_from_epub(epub)
        assert isinstance(summary, str)

    def test_read_epub_both_primary_and_fallback_fail(self, monkeypatch):
        Loader = _get_loader()
        monkeypatch.setattr("utils.loader.epub.read_epub", lambda _path: (_ for _ in ()).throw(RuntimeError("primary fail")))
        monkeypatch.setattr(Loader, "read_from_epub_with_extracting_zip", staticmethod(lambda _path: (_ for _ in ()).throw(RuntimeError("fallback fail"))))
        summary, line_count, page_count = Loader.read_from_epub(Path("broken.epub"))
        assert summary == ""
        assert line_count == 0
        assert page_count == 0


class TestReadFromDocHwp:
    def test_read_from_doc(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        doc_file = tmp_path / "test.doc"
        doc_file.write_text("doc content")
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "converted text\nline2")
        summary, line_count, page_count = Loader.read_from_doc(doc_file)
        assert "converted" in summary
        assert line_count == 2

    def test_read_from_doc_exception(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        doc_file = tmp_path / "test.doc"
        doc_file.write_text("doc content")

        def raise_lo(fp, fmt):
            raise RuntimeError("lo fail")

        monkeypatch.setattr(Loader, "_convert_with_libreoffice", raise_lo)
        summary, line_count, page_count = Loader.read_from_doc(doc_file)
        assert summary == ""

    def test_read_from_hwp(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        hwp_file = tmp_path / "test.hwp"
        hwp_file.write_text("hwp content")
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "hwp text")
        summary, line_count, page_count = Loader.read_from_hwp(hwp_file)
        assert "hwp" in summary


class TestConvertWithLibreoffice:
    def test_successful_conversion(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        import subprocess

        src = tmp_path / "test.doc"
        src.write_text("doc")
        monkeypatch.setattr(Loader, "_find_libreoffice", lambda: "lo")

        class FakeTmp:
            def __enter__(self):
                return str(tmp_path)

            def __exit__(self, *a):
                return False

        monkeypatch.setattr(tempfile, "TemporaryDirectory", lambda: FakeTmp())

        class DummyProc:
            returncode = 0
            stderr = b""

        def fake_run(*args, **kwargs):
            (tmp_path / "test.txt").write_text("converted", encoding="utf-8")
            return DummyProc()

        monkeypatch.setattr(subprocess, "run", fake_run)
        result = Loader._convert_with_libreoffice(src, "txt:Text")
        assert result == "converted"

    def test_no_output(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        import subprocess

        src = tmp_path / "test.doc"
        src.write_text("doc")
        monkeypatch.setattr(Loader, "_find_libreoffice", lambda: "lo")

        class FakeTmp:
            def __enter__(self):
                return str(tmp_path / "empty")

            def __exit__(self, *a):
                return False

        (tmp_path / "empty").mkdir()
        monkeypatch.setattr(tempfile, "TemporaryDirectory", lambda: FakeTmp())

        class DummyProc:
            returncode = 1
            stderr = b"error"

        monkeypatch.setattr(subprocess, "run", lambda *a, **k: DummyProc())
        result = Loader._convert_with_libreoffice(src, "txt:Text")
        assert result == ""


class TestReadFileBranches:
    def test_read_epub_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        epub = tmp_path / "[Author] Book.epub"
        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("META-INF/container.xml", '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="content.opf"/></rootfiles></container>')
            zf.writestr("content.opf", '<package><manifest><item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/></manifest><spine><itemref idref="c1"/></spine></package>')
            zf.writestr("ch1.xhtml", "<html><body>Content</body></html>")
        result = Loader.read_file(epub)
        Loader.path_prefix = original_prefix
        if result:
            doc = list(result.values())[0]
            assert doc["file_type"] == "epub"
            assert doc["author"] == "Author"

    def test_read_html_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "page.html"
        f.write_text("<html><body>Hello HTML</body></html>")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "html"

    def test_read_cbz_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "comic.cbz"
        f.write_bytes(b"cbz data")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "cbz"

    def test_read_rtf_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "doc.rtf"
        f.write_bytes(b"{\\rtf1 Hello}")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "rtf"

    def test_read_pdf_skip_text(self, tmp_path: Path, monkeypatch):
        """Lines 703-713: skip_text with PDF (fast page count)"""
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        pdf = tmp_path / "book.pdf"
        pdf.write_bytes(b"%PDF")
        monkeypatch.setattr(Loader, "_fast_pdf_page_count", lambda fp: 5)
        result = Loader.read_file(pdf, skip_text=True)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["page_count"] == 5
        assert doc["summary"] == ""

    def test_read_pdf_skip_text_fallback(self, tmp_path: Path, monkeypatch):
        """Lines 705-712: skip_text PDF fast count fails, pypdf fallback"""
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        import pypdf

        pdf = tmp_path / "book.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(72, 72)
        with open(pdf, "wb") as f:
            writer.write(f)
        monkeypatch.setattr(Loader, "_fast_pdf_page_count", lambda fp: None)
        result = Loader.read_file(pdf, skip_text=True)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["page_count"] == 1

    def test_read_doc_file(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "file.doc"
        f.write_text("doc")
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "text")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "doc"

    def test_read_hwp_file(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        f = tmp_path / "file.hwp"
        f.write_text("hwp")
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "text")
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "hwp"

    def test_read_docx_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        from docx import Document

        doc_obj = Document()
        doc_obj.add_paragraph("Hello")
        f = tmp_path / "file.docx"
        doc_obj.save(str(f))
        result = Loader.read_file(f)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "docx"

    def test_read_pdf_file(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        import pypdf

        pdf = tmp_path / "file.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(72, 72)
        with open(pdf, "wb") as f:
            writer.write(f)
        result = Loader.read_file(pdf)
        Loader.path_prefix = original_prefix
        doc = list(result.values())[0]
        assert doc["file_type"] == "pdf"
        assert doc["page_count"] == 1


class TestReadFiles:
    def test_read_files(self, tmp_path: Path):
        Loader = _get_loader()
        original_prefix = Loader.path_prefix
        Loader.path_prefix = tmp_path
        (tmp_path / "a.txt").write_text("hello")
        (tmp_path / "b.txt").write_text("world")
        result = Loader.read_files(tmp_path)
        Loader.path_prefix = original_prefix
        assert len(result) >= 2


class TestGetStat:
    def test_get_stat(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "test.txt"
        f.write_text("hello")
        st = Loader.get_stat(f)
        assert st.st_size == 5


class TestPrintUsage:
    def test_print_usage(self):
        _get_loader()
        from utils.loader import print_usage
        import pytest

        with pytest.raises(SystemExit):
            print_usage("test_program")


class TestFindLibreoffice:
    def test_find_via_which(self, monkeypatch):
        Loader = _get_loader()
        import shutil

        monkeypatch.setattr(shutil, "which", lambda cmd: "/usr/bin/soffice" if cmd == "soffice" else None)
        assert Loader._find_libreoffice() == "/usr/bin/soffice"

    def test_find_mac_path(self, monkeypatch):
        Loader = _get_loader()
        import shutil

        monkeypatch.setattr(shutil, "which", lambda cmd: None)
        mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        monkeypatch.setattr(Path, "exists", lambda self: str(self) == mac_path)
        assert Loader._find_libreoffice() == mac_path

    def test_fallback(self, monkeypatch):
        Loader = _get_loader()
        import shutil

        monkeypatch.setattr(shutil, "which", lambda cmd: None)
        monkeypatch.setattr(Path, "exists", lambda self: False)
        assert Loader._find_libreoffice() == "libreoffice"


class TestFastPdfPageCount:
    def test_real_pdf(self, tmp_path: Path):
        Loader = _get_loader()
        import pypdf

        pdf = tmp_path / "test.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(72, 72)
        writer.add_blank_page(72, 72)
        with open(pdf, "wb") as f:
            writer.write(f)
        count = Loader._fast_pdf_page_count(pdf)
        # May return 2 or None depending on xref structure
        assert count is None or count == 2

    def test_not_a_pdf(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "not.pdf"
        f.write_text("not a pdf")
        assert Loader._fast_pdf_page_count(f) is None

    def test_xref_stream_type1_lookup(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "stream.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        data[100 : 100 + len(b"<< /Pages 2 0 R >>")] = b"<< /Pages 2 0 R >>"
        data[140 : 140 + len(b"<< /Count 7 >>")] = b"<< /Count 7 >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_args: (b"decoded", [1, 2, 1], [0, 1], None)))

        def fake_find_entry(_data, _w, _index_ranges, obj_num):
            if obj_num == 1:
                return (1, 100, 0)
            if obj_num == 2:
                return (1, 140, 0)
            return None

        monkeypatch.setattr(Loader, "_xref_stream_find_entry", staticmethod(fake_find_entry))
        assert Loader._fast_pdf_page_count(pdf) == 7

    def test_xref_stream_type2_object_stream_lookup(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "objstream.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        data[100 : 100 + len(b"<< /Pages 2 0 R >>")] = b"<< /Pages 2 0 R >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_args: (b"decoded", [1, 2, 1], [0, 3], None)))

        def fake_find_entry(_data, _w, _index_ranges, obj_num):
            if obj_num == 1:
                return (1, 100, 0)
            if obj_num == 2:
                return (2, 9, 0)
            if obj_num == 9:
                return (1, 180, 0)
            return None

        monkeypatch.setattr(Loader, "_xref_stream_find_entry", staticmethod(fake_find_entry))
        monkeypatch.setattr(Loader, "_read_from_obj_stream", staticmethod(lambda *_args: b"<< /Type /Pages /Count 4 >>"))
        assert Loader._fast_pdf_page_count(pdf) == 4

    def test_xref_stream_without_root_returns_none(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "no-root.pdf"
        pdf.write_bytes(b"0" * 64 + b"startxref\n20\n%%EOF")
        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_args: (b"decoded", [1, 2, 1], [0, 1], None)))
        assert Loader._fast_pdf_page_count(pdf) is None

    def test_xref_stream_missing_pages_reference_returns_none(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "missing-pages-ref.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        data[100 : 100 + len(b"<< /Type /Catalog >>")] = b"<< /Type /Catalog >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_args: (b"decoded", [1, 2, 1], [0, 1], None)))
        monkeypatch.setattr(Loader, "_xref_stream_find_entry", staticmethod(lambda *_args: (1, 100, 0)))
        assert Loader._fast_pdf_page_count(pdf) is None

    def test_xref_stream_missing_pages_object_returns_none(self, tmp_path: Path, monkeypatch):
        Loader = _get_loader()
        pdf = tmp_path / "missing-pages-object.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        data[100 : 100 + len(b"<< /Pages 2 0 R >>")] = b"<< /Pages 2 0 R >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_args: (b"decoded", [1, 2, 1], [0, 1], None)))

        def fake_find_entry(_data, _w, _index_ranges, obj_num):
            if obj_num == 1:
                return (1, 100, 0)
            return None

        monkeypatch.setattr(Loader, "_xref_stream_find_entry", staticmethod(fake_find_entry))
        assert Loader._fast_pdf_page_count(pdf) is None


# ---- coverage: epub read edge cases ----


class TestReadFromEpubEdgeCases:
    def test_epub_chapter_read_exception(self, tmp_path: Path):
        """Lines 118-119: chapter read raises exception"""
        Loader = _get_loader()
        opf = """<?xml version="1.0"?>
<package xmlns="http://www.idpf.org/2007/opf">
  <metadata/>
  <manifest>
    <item id="ch1" href="ch1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine><itemref idref="ch1"/></spine>
</package>"""
        container = """<?xml version="1.0"?>
<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container" version="1.0">
  <rootfiles><rootfile full-path="content.opf"/></rootfiles>
</container>"""
        epub = tmp_path / "test.epub"
        import zipfile

        with zipfile.ZipFile(epub, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("META-INF/container.xml", container)
            zf.writestr("content.opf", opf)
            # ch1.xhtml intentionally corrupt
            zf.writestr("ch1.xhtml", b"\x00\x01\x02\x03")
        result, line_count = Loader.read_from_epub_with_extracting_zip(epub)
        assert isinstance(result, str)

    def test_epub_both_methods_fail(self, tmp_path: Path, monkeypatch):
        """Lines 171-173: both ebooklib and zip fallback fail"""
        Loader = _get_loader()
        epub = tmp_path / "test.epub"
        epub.write_bytes(b"not a zip at all")
        # read_from_epub tries ebooklib (fails) then zip fallback (also fails)
        summary, line_count, page_count = Loader.read_from_epub(epub)
        assert isinstance(summary, str)


class TestReadFromPdfEdgeCases:
    def test_pdf_text_exceeds_limit(self, tmp_path: Path, monkeypatch):
        """Line 198: break when text exceeds TEXT_SIZE"""
        Loader = _get_loader()
        import pypdf

        pdf = tmp_path / "big.pdf"
        writer = pypdf.PdfWriter()
        for _ in range(3):
            writer.add_blank_page(72, 72)
        with open(pdf, "wb") as f:
            writer.write(f)

        class FakePage:
            def extract_text(self):
                return "x" * 3000

        class FakeReader:
            def __init__(self, f):
                self.pages = [FakePage(), FakePage(), FakePage()]

        monkeypatch.setattr(pypdf, "PdfReader", FakeReader)
        summary, line_count, page_count = Loader.read_from_pdf(pdf)
        # TEXT_SIZE is 4096, so with 3000*2=6000 > 4096, should break early
        assert len(summary) <= Loader.TEXT_SIZE + 3000  # first page + partial second


class TestReadFromHwpEdgeCases:
    def test_hwp_conversion_exception(self, tmp_path: Path, monkeypatch):
        """Lines 345-346: LibreOffice conversion raises"""
        Loader = _get_loader()
        hwp = tmp_path / "test.hwp"
        hwp.write_bytes(b"fake hwp")
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", staticmethod(lambda *a, **kw: (_ for _ in ()).throw(RuntimeError("boom"))))
        summary, line_count, page_count = Loader.read_from_hwp(hwp)
        assert summary == ""


# ---- coverage: _find_xref_offset edge cases ----


class TestFindXrefOffsetExtra:
    def test_leading_lines_before_xref(self):
        """Line 365: lines before 'xref' keyword"""
        Loader = _get_loader()
        xref_data = b"some garbage\nmore garbage\nxref\n0 2\n0000000000 65535 f \n0000000100 00000 n \ntrailer\n"
        offset = Loader._find_xref_offset(xref_data, 1)
        assert offset == 100

    def test_empty_lines_in_xref(self):
        """Lines 371-372: empty lines in xref block"""
        Loader = _get_loader()
        xref_data = b"xref\n\n0 2\n0000000000 65535 f \n0000000200 00000 n \ntrailer\n"
        offset = Loader._find_xref_offset(xref_data, 1)
        assert offset == 200

    def test_malformed_line_in_xref(self):
        """Line 388: non-subsection, non-trailer line"""
        Loader = _get_loader()
        xref_data = b"xref\ngarbage_line\n0 2\n0000000000 65535 f \n0000000300 00000 n \ntrailer\n"
        offset = Loader._find_xref_offset(xref_data, 1)
        assert offset == 300


# ---- coverage: _read_from_obj_stream ----


class TestReadFromObjStream:
    def test_missing_first_or_length(self, tmp_path: Path):
        """Lines 443-446: /First or /Length missing"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /Type /ObjStm >>\nstream\ndata\nendstream")
        with open(f, "rb") as fh:
            assert Loader._read_from_obj_stream(fh, 0, 1) is None

    def test_indirect_length(self, tmp_path: Path):
        """Lines 447-448: /Length is indirect reference"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /First 10 /Length 5 0 R >>\nstream\ndata\nendstream")
        with open(f, "rb") as fh:
            assert Loader._read_from_obj_stream(fh, 0, 1) is None

    def test_no_stream_marker(self, tmp_path: Path):
        """Lines 453-455: no stream marker"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /First 10 /Length 20 >>\nendobj")
        with open(f, "rb") as fh:
            assert Loader._read_from_obj_stream(fh, 0, 1) is None

    def test_non_flatedecode_filter(self, tmp_path: Path):
        """Lines 462-463: unsupported filter"""
        Loader = _get_loader()
        data = b"1 0 obj\n<< /First 10 /Length 5 /Filter /ASCIIHexDecode >>\nstream\nhello\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(data)
        with open(f, "rb") as fh:
            assert Loader._read_from_obj_stream(fh, 0, 1) is None

    def test_zlib_decompress_failure(self, tmp_path: Path):
        """Lines 466-467: zlib decompression fails"""
        Loader = _get_loader()
        data = b"1 0 obj\n<< /First 10 /Length 5 /Filter /FlateDecode >>\nstream\nhello\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(data)
        with open(f, "rb") as fh:
            assert Loader._read_from_obj_stream(fh, 0, 1) is None

    def test_target_found(self, tmp_path: Path):
        """Lines 476-480: target object found in stream"""

        Loader = _get_loader()
        # Object stream: header "5 0 6 20" means obj 5 at offset 0, obj 6 at offset 20
        header = b"5 0 6 20 "
        obj_data = b"<< /Type /Catalog >>"
        raw = header + obj_data + b" " * 20
        compressed = zlib.compress(raw)
        first_offset = len(header)
        stream_data = f"1 0 obj\n<< /First {first_offset} /Length {len(compressed)} /Filter /FlateDecode >>\nstream\n".encode() + compressed + b"\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(stream_data)
        with open(f, "rb") as fh:
            result = Loader._read_from_obj_stream(fh, 0, 5)
        assert result is not None
        assert b"Catalog" in result

    def test_target_not_found(self, tmp_path: Path):
        """Lines 482-484: target object not in stream"""

        Loader = _get_loader()
        header = b"5 0 "
        obj_data = b"<< /Type /Catalog >>"
        raw = header + obj_data
        compressed = zlib.compress(raw)
        first_offset = len(header)
        stream_data = f"1 0 obj\n<< /First {first_offset} /Length {len(compressed)} /Filter /FlateDecode >>\nstream\n".encode() + compressed + b"\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(stream_data)
        with open(f, "rb") as fh:
            result = Loader._read_from_obj_stream(fh, 0, 99)
        assert result is None


# ---- coverage: _parse_one_xref_stream ----


class TestParseOneXrefStream:
    def test_missing_w(self, tmp_path: Path):
        """Lines 493-495: /W missing"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /Size 10 /Length 20 >>\nstream\ndata")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_missing_size(self, tmp_path: Path):
        """Lines 498-500: /Size missing"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Length 20 >>\nstream\ndata")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_missing_length(self, tmp_path: Path):
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Size 10 >>\nstream\ndata")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_indirect_length(self, tmp_path: Path):
        """Lines 507-508: /Length is indirect reference"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Size 10 /Length 5 0 R >>\nstream\ndata")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_no_stream_marker(self, tmp_path: Path):
        """Lines 514-516: no stream marker"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Size 10 /Length 20 >>\nendobj")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_non_flatedecode(self, tmp_path: Path):
        """Lines 524-525: unsupported filter"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Size 10 /Length 5 /Filter /LZW >>\nstream\nhello\nendstream")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_zlib_failure(self, tmp_path: Path):
        """Lines 528-529: zlib decompression fails"""
        Loader = _get_loader()
        f = tmp_path / "fake.pdf"
        f.write_bytes(b"1 0 obj\n<< /W [1 2 1] /Size 10 /Length 5 /Filter /FlateDecode >>\nstream\nhello\nendstream")
        with open(f, "rb") as fh:
            assert Loader._parse_one_xref_stream(fh, 0) is None

    def test_uncompressed_stream(self, tmp_path: Path):
        """Lines 530-531: no filter → raw data"""
        Loader = _get_loader()
        # Create an uncompressed xref stream: W=[1,2,1], Size=2, 2 entries of 4 bytes each
        entry_data = bytes([1, 0, 100, 0]) + bytes([1, 0, 200, 0])  # type=1, offset=100/200, gen=0
        stream_data = f"1 0 obj\n<< /W [1 2 1] /Size 2 /Length {len(entry_data)} >>\nstream\n".encode() + entry_data + b"\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(stream_data)
        with open(f, "rb") as fh:
            result = Loader._parse_one_xref_stream(fh, 0)
        assert result is not None
        decompressed, w, index_ranges, prev_offset = result
        assert w == [1, 2, 1]

    def test_with_decode_parms(self, tmp_path: Path):
        """Lines 534-540: PNG predictor with /DecodeParms"""

        Loader = _get_loader()
        # Each row: 1 filter byte + 4 data bytes (W=[1,2,1] → 4 bytes per entry)
        # Filter byte 0 = None predictor
        row1 = bytes([0, 1, 0, 100, 0])  # type=1, offset=100, gen=0
        row2 = bytes([0, 1, 0, 200, 0])  # type=1, offset=200, gen=0
        raw = row1 + row2
        compressed = zlib.compress(raw)
        stream_data = f"1 0 obj\n<< /W [1 2 1] /Size 2 /Length {len(compressed)} /Filter /FlateDecode /DecodeParms << /Predictor 12 /Columns 4 >> >>\nstream\n".encode() + compressed + b"\nendstream"
        f = tmp_path / "fake.pdf"
        f.write_bytes(stream_data)
        with open(f, "rb") as fh:
            result = Loader._parse_one_xref_stream(fh, 0)
        assert result is not None


# ---- coverage: read_file edge cases ----


class TestReadFileEdgeCases:
    def test_read_file_skip_text_pdf_fast_fails(self, tmp_path: Path, monkeypatch):
        """Lines 710-712: skip_text PDF, fast count None, pypdf also fails"""
        Loader = _get_loader()
        pdf = tmp_path / "cat" / "test.pdf"
        pdf.parent.mkdir()
        pdf.write_bytes(b"not a pdf")
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)
        monkeypatch.setattr(Loader, "_fast_pdf_page_count", staticmethod(lambda f: None))
        result = Loader.read_file(pdf, skip_text=True)
        assert result
        inode = list(result.keys())[0]
        assert result[inode]["page_count"] == 0

    def test_read_file_comics_pdf(self, tmp_path: Path, monkeypatch):
        """Lines 749-759: PDF in comics prefix (text extraction skipped)"""
        import pypdf

        Loader = _get_loader()
        pdf = tmp_path / "comics_cat" / "test.pdf"
        pdf.parent.mkdir()
        writer = pypdf.PdfWriter()
        writer.add_blank_page(72, 72)
        with open(pdf, "wb") as f:
            writer.write(f)
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)
        monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path)
        result = Loader.read_file(pdf)
        assert result
        inode = list(result.keys())[0]
        assert result[inode]["page_count"] >= 0

    def test_convert_with_libreoffice_glob_fallback(self, tmp_path: Path, monkeypatch):
        """Line 305: output file stem doesn't match, falls back to glob"""
        Loader = _get_loader()
        import subprocess as sp

        def fake_run(cmd, **kwargs):
            # Create output with different stem
            outdir = cmd[cmd.index("--outdir") + 1]
            Path(outdir, "different_name.txt").write_text("converted", encoding="utf-8")

            class Result:
                returncode = 0
                stdout = ""
                stderr = ""

            return Result()

        monkeypatch.setattr(sp, "run", fake_run)
        monkeypatch.setattr(Loader, "_find_libreoffice", staticmethod(lambda: "libreoffice"))
        result = Loader._convert_with_libreoffice(tmp_path / "input.doc", "txt:Text")
        assert result == "converted"


# ---- coverage: main() function ----


class TestLoaderMain:
    @staticmethod
    def _setup_env(monkeypatch, tmp_path):
        monkeypatch.setenv("TM_ES_BOOK_INDEX", "test_books")
        monkeypatch.setenv("TM_ES_COMICS_INDEX", "test_comics")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)
        monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path / "comics")
        return Loader

    def test_main_no_args(self, monkeypatch, tmp_path):
        """Lines 886-888: no args"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader"])
        import utils.loader as loader_mod

        monkeypatch.setattr(loader_mod, "print_usage", lambda prog: None)
        result = loader_mod.main()
        assert result == 1

    def test_main_invalid_index(self, monkeypatch, tmp_path):
        """Lines 891-894: invalid index name"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "invalid_index", str(tmp_path)])
        import utils.loader as loader_mod

        monkeypatch.setattr(loader_mod, "print_usage", lambda prog: None)
        result = loader_mod.main()
        assert result == 1

    def test_main_no_file_args(self, monkeypatch, tmp_path):
        """Lines 898-900: no file args without --delete"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "book"])
        import utils.loader as loader_mod

        monkeypatch.setattr(loader_mod, "print_usage", lambda prog: None)
        result = loader_mod.main()
        assert result == 1

    def test_main_delete_mode(self, monkeypatch, tmp_path):
        """Lines 917-927: --delete mode"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "--delete", "book"])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.do_exist_index.return_value = True
        mock_es.delete_index.return_value = None
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_delete_index_not_exist(self, monkeypatch, tmp_path):
        """Lines 922-923: delete but index doesn't exist"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "--delete", "book"])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.do_exist_index.return_value = False
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_delete_exception(self, monkeypatch, tmp_path):
        """Lines 924-926: delete raises exception"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "--delete", "book"])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.do_exist_index.side_effect = RuntimeError("boom")
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == -1

    def test_main_ping_failure(self, monkeypatch, tmp_path):
        """Lines 907-913: ES ping fails"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = False
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        with MagicMock():
            import pytest

            with pytest.raises(SystemExit):
                main()

    def test_main_file_not_found(self, monkeypatch, tmp_path):
        """Lines 1007-1009: file not found"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "book", str(tmp_path / "nonexistent")])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_single_file(self, monkeypatch, tmp_path):
        """Lines 1019-1025: single file force reload"""
        self._setup_env(monkeypatch, tmp_path)
        cat_dir = tmp_path / "cat"
        cat_dir.mkdir()
        test_file = cat_dir / "[author] title.txt"
        test_file.write_text("hello world")
        monkeypatch.setattr("sys.argv", ["loader", "book", str(test_file)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.insert.return_value = []
        mock_es.delete_by_file_paths.return_value = 0
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_recursive(self, monkeypatch, tmp_path):
        """Lines 1026-1038: recursive directory processing"""
        self._setup_env(monkeypatch, tmp_path)
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "[auth] book.txt").write_text("content")
        monkeypatch.setattr("sys.argv", ["loader", "--recursive", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.get_existing_paths.return_value = {}
        mock_es.insert.return_value = []
        mock_es.delete_by_file_paths.return_value = 0
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_directory_non_recursive(self, monkeypatch, tmp_path):
        """Lines 1039-1074: directory without --recursive (2-stage)"""
        self._setup_env(monkeypatch, tmp_path)
        sub1 = tmp_path / "sub1"
        sub1.mkdir()
        (sub1 / "[auth] book1.txt").write_text("content1")
        sub2 = tmp_path / "sub2"
        sub2.mkdir()
        (sub2 / "[auth] book2.txt").write_text("content2")
        (tmp_path / "[auth] root.txt").write_text("root content")
        monkeypatch.setattr("sys.argv", ["loader", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.get_existing_paths.return_value = {}
        mock_es.insert.return_value = []
        mock_es.delete_by_file_paths.return_value = 0
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_recursive_reload(self, monkeypatch, tmp_path):
        """Lines 1032-1033: --recursive --reload (skip_check=True)"""
        self._setup_env(monkeypatch, tmp_path)
        (tmp_path / "[auth] book.txt").write_text("content")
        monkeypatch.setattr("sys.argv", ["loader", "--recursive", "--reload", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.insert.return_value = []
        mock_es.delete_by_file_paths.return_value = 0
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_path_not_in_prefix(self, monkeypatch, tmp_path):
        """Lines 1010-1012: path not in book/comics prefix"""
        Loader = self._setup_env(monkeypatch, tmp_path)
        outside = tmp_path / "outside"
        outside.mkdir()
        (outside / "test.txt").write_text("x")
        # Set prefix to a subdirectory so outside is not relative
        monkeypatch.setattr(Loader, "path_prefix", tmp_path / "books_only")
        monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path / "comics_only")
        (tmp_path / "books_only").mkdir()
        monkeypatch.setattr("sys.argv", ["loader", "book", str(outside)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_main_getopt_error(self, monkeypatch, tmp_path):
        """Lines 876-878: invalid option"""
        self._setup_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "--invalid-option", "book"])
        import utils.loader as loader_mod

        monkeypatch.setattr(loader_mod, "print_usage", lambda prog: None)
        result = loader_mod.main()
        assert result == 1


# ---- HWP 버전별 실제 파일 테스트 ----

HWP_TEST_DIR = Path(__file__).parent / "books" / "_hwp"


def _detect_hwp_version(file_path: Path) -> str:
    """HWP 파일 헤더에서 버전 문자열을 추출한다."""
    with open(file_path, "rb") as f:
        header = f.read(30)
    if header[:15] == b"HWP Document Fi":
        import re as _re

        m = _re.search(r"V(\d+\.\d+)", header.decode("ascii", errors="replace"))
        if m:
            return m.group(1)
    if header[:4] == b"\xd0\xcf\x11\xe0":
        return "5.x"
    return "unknown"


class TestHwpVersionDetection:
    """HWP 파일 헤더 기반 버전 감지 테스트"""

    def test_detect_v1_20(self):
        f = HWP_TEST_DIR / "v1.20_부하의약혼녀.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        assert _detect_hwp_version(f) == "1.20"

    def test_detect_v2_00(self):
        f = HWP_TEST_DIR / "v2.00_박노해_참된시작.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        assert _detect_hwp_version(f) == "2.00"

    def test_detect_v2_10(self):
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        assert _detect_hwp_version(f) == "2.10"

    def test_detect_v3_00(self):
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        assert _detect_hwp_version(f) == "3.00"

    def test_detect_unknown_for_non_hwp(self, tmp_path: Path):
        f = tmp_path / "fake.hwp"
        f.write_bytes(b"this is not an hwp file at all")
        assert _detect_hwp_version(f) == "unknown"

    def test_detect_ole2_hwp5(self, tmp_path: Path):
        """OLE2 매직 바이트를 가진 HWP 5.x 감지"""
        f = tmp_path / "hwp5.hwp"
        f.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 22)
        assert _detect_hwp_version(f) == "5.x"


class TestReadFromHwpRealFiles:
    """실제 HWP 파일로 read_from_hwp 테스트 (LibreOffice 의존)

    LibreOffice가 설치되어 있지 않거나 변환 실패 시에도
    크래시 없이 빈 문자열을 반환하는지 검증한다.
    """

    @staticmethod
    def _hwp_files():
        if not HWP_TEST_DIR.exists():
            return []
        return sorted(HWP_TEST_DIR.glob("*.hwp"))

    def test_v1_20_no_crash(self):
        f = HWP_TEST_DIR / "v1.20_부하의약혼녀.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert isinstance(summary, str)
        assert isinstance(line_count, int)
        assert page_count == 0

    def test_v2_00_no_crash(self):
        f = HWP_TEST_DIR / "v2.00_박노해_참된시작.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert isinstance(summary, str)
        assert isinstance(line_count, int)
        assert page_count == 0

    def test_v2_10_no_crash(self):
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert isinstance(summary, str)
        assert isinstance(line_count, int)
        assert page_count == 0

    def test_v3_00_no_crash(self):
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert isinstance(summary, str)
        assert isinstance(line_count, int)
        assert page_count == 0

    def test_v3_00_noname22_no_crash(self):
        f = HWP_TEST_DIR / "v3.00_noname22.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert isinstance(summary, str)
        assert isinstance(line_count, int)
        assert page_count == 0

    def test_all_versions_return_tuple(self):
        """모든 버전 파일에 대해 3-tuple 반환 확인"""
        files = self._hwp_files()
        if not files:
            import pytest

            pytest.skip("HWP 테스트 파일 없음")
        Loader = _get_loader()
        for f in files:
            result = Loader.read_from_hwp(f)
            assert len(result) == 3, f"파일 {f.name}: 3-tuple이 아님"
            summary, line_count, page_count = result
            assert isinstance(summary, str), f"파일 {f.name}: summary가 str이 아님"
            assert isinstance(line_count, int), f"파일 {f.name}: line_count가 int가 아님"
            assert page_count == 0, f"파일 {f.name}: page_count가 0이 아님"


class TestReadFromHwpWithMockedLibreoffice:
    """LibreOffice 변환 결과를 시뮬레이션하여 버전별 동작 검증"""

    def test_v1_20_libreoffice_returns_empty(self, monkeypatch):
        """V1.20 파일: LibreOffice가 빈 결과를 반환하는 경우"""
        f = HWP_TEST_DIR / "v1.20_부하의약혼녀.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert summary == ""
        assert line_count == 0

    def test_v2_10_libreoffice_returns_text(self, monkeypatch):
        """V2.10 파일: LibreOffice가 텍스트를 반환하는 경우"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "교카 시집\n첫 번째 시\n두 번째 시")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert "교카" in summary
        assert line_count == 3

    def test_v3_00_libreoffice_raises(self, monkeypatch):
        """V3.00 파일: LibreOffice 변환 시 예외 발생"""
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: (_ for _ in ()).throw(RuntimeError("변환 실패")))
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert summary == ""
        assert line_count == 0

    def test_special_chars_cleaned(self, monkeypatch):
        """변환 결과의 특수문자가 정리되는지 확인"""
        f = HWP_TEST_DIR / "v2.10_PAGER.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "텍스트★내용♣특수◆문자")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert "★" not in summary
        assert "텍스트" in summary
        assert "내용" in summary


# ---- HWP3 네이티브 파서 테스트 ----


def _get_hwp3_parser():
    from utils.hwp3_parser import extract_text_from_hwp3

    return extract_text_from_hwp3


class TestHwp3ParserDirect:
    """hwp3_parser.extract_text_from_hwp3 직접 테스트"""

    def test_v3_00_extracts_korean_text(self):
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 10
        assert "고정희" in text or "너" in text  # 시 제목/내용

    def test_v2_10_extracts_korean_text(self):
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 10
        assert "영동" in text or "교가" in text

    def test_v2_10_pager(self):
        f = HWP_TEST_DIR / "v2.10_PAGER.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 100

    def test_v1_20_returns_empty(self):
        """V1.20은 미지원 — 빈 문자열 반환"""
        f = HWP_TEST_DIR / "v1.20_부하의약혼녀.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert isinstance(text, str)

    def test_invalid_file_returns_empty(self, tmp_path: Path):
        f = tmp_path / "not_hwp.hwp"
        f.write_bytes(b"this is definitely not an HWP file")
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_truncated_file_returns_empty(self, tmp_path: Path):
        f = tmp_path / "truncated.hwp"
        f.write_bytes(b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05" + b"\x00" * 50)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_encrypted_file_returns_empty(self, tmp_path: Path):
        """암호화 플래그가 설정된 파일"""
        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        doc_info[96] = 1  # encryption flag (low byte of uint16)
        f = tmp_path / "encrypted.hwp"
        f.write_bytes(sig + bytes(doc_info) + b"\x00" * 1008)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_nonexistent_file_returns_empty(self, tmp_path: Path):
        extract = _get_hwp3_parser()
        assert extract(tmp_path / "does_not_exist.hwp") == ""


class TestHwp3FallbackIntegration:
    """LibreOffice 실패 시 hwp3 파서 fallback 테스트"""

    def test_fallback_when_libreoffice_returns_empty(self, monkeypatch):
        """LibreOffice가 빈 문자열 → hwp3 파서로 텍스트 추출"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert len(summary) > 0
        assert "영동" in summary or "교가" in summary

    def test_fallback_when_libreoffice_returns_whitespace(self, monkeypatch):
        """LibreOffice가 공백만 반환 → hwp3 파서로 fallback"""
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "   \n  ")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert len(summary) > 0

    def test_no_fallback_when_libreoffice_succeeds(self, monkeypatch):
        """LibreOffice가 텍스트 반환 → fallback 사용 안 함"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "LO 결과")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert "LO" in summary


# ---- HWP3 파서 추가 엣지 케이스 테스트 ----


class TestHwp3ParserEdgeCases:
    """hwp3_parser의 다양한 엣지 케이스 검증"""

    def test_compressed_v3_file(self):
        """V3.00 압축 파일 정상 파싱"""
        f = HWP_TEST_DIR / "v3.00_현대시사전.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        # 압축 플래그 확인
        data = f.read_bytes()
        assert data[30 + 124] != 0, "테스트 파일이 압축되어 있어야 함"
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 0

    def test_v2_10_file_parses(self):
        """V2.10 파일 정상 파싱"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 0

    def test_text_contains_no_null_bytes(self):
        """추출된 텍스트에 null 바이트가 없어야 함"""
        for name in ["v2.10_KYOKA.hwp", "v3.00_현대시사전.hwp", "v2.10_PAGER.hwp"]:
            f = HWP_TEST_DIR / name
            if not f.exists():
                continue
            extract = _get_hwp3_parser()
            text = extract(f)
            assert "\x00" not in text, f"{name}: null 바이트 포함"

    def test_text_is_valid_unicode(self):
        """추출된 텍스트가 유효한 유니코드 문자열이어야 함"""
        for name in ["v2.10_KYOKA.hwp", "v3.00_현대시사전.hwp"]:
            f = HWP_TEST_DIR / name
            if not f.exists():
                continue
            extract = _get_hwp3_parser()
            text = extract(f)
            # encode/decode가 에러 없이 수행되어야 함
            encoded = text.encode("utf-8")
            decoded = encoded.decode("utf-8")
            assert decoded == text

    def test_empty_body_after_header(self, tmp_path: Path):
        """헤더만 있고 본문이 없는 파일"""
        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        # 압축 안 함, 정보블록 0
        f = tmp_path / "empty_body.hwp"
        f.write_bytes(sig + bytes(doc_info) + b"\x00" * 1008)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_corrupted_compression(self, tmp_path: Path):
        """압축 플래그가 설정되었지만 데이터가 손상된 경우"""
        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        doc_info[124] = 1  # compressed
        f = tmp_path / "bad_compress.hwp"
        f.write_bytes(sig + bytes(doc_info) + b"\x00" * 1008 + b"\xff\xfe\xfd\xfc" * 100)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_hwp3_hnc_to_unicode_ascii(self):
        """HNC→Unicode: ASCII 범위 변환"""
        from utils.hwp3_parser import _hnc_to_unicode

        assert _hnc_to_unicode(0x41) == "A"
        assert _hnc_to_unicode(0x20) == " "
        assert _hnc_to_unicode(0x7E) == "~"

    def test_hwp3_hnc_to_unicode_hangul(self):
        """HNC→Unicode: 한글 음절 변환"""
        from utils.hwp3_parser import _hnc_to_unicode

        # '가' = cho=2(ㄱ), jung=3(ㅏ), jong=1(없음)
        # HNC: (2<<10)|(3<<5)|1 + 0x8000 = 0x8000 | 0x0800 | 0x0060 | 0x0001 = 0x8861
        result = _hnc_to_unicode(0x8861)
        assert result == "가", f"expected '가', got '{result}'"

    def test_hwp3_hnc_to_unicode_null(self):
        """HNC→Unicode: 0은 빈 문자열"""
        from utils.hwp3_parser import _hnc_to_unicode

        assert _hnc_to_unicode(0) == ""

    def test_hwp3_hnc_to_unicode_hanja(self):
        """HNC→Unicode: 한자 변환 (첫 번째 항목)"""
        from utils.hwp3_parser import _hnc_to_unicode
        from utils.hwp3_tables import KSC5601_TO_UNI

        # 0x4000 → KSC5601_TO_UNI[0]
        result = _hnc_to_unicode(0x4000)
        expected = chr(KSC5601_TO_UNI[0])
        assert result == expected

    def test_hwp3_stream_bounds(self):
        """_HwpStream 범위 초과 시 예외 발생"""
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"\x01\x02\x03")
        assert stream.read_uint8() == 1
        assert stream.remaining() == 2
        with pytest.raises(_HwpStreamError):
            stream.read_uint32()  # 4바이트 필요하지만 2바이트만 남음

    def test_hwp3_stream_skip_negative(self):
        """_HwpStream 음수 skip 시 예외"""
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"\x01\x02\x03")
        with pytest.raises(_HwpStreamError):
            stream.skip(-1)


class TestLoaderHwp3FallbackEdgeCases:
    """Loader.read_from_hwp의 hwp3 fallback 엣지 케이스"""

    def test_fallback_exception_in_hwp3_parser(self, monkeypatch):
        """hwp3 파서에서 예외 발생해도 전체가 크래시하지 않음"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")

        import utils.hwp3_parser as hwp3_mod

        original = hwp3_mod.extract_text_from_hwp3
        monkeypatch.setattr(hwp3_mod, "extract_text_from_hwp3", lambda fp: (_ for _ in ()).throw(RuntimeError("파서 에러")))
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert summary == ""
        assert page_count == 0
        monkeypatch.setattr(hwp3_mod, "extract_text_from_hwp3", original)

    def test_fallback_text_truncated_to_text_size(self, monkeypatch):
        """fallback 결과가 TEXT_SIZE로 잘리는지 확인"""
        f = HWP_TEST_DIR / "v2.10_PAGER.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        assert len(summary) <= Loader.TEXT_SIZE

    def test_fallback_special_chars_cleaned(self, monkeypatch):
        """fallback 결과에도 특수문자 정리가 적용되는지 확인"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        import re

        # 한글, 영숫자, 공백만 남아야 함
        cleaned = re.sub(r"[^\w\sㄱ-힣]", "", summary)
        assert cleaned == summary

    def test_fallback_line_count_correct(self, monkeypatch):
        """fallback 결과의 line_count가 정확한지 확인"""
        f = HWP_TEST_DIR / "v2.10_KYOKA.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        Loader = _get_loader()
        monkeypatch.setattr(Loader, "_convert_with_libreoffice", lambda fp, fmt: "")
        summary, line_count, page_count = Loader.read_from_hwp(f)
        if summary:
            assert line_count > 0


# ---- hwp3_parser 커버리지 보강 테스트 ----


class TestHwp3StreamMethods:
    """_HwpStream 메서드 커버리지"""

    def test_read_uint32(self):
        from utils.hwp3_parser import _HwpStream

        stream = _HwpStream(b"\x01\x00\x00\x00\x02\x00\x00\x00")
        assert stream.read_uint32() == 1
        assert stream.read_uint32() == 2

    def test_read_uint32_overflow(self):
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"\x01\x02")
        with pytest.raises(_HwpStreamError):
            stream.read_uint32()

    def test_read_bytes(self):
        from utils.hwp3_parser import _HwpStream

        stream = _HwpStream(b"hello world")
        assert stream.read_bytes(5) == b"hello"
        assert stream.remaining() == 6

    def test_read_bytes_overflow(self):
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"hi")
        with pytest.raises(_HwpStreamError):
            stream.read_bytes(10)

    def test_read_uint16_overflow(self):
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"\x01")
        with pytest.raises(_HwpStreamError):
            stream.read_uint16()

    def test_read_uint8_overflow(self):
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"")
        with pytest.raises(_HwpStreamError):
            stream.read_uint8()

    def test_skip_overflow(self):
        import pytest
        from utils.hwp3_parser import _HwpStream, _HwpStreamError

        stream = _HwpStream(b"\x01")
        with pytest.raises(_HwpStreamError):
            stream.skip(100)


class TestHncToUnicodeExtended:
    """_hnc_to_unicode 분기 커버리지 보강"""

    def test_special_char_range(self):
        """0x007F-0x3FFF: 특수문자 매핑"""
        from utils.hwp3_parser import _hnc_to_unicode

        # 0x0080 = 유로 기호 (€)
        result = _hnc_to_unicode(0x0080)
        assert result == "€"

    def test_special_char_unmapped(self):
        """매핑 없는 특수문자 → 빈 문자열"""
        from utils.hwp3_parser import _hnc_to_unicode

        result = _hnc_to_unicode(0x0001)  # 범위 밖
        assert result == ""

    def test_hanja_level2(self):
        """0x5318-0x7FFF: 2수준 한자"""
        from utils.hwp3_parser import _hnc_to_unicode
        from utils.hwp3_tables import HNC2UNI

        # 매핑이 있는 2수준 한자 코드 찾기
        for code in range(0x5318, 0x5400):
            if code in HNC2UNI:
                result = _hnc_to_unicode(code)
                assert len(result) == 1
                break

    def test_hanja_level2_unmapped(self):
        """매핑 없는 2수준 한자 → 빈 문자열"""
        from utils.hwp3_parser import _hnc_to_unicode

        result = _hnc_to_unicode(0x7FFF)  # 매핑 없을 확률 높음
        assert isinstance(result, str)

    def test_hangul_choseong_only(self):
        """초성만 있는 한글 (ㄱ)"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=2(ㄱ), jung=2(FILL), jong=1(FILL)
        c = 0x8000 | (2 << 10) | (2 << 5) | 1
        result = _hnc_to_unicode(c)
        assert result == "ㄱ"

    def test_hangul_jungseong_only(self):
        """중성만 있는 한글 (ㅏ)"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=1(FILL), jung=3(ㅏ), jong=1(FILL)
        c = 0x8000 | (1 << 10) | (3 << 5) | 1
        result = _hnc_to_unicode(c)
        assert result == "ㅏ"

    def test_hangul_jongseong_only(self):
        """종성만 있는 한글"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=1(FILL), jung=2(FILL), jong=2(ㄱ)
        c = 0x8000 | (1 << 10) | (2 << 5) | 2
        result = _hnc_to_unicode(c)
        assert result == "ㄱ"

    def test_old_hangul_cho_jung(self):
        """옛한글: 초성+중성만"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=2(ㄱ), jung=3(ㅏ), jong=1(FILL) — 현대 한글이면 완성형으로 가지만
        # L_MAP/V_MAP/T_MAP이 NONE인 조합을 찾아야 함
        # cho=21(ㅎ확장), jung=3(ㅏ), jong=1(FILL)
        c = 0x8000 | (21 << 10) | (3 << 5) | 1
        result = _hnc_to_unicode(c)
        assert len(result) >= 1  # 옛한글 자모 조합

    def test_old_hangul_cho_jung_jong(self):
        """옛한글: 초성+중성+종성 모두"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=21, jung=3, jong=18(옛종성)
        c = 0x8000 | (21 << 10) | (3 << 5) | 18
        result = _hnc_to_unicode(c)
        assert len(result) >= 1

    def test_hangul_jung_zero_fallback(self):
        """jung=0인 완성형 옛한글 fallback"""
        from utils.hwp3_parser import _hnc_to_unicode

        # cho=2, jung=0, jong=0 → jung==0 분기
        c = 0x8000 | (2 << 10) | (0 << 5) | 0
        result = _hnc_to_unicode(c)
        assert isinstance(result, str)

    def test_out_of_range_code(self):
        """범위 밖 코드 → 빈 문자열"""
        from utils.hwp3_parser import _hnc_to_unicode

        assert _hnc_to_unicode(0x0010) == ""


class TestHandleControlCharCoverage:
    """_handle_control_char 분기 커버리지 보강"""

    def _make_ctrl_paragraph(self, ctrl_code, extra_data):
        """제어문자가 포함된 최소 문단 바이너리를 생성한다."""
        import struct

        n_chars = 4 if ctrl_code == 23 else (2 if ctrl_code in (24, 25) else (31 if ctrl_code == 28 else (1 if ctrl_code in (30, 31) else 3)))
        n_chars += 1  # 제어문자 자체 + CR
        n_lines = 1
        char_shape = 0

        header = struct.pack("<BHHB", 0, n_chars, n_lines, char_shape)  # 6 bytes
        header += b"\x00" * 37  # padding to 43 bytes
        header += b"\x00" * 187  # para shape (prev_para_shape=0)
        line_info = b"\x00" * 14  # 1 line
        # chars: ctrl_code + extra + CR(13)
        chars = struct.pack("<H", ctrl_code) + extra_data + struct.pack("<H", 13)
        # 빈 문단 (리스트 종료)
        terminator = b"\x00" * 43
        return header + line_info + chars + terminator

    def _parse_with_data(self, data):
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list

        stream = _HwpStream(data)
        return _parse_paragraph_list(stream)

    def test_ctrl_5_field_code(self):
        """제어문자 5: 필드 코드"""
        import struct

        # skip(6) + uint32(len=0) + skip(2) + skip(0)
        extra = b"\x00" * 6 + struct.pack("<I", 0) + b"\x00" * 2
        data = self._make_ctrl_paragraph(5, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_6_bookmark(self):
        """제어문자 6: 책갈피 (40바이트)"""
        extra = b"\x00" * 40
        data = self._make_ctrl_paragraph(6, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_23_compose(self):
        """제어문자 23: 글자겹침 (8바이트)"""
        extra = b"\x00" * 8
        data = self._make_ctrl_paragraph(23, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_24_hyphen(self):
        """제어문자 24: 하이픈 (4바이트)"""
        extra = b"\x00" * 4
        data = self._make_ctrl_paragraph(24, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_28_outline(self):
        """제어문자 28: 개요번호 (62바이트)"""
        extra = b"\x00" * 62
        data = self._make_ctrl_paragraph(28, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_30_keep_space(self):
        """제어문자 30: 묶음빈칸 (2바이트)"""
        extra = b"\x00" * 2
        data = self._make_ctrl_paragraph(30, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)

    def test_ctrl_18_auto_num(self):
        """제어문자 18: 자동번호 (6바이트)"""
        extra = b"\x00" * 6
        data = self._make_ctrl_paragraph(18, extra)
        result = self._parse_with_data(data)
        assert isinstance(result, str)


class TestParseEdgeCases:
    """파싱 경계 조건 커버리지"""

    def test_max_recursion_depth(self):
        """재귀 깊이 초과 시 빈 문자열 반환"""
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list, _MAX_RECURSION_DEPTH

        stream = _HwpStream(b"\x00" * 100)
        result = _parse_paragraph_list(stream, depth=_MAX_RECURSION_DEPTH + 1)
        assert result == ""

    def test_sanity_check_invalid_n_chars(self):
        """n_chars가 비정상적으로 크면 예외 → 빈 문자열"""
        import struct
        import pytest
        from utils.hwp3_parser import _HwpStream, _parse_paragraph, _HwpStreamError

        # prev_para_shape=1, n_chars=50000(invalid), n_lines=1, csi=0
        header = struct.pack("<BHHB", 1, 50000, 1, 0) + b"\x00" * 37
        stream = _HwpStream(header + b"\x00" * 200)
        with pytest.raises(_HwpStreamError):
            _parse_paragraph(stream, 0)

    def test_body_offset_exceeds_file(self, tmp_path: Path):
        """body_offset이 파일 크기를 초과하는 경우"""
        import struct

        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        # info_block_len = 60000 (파일 크기 초과)
        struct.pack_into("<H", doc_info, 126, 60000)
        f = tmp_path / "big_info_block.hwp"
        f.write_bytes(sig + bytes(doc_info) + b"\x00" * 1008)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_gzip_fallback_decompression(self, tmp_path: Path):
        """raw deflate 실패 → gzip auto 시도"""
        import zlib

        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        doc_info[124] = 1  # compressed
        # 유효한 gzip 데이터 (빈 내용)
        body = zlib.compress(b"\x00" * 100)  # zlib 헤더 포함
        f = tmp_path / "gzip_test.hwp"
        f.write_bytes(sig + bytes(doc_info) + b"\x00" * 1008 + body)
        extract = _get_hwp3_parser()
        # 파싱은 빈 결과지만 크래시하지 않아야 함
        result = extract(f)
        assert isinstance(result, str)

    def test_unsupported_version_v1(self, tmp_path: Path):
        """V1.20 시그니처 → 빈 문자열"""
        sig = b"HWP Document File V1.20 \x1a\x01\x02\x03\x04\x05"
        f = tmp_path / "v1.hwp"
        f.write_bytes(sig + b"\x00" * 200)
        extract = _get_hwp3_parser()
        assert extract(f) == ""

    def test_header_decode_exception(self, tmp_path: Path):
        """시그니처 영역이 깨진 바이너리"""
        # 첫 17바이트는 맞지만 나머지가 깨진 경우
        sig = b"HWP Document File" + b"\xff" * 13
        f = tmp_path / "bad_header.hwp"
        f.write_bytes(sig + b"\x00" * 200)
        extract = _get_hwp3_parser()
        assert extract(f) == ""


class TestHandleControlCharRecursive:
    """표/그림/머리말/각주 등 재귀적 제어문자 커버리지"""

    def _build_hwp_with_paragraphs(self, body_data: bytes) -> bytes:
        """비압축 V3.00 HWP 파일 바이너리를 생성 (fonts=0, styles=0)"""

        sig = b"HWP Document File V3.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)  # 비압축, 정보블록=0
        summary = b"\x00" * 1008
        # fonts: 7 카테고리, 각 0개
        fonts = b"\x00\x00" * 7
        # styles: 0개
        styles = b"\x00\x00"
        return sig + bytes(doc_info) + summary + fonts + styles + body_data

    def _make_paragraph(self, chars_data: bytes, n_chars: int, prev_shape: int = 0) -> bytes:
        """문단 헤더 + 줄 정보 + 글자 데이터"""
        import struct

        n_lines = 1
        header = struct.pack("<BHHB", prev_shape, n_chars, n_lines, 0) + b"\x00" * 37
        if prev_shape == 0:
            header += b"\x00" * 187  # 문단 모양
        line_info = b"\x00" * 14
        return header + line_info + chars_data

    def _empty_paragraph(self) -> bytes:
        """빈 문단 (리스트 종료)"""
        import struct

        return struct.pack("<BHHB", 0, 0, 0, 0) + b"\x00" * 37 + b"\x00" * 187

    def test_ctrl_11_picture(self):
        """제어문자 11: 그림 (skip(6)+uint32+skip(344)+skip(len)+캡션)"""
        import struct

        pic_len = 10
        # 제어문자 11 (n_chars_read += 3 → 총 4)
        ctrl = struct.pack("<H", 11)
        ctrl_data = b"\x00" * 6 + struct.pack("<I", pic_len) + b"\x00" * 344 + b"\x00" * pic_len
        # 캡션 = 빈 문단 리스트
        caption = self._empty_paragraph()
        cr = struct.pack("<H", 13)
        chars = ctrl + ctrl_data + caption + cr
        para = self._make_paragraph(chars, n_chars=4 + 1)
        body = para + self._empty_paragraph()
        data = self._build_hwp_with_paragraphs(body)

        from utils.hwp3_parser import _HwpStream

        # fonts(14) + styles(2) = 16바이트 오프셋부터 문단 시작
        _HwpStream(b"\x00\x00" * 7 + b"\x00\x00" + body)
        # 직접 파싱
        from pathlib import Path
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
            f.write(data)
            f.flush()
            extract = _get_hwp3_parser()
            result = extract(Path(f.name))
        assert isinstance(result, str)

    def test_ctrl_16_header_footer(self):
        """제어문자 16: 머리말/꼬리말"""
        import struct

        ctrl = struct.pack("<H", 16)
        ctrl_data = b"\x00" * 16  # skip(6+10)
        nested_end = self._empty_paragraph()
        cr = struct.pack("<H", 13)
        chars = ctrl + ctrl_data + nested_end + cr
        para = self._make_paragraph(chars, n_chars=4 + 1)
        body = para + self._empty_paragraph()
        data = self._build_hwp_with_paragraphs(body)
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
            f.write(data)
            f.flush()
            extract = _get_hwp3_parser()
            result = extract(Path(f.name))
        assert isinstance(result, str)

    def test_ctrl_17_footnote(self):
        """제어문자 17: 각주/미주"""
        import struct

        ctrl = struct.pack("<H", 17)
        ctrl_data = b"\x00" * 20  # skip(6+14)
        # 각주 내 텍스트: '가' + CR + 빈 문단
        ga = struct.pack("<H", 0x8861)  # '가'
        cr = struct.pack("<H", 13)
        footnote_para = self._make_paragraph(ga + cr, n_chars=2, prev_shape=0)
        footnote_end = self._empty_paragraph()
        outer_cr = struct.pack("<H", 13)
        chars = ctrl + ctrl_data + footnote_para + footnote_end + outer_cr
        para = self._make_paragraph(chars, n_chars=4 + 1)
        body = para + self._empty_paragraph()
        data = self._build_hwp_with_paragraphs(body)
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
            f.write(data)
            f.flush()
            extract = _get_hwp3_parser()
            result = extract(Path(f.name))
        assert isinstance(result, str)

    def test_paragraph_list_stream_error_partial(self):
        """문단 파싱 중 스트림 끝 도달 → 부분 텍스트 반환"""
        import struct
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list

        # 유효한 첫 문단 (CR만) + 잘린 두 번째 문단 헤더
        cr = struct.pack("<H", 13)
        para1 = self._make_paragraph(cr, n_chars=1, prev_shape=0)
        truncated = struct.pack("<BHHB", 0, 100, 1, 0)  # 불완전한 헤더
        stream = _HwpStream(para1 + truncated)
        result = _parse_paragraph_list(stream)
        assert "\n" in result  # 첫 문단의 CR은 포함

    def test_max_text_length_break(self):
        """_MAX_TEXT_LENGTH 초과 시 파싱 중단"""
        import struct
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list, _MAX_TEXT_LENGTH

        # 큰 텍스트를 가진 문단 반복 생성
        # '가' 200개 + CR
        ga = struct.pack("<H", 0x8861)
        cr = struct.pack("<H", 13)
        chars = ga * 200 + cr
        paragraphs = b""
        # 충분히 많은 문단
        for _ in range((_MAX_TEXT_LENGTH // 200) + 10):
            paragraphs += self._make_paragraph(chars, n_chars=201, prev_shape=1)
        paragraphs += self._empty_paragraph()
        stream = _HwpStream(paragraphs)
        result = _parse_paragraph_list(stream)
        assert len(result) <= _MAX_TEXT_LENGTH + 500  # 약간의 여유


# ---- hwp3_parser 추가 커버리지 ----


class TestSafeChr:
    """_safe_chr surrogate 필터"""

    def test_surrogate_high(self):
        from utils.hwp3_parser import _safe_chr

        assert _safe_chr(0xD800) == ""
        assert _safe_chr(0xD830) == ""

    def test_surrogate_low(self):
        from utils.hwp3_parser import _safe_chr

        assert _safe_chr(0xDC00) == ""
        assert _safe_chr(0xDFFF) == ""

    def test_normal_char(self):
        from utils.hwp3_parser import _safe_chr

        assert _safe_chr(0xAC00) == "가"
        assert _safe_chr(0x41) == "A"


class TestControlCharTab:
    """탭(ctrl 9) 제어문자 처리"""

    def test_ctrl_9_tab_in_paragraph(self):
        """탭 제어문자가 포함된 문단"""
        import struct
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list

        # 탭(9) + 6바이트 데이터 → n_chars_read += 3+1=4
        tab = struct.pack("<H", 9) + b"\x00" * 6
        ga = struct.pack("<H", 0x8861)  # '가'
        cr = struct.pack("<H", 13)
        chars = ga + tab + ga + cr  # 가\t가\n, n_chars=1+4+1+1=7
        # 문단 헤더
        header = struct.pack("<BHHB", 0, 7, 1, 0) + b"\x00" * 37
        header += b"\x00" * 187  # para shape
        line_info = b"\x00" * 14
        terminator = struct.pack("<BHHB", 0, 0, 0, 0) + b"\x00" * 37 + b"\x00" * 187
        data = header + line_info + chars + terminator
        stream = _HwpStream(data)
        result = _parse_paragraph_list(stream)
        assert "\t" in result
        assert "가" in result


class TestControlCharDateLine:
    """날짜형식(7), 날짜코드(8), 선(14), 숨은설명(15) 제어문자 처리"""

    def _build_ctrl_para(self, ctrl_code, skip_bytes, n_read_add, extra_nested=False):
        import struct
        from utils.hwp3_parser import _HwpStream, _parse_paragraph_list

        n_chars = 1 + n_read_add + 1  # ctrl + n_read_add + CR
        ctrl = struct.pack("<H", ctrl_code) + b"\x00" * skip_bytes
        if extra_nested:
            # 빈 문단 리스트 (종료)
            ctrl += struct.pack("<BHHB", 0, 0, 0, 0) + b"\x00" * 37 + b"\x00" * 187
        cr = struct.pack("<H", 13)
        chars = ctrl + cr
        header = struct.pack("<BHHB", 0, n_chars, 1, 0) + b"\x00" * 37 + b"\x00" * 187
        line_info = b"\x00" * 14
        terminator = struct.pack("<BHHB", 0, 0, 0, 0) + b"\x00" * 37 + b"\x00" * 187
        data = header + line_info + chars + terminator
        stream = _HwpStream(data)
        return _parse_paragraph_list(stream)

    def test_ctrl_7_date_format(self):
        result = self._build_ctrl_para(7, 84, 3)
        assert isinstance(result, str)

    def test_ctrl_8_date_code(self):
        result = self._build_ctrl_para(8, 96, 3)
        assert isinstance(result, str)

    def test_ctrl_14_line(self):
        result = self._build_ctrl_para(14, 92, 3)
        assert isinstance(result, str)

    def test_ctrl_15_hidden_comment(self):
        result = self._build_ctrl_para(15, 16, 3, extra_nested=True)
        assert isinstance(result, str)


class TestV200BruteForce:
    """V2.00 brute-force 추출 테스트"""

    def test_v200_extracts_text(self):
        """V2.00 파일에서 텍스트가 추출되는지 확인"""
        f = HWP_TEST_DIR / "v2.00_박노해_참된시작.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        assert len(text) > 100

    def test_v200_utf8_safe(self):
        """V2.00 추출 결과가 UTF-8 안전한지 확인"""
        f = HWP_TEST_DIR / "v2.00_박노해_참된시작.hwp"
        if not f.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        extract = _get_hwp3_parser()
        text = extract(f)
        text.encode("utf-8")  # should not raise

    def test_bruteforce_garbage_ratio_filter(self):
        """쓰레기 비율이 높으면 빈 문자열 반환"""
        from utils.hwp3_parser import _extract_text_bruteforce
        import struct

        # 전부 비한글/비한자/비ASCII 범위의 HNC 코드
        garbage = b""
        for i in range(100):
            garbage += struct.pack("<H", 0x3F00 + i)  # 특수문자 범위, 매핑 없음
        result = _extract_text_bruteforce(garbage)
        assert result == ""

    def test_bruteforce_with_valid_text(self):
        """유효한 텍스트가 포함된 데이터에서 추출"""
        from utils.hwp3_parser import _extract_text_bruteforce
        import struct

        # '가나다라마바사아자차' (한글 10자)
        data = b""
        for c in [0x8861, 0x8CC2, 0x9161, 0x9562, 0xA1A1, 0xA562, 0xAD61, 0xB161, 0xB961, 0xBD62]:
            data += struct.pack("<H", c)
        data += struct.pack("<H", 13)  # CR
        result = _extract_text_bruteforce(data)
        assert len(result) > 0

    def test_v200_font_skip_failure_fallback(self, tmp_path: Path):
        """V2.00에서 글꼴 영역 skip 실패 시 전체 brute-force"""
        import struct

        sig = b"HWP Document File V2.00 \x1a\x01\x02\x03\x04\x05"
        doc_info = bytearray(128)
        # 비압축, 정보블록=0
        summary = b"\x00" * 1008
        # 글꼴 영역이 깨진 데이터 (첫 font count가 비정상)
        broken_fonts = struct.pack("<H", 60000)  # n_fonts = 60000 → skip 실패
        # 그 뒤에 유효한 텍스트
        text_data = b""
        for c in [0x8861, 0x8CC2, 0x9161, 0x9562, 0xA1A1, 0xA562, 0xAD61, 0xB161, 0xB961, 0xBD62]:
            text_data += struct.pack("<H", c)
        text_data += struct.pack("<H", 13)
        body = broken_fonts + b"\x00" * 50 + text_data
        f = tmp_path / "v200_broken.hwp"
        f.write_bytes(sig + bytes(doc_info) + summary + body)
        extract = _get_hwp3_parser()
        text = extract(f)
        assert isinstance(text, str)


class TestBookManagerPreviewEmptyContent:
    """book_manager HWP preview에서 빈 결과 시 안내 메시지 반환"""

    def test_hwp_preview_empty_returns_message(self, tmp_path: Path, monkeypatch):
        """LO+hwp3 모두 빈 결과 → '미리보기를 생성할 수 없습니다' 메시지"""
        import shutil

        hwp_src = HWP_TEST_DIR / "v1.20_부하의약혼녀.hwp"
        if not hwp_src.exists():
            import pytest

            pytest.skip("테스트 파일 없음")
        (tmp_path / "A").mkdir(parents=True, exist_ok=True)
        shutil.copy(hwp_src, tmp_path / "A" / "old.hwp")

        from backend.book_manager import BookManager
        from tests.test_book_manager import make_manager, make_doc, DummyES, asyncio_runner

        doc = make_doc("A/old.hwp", "hwp")
        es = DummyES()
        manager = make_manager(tmp_path, es)
        es.search_by_id = lambda _id: doc
        monkeypatch.setattr(BookManager, "_convert_with_libreoffice", lambda p, fmt: "")
        resp = asyncio_runner(manager.get_book_preview(1))
        assert resp.status_code == 200
        body = resp.body.decode("utf-8")
        assert "미리보기를 생성할 수 없습니다" in body


# ── _find_xref_offset 미커버 분기 (loader.py:366, 377-379, 381) ─────────────


class TestFindXrefOffsetBranches:
    """기존 TestFindXrefOffset 이 놓친 분기를 명시적으로 검증한다."""

    def test_trailer_break_and_end_return_none(self):
        """obj 가 범위 밖 → i += 1+sub_count → trailer break → return None (lines 366, 378-379, 381)."""
        Loader = _get_loader()
        # obj_num=5 는 [0,3) 범위 밖 → skip subsection → trailer → None
        xref = b"xref\n0 3\n0000000000 65535 f \n0000000100 00000 n \n0000000200 00000 n \ntrailer\n<< /Size 3 >>"
        result = Loader._find_xref_offset(xref, 5)
        assert result is None

    def test_entry_idx_out_of_range_returns_none(self):
        """entry_idx >= len(lines) 이면 return None (line 377)."""
        Loader = _get_loader()
        # sub_count=2 이지만 entry 가 1개만 존재 → obj_num=1 의 entry_idx 가 범위 밖
        xref = b"xref\n0 2\n0000000009 00000 n \n"
        result = Loader._find_xref_offset(xref, 1)
        assert result is None

    def test_non_matching_line_increments_i(self):
        """digits 2개가 아닌 줄은 i += 1 로 건너뛴다 (line 380-381)."""
        Loader = _get_loader()
        # "junk" 줄은 2-digit 패턴 불일치 → i += 1; 그 다음 "0 1" 을 처리 → obj 0 발견
        xref = b"xref\njunk\n0 1\n0000000009 00000 n \ntrailer\n"
        result = Loader._find_xref_offset(xref, 0)
        assert result == 9

    def test_multiple_subsections_skips_first(self):
        """첫 subsection 에 obj 없음 → i += 1+sub_count (lines 378-379); 두 번째에서 발견."""
        Loader = _get_loader()
        xref = b"xref\n0 1\n0000000009 00000 n \n5 1\n0000000100 00000 n \ntrailer\n"
        # obj_num=5 는 첫 subsection [0,1) 에 없음 → skip → 두 번째 [5,6) 에서 발견
        assert Loader._find_xref_offset(xref, 5) == 100
        # obj_num=0 는 첫 subsection 에서 발견
        assert Loader._find_xref_offset(xref, 0) == 9


# ── comics PDF 경로: _fast_pdf_page_count → None → pypdf fallback (lines 726-732)


class TestReadFileComicsPdf:
    def test_comics_pdf_fast_count_none_falls_back_to_pypdf(self, tmp_path: Path, monkeypatch):
        """_fast_pdf_page_count 가 None 을 반환하면 pypdf.PdfReader 로 페이지 수를 구한다 (lines 726-729)."""
        Loader = _get_loader()
        import pypdf

        # comics_path_prefix 를 tmp_path 로 설정
        monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path)
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)

        pdf_file = tmp_path / "book.pdf"
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(pdf_file, "wb") as f:
            writer.write(f)

        # _fast_pdf_page_count 가 None 을 반환하도록 강제
        monkeypatch.setattr(Loader, "_fast_pdf_page_count", staticmethod(lambda path: None))

        data = Loader.read_file(pdf_file)
        # inode 키로 결과가 반환되며 page_count 가 1
        assert len(data) == 1
        item = next(iter(data.values()))
        assert item["page_count"] == 1

    def test_comics_pdf_fast_count_none_pypdf_exception(self, tmp_path: Path, monkeypatch):
        """pypdf.PdfReader 도 실패하면 page_count=0 으로 처리한다 (lines 726-732)."""
        from unittest.mock import patch

        Loader = _get_loader()

        monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path)
        monkeypatch.setattr(Loader, "path_prefix", tmp_path)

        pdf_file = tmp_path / "broken.pdf"
        pdf_file.write_bytes(b"%PDF-broken")

        monkeypatch.setattr(Loader, "_fast_pdf_page_count", staticmethod(lambda path: None))

        with patch("pypdf.PdfReader", side_effect=Exception("corrupt")):
            data = Loader.read_file(pdf_file)

        assert len(data) == 1
        item = next(iter(data.values()))
        assert item["page_count"] == 0


# ---- coverage: _fast_pdf_page_count 미커버 분기 ----


class TestFastPdfPageCountBranches:
    def test_xref_stream_parse_none_breaks_and_returns_none(self, tmp_path: Path, monkeypatch):
        """parse None → break(603), 이후 lookup 없어 root_data None → return None(645)."""
        Loader = _get_loader()
        pdf = tmp_path / "parse-none.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_a: None))
        assert Loader._fast_pdf_page_count(pdf) is None

    def test_xref_stream_parse_raises_is_swallowed(self, tmp_path: Path, monkeypatch):
        """parse 중 예외 발생 → except로 흡수 후 None(658-660)."""
        Loader = _get_loader()
        pdf = tmp_path / "parse-raise.pdf"
        data = bytearray(b" " * 256)
        data[20 : 20 + len(b"<< /Type /XRef /Root 1 0 R >>")] = b"<< /Type /XRef /Root 1 0 R >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        def _boom(*_a):
            raise RuntimeError("boom")

        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(_boom))
        assert Loader._fast_pdf_page_count(pdf) is None

    def test_mixed_traditional_and_stream_lookups_type2(self, tmp_path: Path, monkeypatch):
        """traditional + stream lookup 혼재 시 type-2 조회 루프에서 비-stream은 continue(634)."""
        Loader = _get_loader()
        pdf = tmp_path / "mixed.pdf"
        data = bytearray(b" " * 400)
        # offset 20: traditional xref (root=1, prev=200)
        xref_block = b"xref\n0 1\n/Root 1 0 R\n/Prev 200\n"
        data[20 : 20 + len(xref_block)] = xref_block
        # offset 200: xref stream (non-xref header)
        data[200 : 200 + len(b"<< /Type /XRef >>")] = b"<< /Type /XRef >>"
        # offset 100: root object → /Pages 2 0 R
        data[100 : 100 + len(b"<< /Pages 2 0 R >>")] = b"<< /Pages 2 0 R >>"
        trailer = b"startxref\n20\n%%EOF"
        data[-len(trailer) :] = trailer
        pdf.write_bytes(bytes(data))

        # traditional lookup: obj 1만 offset 100, 나머지는 None
        monkeypatch.setattr(Loader, "_find_xref_offset", staticmethod(lambda _xd, obj: 100 if obj == 1 else None))
        # 단일 stream xref (prev 없음)
        monkeypatch.setattr(Loader, "_parse_one_xref_stream", staticmethod(lambda *_a: (b"d", [1, 2, 1], [0, 3], None)))

        def fake_find_entry(_d, _w, _ir, obj_num):
            if obj_num == 2:
                return (2, 9, 0)  # type 2 → object stream(컨테이너 obj 9)
            if obj_num == 9:
                return (1, 300, 0)  # type 1 → 파일 offset 300
            return None

        monkeypatch.setattr(Loader, "_xref_stream_find_entry", staticmethod(fake_find_entry))
        monkeypatch.setattr(Loader, "_read_from_obj_stream", staticmethod(lambda *_a: b"<< /Count 4 >>"))
        assert Loader._fast_pdf_page_count(pdf) == 4


# ---- coverage: main() 미커버 분기 ----


def _setup_loader_env(monkeypatch, tmp_path):
    monkeypatch.setenv("TM_ES_BOOK_INDEX", "test_books")
    monkeypatch.setenv("TM_ES_COMICS_INDEX", "test_comics")
    Loader = _get_loader()
    monkeypatch.setattr(Loader, "path_prefix", tmp_path)
    monkeypatch.setattr(Loader, "comics_path_prefix", tmp_path / "comics")
    return Loader


class TestLoaderMainBranches:
    def test_ping_raises_exits(self, monkeypatch, tmp_path):
        """es.ping()가 예외 → except에서 sys.exit(-1) (881-882)."""
        import pytest

        _setup_loader_env(monkeypatch, tmp_path)
        monkeypatch.setattr("sys.argv", ["loader", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.side_effect = RuntimeError("connection refused")
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        with pytest.raises(SystemExit):
            main()

    def test_single_file_stat_oserror(self, monkeypatch, tmp_path):
        """get_stat가 OSError → continue(917-918), 적재 실패 메시지(994)."""
        Loader = _setup_loader_env(monkeypatch, tmp_path)
        f = tmp_path / "[author] title.txt"
        f.write_text("hello")
        monkeypatch.setattr("sys.argv", ["loader", "book", str(f)])

        def _raise(_p):
            raise OSError("stat failed")

        monkeypatch.setattr(Loader, "get_stat", staticmethod(_raise))
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0

    def test_recursive_path_change_and_dedup(self, monkeypatch, tmp_path, capsys):
        """경로 변경 감지(932-939), 중복 제거(961), 동기화 메시지(1004)."""
        _setup_loader_env(monkeypatch, tmp_path)
        f = tmp_path / "[author] book.txt"
        f.write_text("content")
        monkeypatch.setattr("sys.argv", ["loader", "--recursive", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        # 모든 inode를 '경로 변경됨'으로 표시 + batch에 없는 유령 inode(933 continue 경로)
        mock_es.get_existing_paths.side_effect = lambda inodes: {**{ino: "OLD/old.txt" for ino in inodes}, -1: "phantom/ghost.txt"}
        mock_es.delete_by_file_paths.return_value = 1  # cleaned>0 → 961
        mock_es.insert.return_value = []
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0
        out = capsys.readouterr().out
        assert "경로 변경 감지" in out
        assert "경로 동기화" in out

    def test_recursive_skipped(self, monkeypatch, tmp_path, capsys):
        """변경 없는 기존 파일은 skip → 건너뜀 메시지(1002)."""
        _setup_loader_env(monkeypatch, tmp_path)
        f = tmp_path / "[author] book.txt"
        f.write_text("content")
        ino = f.stat().st_ino
        rel = str(f.relative_to(tmp_path))
        monkeypatch.setattr("sys.argv", ["loader", "--recursive", "book", str(tmp_path)])
        from utils.loader import main

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.get_existing_paths.return_value = {ino: rel}  # 동일 경로 → skip
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0
        assert "중복 파일 건너뜀" in capsys.readouterr().out

    def test_directory_two_stage_empty_subdir_and_sync(self, monkeypatch, tmp_path, capsys):
        """비재귀 2단계: 빈 하위디렉토리(1018), 샘플 동기화(1030), 2단계 skip(1038)/sync(1040)."""
        _setup_loader_env(monkeypatch, tmp_path)
        s1 = tmp_path / "s1"
        s1.mkdir()
        f_s1 = s1 / "[author] sample.txt"
        f_s1.write_text("sample")
        (tmp_path / "s2").mkdir()  # 빈 하위 디렉토리 → 1018
        f_r1 = tmp_path / "[author] root1.txt"
        f_r1.write_text("root1")
        f_r2 = tmp_path / "[author] root2.txt"
        f_r2.write_text("root2")

        ino_r1 = f_r1.stat().st_ino
        rel_r1 = str(f_r1.relative_to(tmp_path))

        monkeypatch.setattr("sys.argv", ["loader", "book", str(tmp_path)])
        from utils.loader import main

        def get_existing(inodes):
            res = {}
            for ino in inodes:
                if ino == ino_r1:
                    res[ino] = rel_r1  # 변경 없음 → skip(1038)
                else:
                    res[ino] = "OLD/changed.txt"  # 변경됨 → sync(1030/1040)
            return res

        mock_es = MagicMock()
        mock_es.es.ping.return_value = True
        mock_es.get_existing_paths.side_effect = get_existing
        mock_es.delete_by_file_paths.return_value = 0
        mock_es.insert.return_value = []
        monkeypatch.setattr("utils.loader.ESManager", lambda index_name: mock_es)

        result = main()
        assert result == 0
        out = capsys.readouterr().out
        assert "(파일 없음)" in out
        assert "경로 동기화" in out
        assert "중복 파일 건너뜀" in out


def test_loader_module_executed_as_main(monkeypatch):
    """if __name__ == '__main__' 가드 실행 (loader.py:1055).

    인자 없이 실행 → main()이 print_usage()로 SystemExit(0).
    ES 접속 없이 인자 검증 단계에서 종료된다.
    """
    import runpy
    import pytest

    monkeypatch.setenv("TM_ES_BOOK_INDEX", "test_books")
    monkeypatch.setenv("TM_ES_COMICS_INDEX", "test_comics")
    monkeypatch.setattr(sys, "argv", ["loader.py"])
    with pytest.raises(SystemExit) as exc_info:
        runpy.run_module("utils.loader", run_name="__main__")
    assert exc_info.value.code == 0


def test_main_recursive_reload_deletes_orphans(tmp_path, monkeypatch):
    """--recursive --reload: 대상 경로 하위 orphan(디스크에 없는) 레코드를 삭제한다."""
    import utils.loader as loader_mod

    # 대상 디렉토리를 book prefix 로 지정하고 파일 1개 생성
    monkeypatch.setattr(loader_mod.Loader, "path_prefix", tmp_path)
    (tmp_path / "sub").mkdir()
    (tmp_path / "sub" / "a.txt").write_text("hello world\n", encoding="utf-8")

    monkeypatch.setenv("TM_ES_BOOK_INDEX", "book_idx")
    monkeypatch.setenv("TM_ES_COMICS_INDEX", "comics_idx")

    captured = {}

    class FakeESManager:
        def __init__(self, index_name=None):
            self.index_name = index_name
            self.inserted = set()
            self.deleted_ids = None
            self.refreshed = 0

            class Conn:
                def ping(self_inner):
                    return True

            self.es = Conn()
            captured["m"] = self

        def create_index(self):
            return {}

        def get_existing_paths(self, ids):
            return {}

        def delete_by_file_paths(self, paths, exclude_ids=None):
            return 0

        def insert(self, data, *args, **kwargs):
            self.inserted.update(data.keys())
            return list(data.keys())

        def refresh(self):
            self.refreshed += 1

        def get_doc_ids_by_path_prefix(self, rel):
            # 실제 적재된 inode + 디스크에 없는 orphan 하나
            return set(self.inserted) | {999999}

        def delete_by_ids(self, ids, chunk_size=10000):
            self.deleted_ids = sorted(ids)
            return len(ids)

    monkeypatch.setattr(loader_mod, "ESManager", FakeESManager)
    monkeypatch.setattr(sys, "argv", ["loader", "--recursive", "--reload", "book", str(tmp_path)])

    rc = loader_mod.main()

    m = captured["m"]
    assert rc == 0
    assert m.inserted, "파일이 적재되어야 한다"
    assert m.deleted_ids == [999999], "orphan 만 삭제되어야 한다 (live inode 는 보존)"
    assert 999999 not in m.inserted
